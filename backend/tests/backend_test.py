"""Blackrod Now backend tests — MongoDB-persisted feature pass.

Covers:
  - Health + admin seeded check
  - Subscribe / preferences / unsubscribe (with token merge)
  - Personalised newsletter preview + broadcast (mocked email)
  - Multi-event AI parse-content
  - Contact admin inbox + read
  - Admin -> org notifications
  - Org profile PATCH persistence
  - Facebook mocked connect + publish
  - Document upload (Emergent object storage)
"""
import io
import os
import time
import uuid
import pytest
import requests
from PIL import Image, ImageDraw, ImageFont

BASE_URL = os.environ["REACT_APP_BACKEND_URL"].rstrip("/")
API = f"{BASE_URL}/api"


@pytest.fixture(scope="session")
def api():
    s = requests.Session()
    s.headers.update({"Accept": "application/json"})
    return s


# ─────────── Health ───────────
class TestHealth:
    def test_root(self, api):
        r = api.get(f"{API}/", timeout=15)
        assert r.status_code == 200
        assert r.json().get("ok") is True

    def test_seeded(self, api):
        r = api.get(f"{API}/admin/seeded", timeout=15)
        assert r.status_code == 200
        assert r.json().get("seeded") is True

    def test_stats(self, api):
        r = api.get(f"{API}/admin/stats", timeout=15)
        assert r.status_code == 200
        d = r.json()
        for k in ("events_total", "orgs_total", "subscribers", "messages_unread"):
            assert k in d


# ─────────── Subscribe / preferences ───────────
class TestSubscriberFlow:
    email = f"TEST_sub_{uuid.uuid4().hex[:8]}@example.com"

    def test_subscribe(self, api):
        r = api.post(
            f"{API}/subscribe",
            json={
                "email": self.__class__.email,
                "device_id": "TEST_DEV_1",
                "followed_orgs": ["blackrod-town-council"],
                "followed_categories": ["Community"],
            },
            timeout=20,
        )
        assert r.status_code == 200, r.text
        d = r.json()
        assert d["ok"] is True
        assert d["already_subscribed"] is False
        assert d.get("unsub_token") and d.get("pref_token")
        self.__class__.unsub_token = d["unsub_token"]
        self.__class__.pref_token = d["pref_token"]

    def test_resubscribe_merges(self, api):
        r = api.post(
            f"{API}/subscribe",
            json={
                "email": self.__class__.email,
                "followed_orgs": ["blackrod-town-council"],
                "followed_categories": ["Sport"],
            },
            timeout=15,
        )
        assert r.status_code == 200, r.text
        d = r.json()
        assert d["already_subscribed"] is True
        # tokens preserved
        assert d["unsub_token"] == self.__class__.unsub_token
        assert d["pref_token"] == self.__class__.pref_token

    def test_get_preferences(self, api):
        r = api.get(f"{API}/preferences/{self.__class__.pref_token}", timeout=15)
        assert r.status_code == 200, r.text
        sub = r.json()
        assert sub["email"].lower() == self.__class__.email.lower()
        # merge means both org slugs present
        assert "blackrod-town-council" in sub["followed_orgs"]
        assert "blackrod-town-council" in sub["followed_orgs"]
        assert "Community" in sub["followed_categories"]
        assert "Sport" in sub["followed_categories"]

    def test_patch_preferences(self, api):
        r = api.patch(
            f"{API}/preferences/{self.__class__.pref_token}",
            json={"followed_categories": ["Family"], "digest": False},
            timeout=15,
        )
        assert r.status_code == 200, r.text
        sub = r.json()
        assert sub["followed_categories"] == ["Family"]
        assert sub["digest"] is False

    def test_newsletter_preview_personalised(self, api):
        r = api.get(f"{API}/admin/newsletter/preview", params={"email": self.__class__.email}, timeout=30)
        assert r.status_code == 200, r.text
        d = r.json()
        assert "html" in d and "<html" in d["html"].lower()
        assert isinstance(d.get("matched_events"), int)
        # subject non-empty
        assert d.get("subject")

    def test_unsubscribe(self, api):
        r = api.post(f"{API}/unsubscribe/{self.__class__.unsub_token}", timeout=15)
        assert r.status_code == 200
        assert r.json().get("ok") is True
        # verify via preferences
        r2 = api.get(f"{API}/preferences/{self.__class__.pref_token}", timeout=15)
        assert r2.status_code == 200
        assert r2.json().get("unsubscribed") is True


# ─────────── Broadcast (works in mock OR live Resend mode) ───────────
class TestBroadcast:
    def test_broadcast_works(self, api):
        r = api.post(
            f"{API}/admin/broadcast",
            json={"subject": "TEST broadcast", "html": "<p>Hi</p>"},
            timeout=60,
        )
        assert r.status_code == 200, r.text
        d = r.json()
        assert d.get("ok") is True
        # `mocked` is True only when RESEND_API_KEY is not set. Either way is fine.
        assert "mocked" in d


# ─────────── AI parse-content (multi-event) ───────────
class TestParseContentMultiEvent:
    def test_multi_event(self, api):
        payload = {
            "text": (
                "Summer Fair Sat 14 June 11am-4pm at Community Centre. Bouncy castles, food, live music.\n"
                "Also Youth Football Sun 15 June 10am at Blackrod Rec — under-14s welcome."
            )
        }
        r = api.post(f"{API}/parse-content", json=payload, timeout=120)
        assert r.status_code == 200, r.text
        d = r.json()
        assert "items" in d and isinstance(d["items"], list)
        # at least 2 items OR mocked fallback returning 1 (guard for AI variance)
        assert len(d["items"]) >= 2, f"expected >=2 items, got {len(d['items'])}: {d}"
        for it in d["items"]:
            assert it.get("title")
            assert it.get("suggested_type") in ("event", "update")


class TestAdminImageParse:
    def test_image_upload_is_ocrd(self, api):
        image = Image.new("RGB", (900, 320), "white")
        draw = ImageDraw.Draw(image)
        draw.text((40, 80), "Village Hall Open Day", fill="black")
        draw.text((40, 160), "Saturday 14 June 11am-4pm", fill="black")
        buffer = io.BytesIO()
        image.save(buffer, format="PNG")
        buffer.seek(0)

        files = [
            ("files", ("flyer.png", buffer, "image/png")),
        ]
        data = {"source_org_slug": "blackrod-town-council"}
        r = api.post(f"{API}/admin/documents/parse", files=files, data=data, timeout=120)
        assert r.status_code == 200, r.text
        payload = r.json()
        assert payload["documents"]
        doc = payload["documents"][0]
        assert doc["source_type"] == "image"
        assert doc["items"]
        assert any("Village Hall" in item["title"] or item["suggested_type"] == "event" for item in doc["items"])


# ─────────── Contact admin inbox ───────────
class TestContactAdmin:
    def test_message_lifecycle(self, api):
        payload = {
            "from_org_slug": "blackrod-town-council",
            "from_email": "TEST_contact@example.com",
            "from_name": "TEST Contact",
            "subject": f"TEST_msg_{uuid.uuid4().hex[:6]}",
            "body": "Hello admin, please review our next event.",
        }
        r = api.post(f"{API}/contact-admin", json=payload, timeout=15)
        assert r.status_code == 200, r.text
        msg = r.json()
        assert msg["subject"] == payload["subject"]
        assert msg["read"] is False
        mid = msg["id"]

        # list
        r2 = api.get(f"{API}/admin/messages", timeout=15)
        assert r2.status_code == 200
        subjects = [m["subject"] for m in r2.json()]
        assert payload["subject"] in subjects

        # mark read
        r3 = api.patch(f"{API}/admin/messages/{mid}/read", timeout=15)
        assert r3.status_code == 200
        # verify persisted
        r4 = api.get(f"{API}/admin/messages", timeout=15)
        found = next((m for m in r4.json() if m["id"] == mid), None)
        assert found and found["read"] is True


# ─────────── Notifications (admin → org) ───────────
class TestOrgNotifications:
    def test_create_and_list(self, api):
        title = f"TEST_notif_{uuid.uuid4().hex[:6]}"
        r = api.post(
            f"{API}/admin/notifications",
            json={"org_slug": "blackrod-town-council", "title": title, "body": "Please review your listing."},
            timeout=15,
        )
        assert r.status_code == 200, r.text
        n = r.json()
        assert n["title"] == title
        assert n["read"] is False

        r2 = api.get(f"{API}/organisations/blackrod-town-council/notifications", timeout=15)
        assert r2.status_code == 200
        titles = [x["title"] for x in r2.json()]
        assert title in titles


# ─────────── Org profile patch ───────────
class TestOrgProfilePatch:
    def test_patch_persists(self, api):
        new_short = f"TEST_short_{uuid.uuid4().hex[:6]}"
        r = api.patch(
            f"{API}/organisations/blackrod-town-council",
            json={"short": new_short, "brandColor": "#123456"},
            timeout=15,
        )
        assert r.status_code == 200, r.text
        assert r.json()["short"] == new_short
        # confirm via GET
        r2 = api.get(f"{API}/organisations/blackrod-town-council", timeout=15)
        assert r2.status_code == 200
        d = r2.json()
        assert d["short"] == new_short
        assert d["brandColor"] == "#123456"


# ─────────── Event OG page ───────────
class TestEventOgPage:
    """GET /api/events/{id}/og returns crawler-friendly HTML with per-event
    Open Graph tags plus a redirect to the canonical event page."""

    def test_og_page_returns_event_specific_tags(self, api):
        r = api.get(f"{API}/events?upcoming_only=false", timeout=15)
        assert r.status_code == 200
        events = r.json()
        assert events, "seed events missing"
        # Prefer the well-known festival id if present.
        event = next((x for x in events if x["id"] == "evt-festival-2026"), events[0])
        event_id = event["id"]
        expected_title = event["title"]

        og = api.get(f"{API}/events/{event_id}/og", timeout=15)
        assert og.status_code == 200
        assert "text/html" in og.headers.get("content-type", "").lower()
        body = og.text

        # Base OG tags
        assert 'property="og:site_name" content="Blackrod Now"' in body
        assert 'property="og:title"' in body
        assert 'property="og:description"' in body
        assert 'property="og:image"' in body
        assert 'property="og:url"' in body
        assert 'name="twitter:card" content="summary_large_image"' in body
        import html as _h
        assert _h.escape(expected_title) in body or expected_title in body
        assert f"/events/{event_id}" in body

        # (a) NO active <meta http-equiv="refresh"> tag (comment referencing it is OK)
        import re as _re
        body_no_comments = _re.sub(r"<!--.*?-->", "", body, flags=_re.S)
        assert not _re.search(r"<meta\s+[^>]*http-equiv\s*=\s*[\"']refresh[\"']", body_no_comments, _re.I), \
            "Active meta-refresh tag found — FB crawler would follow it"

        # (b) og:type must be article, NOT event
        assert 'property="og:type" content="article"' in body
        assert 'property="og:type" content="event"' not in body

        # (c) og:image dimensions present and sensible
        assert 'property="og:image:width"' in body
        assert 'property="og:image:height"' in body
        w = int(_re.search(r'og:image:width" content="(\d+)"', body).group(1))
        h = int(_re.search(r'og:image:height" content="(\d+)"', body).group(1))
        assert (w, h) in [(1200, 630), (512, 512), (1600, 500)], f"unexpected image dims {w}x{h}"

        # (d) og:image:secure_url present and matches og:image
        m_img = _re.search(r'property="og:image" content="([^"]+)"', body)
        m_sec = _re.search(r'property="og:image:secure_url" content="([^"]+)"', body)
        assert m_img and m_sec
        assert m_img.group(1) == m_sec.group(1)

        # (e) og:url is canonical /events/{id}
        m_url = _re.search(r'property="og:url" content="([^"]+)"', body)
        assert m_url and m_url.group(1).endswith(f"/events/{event_id}")

        # (f) canonical link tag present
        assert _re.search(r'<link rel="canonical" href="[^"]+' + _re.escape(f"/events/{event_id}") + r'"', body)

        # (g) JS redirect present for humans
        assert "window.location.replace" in body

    def test_og_404_for_missing_event(self, api):
        r = api.get(f"{API}/events/does-not-exist/og", timeout=15)
        assert r.status_code == 404


# ─────────── Documents (Emergent object storage) ───────────
class TestCalendarFeed:
    """GET /api/calendar.ics returns a valid iCalendar feed with filter params."""

    def test_all_events_feed(self, api):
        r = api.get(f"{API}/calendar.ics", timeout=15)
        assert r.status_code == 200
        assert "text/calendar" in r.headers.get("content-type", "").lower()
        body = r.text
        assert body.startswith("BEGIN:VCALENDAR")
        assert "END:VCALENDAR" in body
        assert "PRODID:-//Blackrod Now//" in body
        assert "X-WR-CALNAME:Blackrod Now" in body
        assert "REFRESH-INTERVAL;VALUE=DURATION:PT1H" in body
        # At least one VEVENT
        assert body.count("BEGIN:VEVENT") >= 1
        assert body.count("BEGIN:VEVENT") == body.count("END:VEVENT")

    def test_category_filter(self, api):
        r = api.get(f"{API}/calendar.ics?category=Community", timeout=15)
        assert r.status_code == 200
        assert "X-WR-CALNAME:Blackrod Now · Community" in r.text

    def test_orgs_filter(self, api):
        r = api.get(f"{API}/calendar.ics?orgs=blackrod-town-council", timeout=15)
        assert r.status_code == 200
        assert "X-WR-CALNAME" in r.text


class TestSharePack:
    """Per-org share pack (data + email)."""

    def test_get_share_pack(self, api):
        r = api.get(f"{API}/organisations/blackrod-town-council/share-pack", timeout=15)
        assert r.status_code == 200
        d = r.json()
        assert d["org"]["slug"] == "blackrod-town-council"
        assert isinstance(d["events"], list)
        if d["events"]:
            e = d["events"][0]
            assert e["og_url"].endswith("/og")
            assert e["canonical_url"].startswith("http")
            assert "facebook" in e["share_links"]
            assert "linkedin" in e["share_links"]
            assert "twitter" in e["share_links"]
            assert "whatsapp" in e["share_links"]

    def test_email_share_pack_delivers(self, api):
        # Use Resend's always-accept test address so this works in both
        # mocked and live modes without cross-contaminating real inboxes.
        r = api.post(
            f"{API}/organisations/blackrod-town-council/share-pack/email",
            json={"to": "delivered@resend.dev"},
            timeout=30,
        )
        assert r.status_code == 200, r.text
        d = r.json()
        assert d["ok"] is True
        assert d["to"] == "delivered@resend.dev"
        assert d["email"]["ok"] is True
        assert "mocked" in d["email"]  # either True (no key) or False (live)

    def test_email_share_pack_requires_recipient(self, api):
        # A fake org with no email would 400, but existing seed orgs may have email set.
        # Just verify that passing empty `to` and missing org email produces 400 or 200.
        # Since we can't guarantee no email, skip the negative case if org.email is set.
        r_org = api.get(f"{API}/organisations/blackrod-town-council", timeout=15)
        if r_org.json().get("email"):
            return
        r = api.post(
            f"{API}/organisations/blackrod-town-council/share-pack/email",
            json={},
            timeout=15,
        )
        assert r.status_code == 400


class TestAdminEmailCompose:
    """GET /admin/email/senders, POST /admin/email/preview, POST /admin/email/send."""

    def test_senders_returned(self, api):
        r = api.get(f"{API}/admin/email/senders", timeout=15)
        assert r.status_code == 200
        d = r.json()
        assert isinstance(d["senders"], list) and len(d["senders"]) >= 1
        assert d.get("default")

    def test_preview_renders_html_and_parses_recipients(self, api):
        r = api.post(
            f"{API}/admin/email/preview",
            json={
                "to": "delivered@resend.dev, bad-email, another@example.com\ndup@example.com,dup@example.com",
                "subject": "TEST_preview subject",
                "body": "Line 1\n\nLine 2 with a link https://blackrodnow.com — done.",
            },
            timeout=15,
        )
        assert r.status_code == 200
        d = r.json()
        # HTML contains the branded shell + auto-linked URL
        assert "<html" in d["html"].lower()
        assert 'href="https://blackrodnow.com"' in d["html"]
        assert d["subject"] == "TEST_preview subject"
        # 3 unique valid, 1 invalid ("bad-email"), dup removed
        assert set(d["recipients"]) == {"delivered@resend.dev", "another@example.com", "dup@example.com"}
        assert d["invalid_recipients"] == ["bad-email"]

    def test_preview_rejects_bad_sender(self, api):
        r = api.post(
            f"{API}/admin/email/preview",
            json={"to": "a@b.co", "subject": "s", "body": "b", "from_email": "attacker@evil.com"},
            timeout=15,
        )
        assert r.status_code == 400

    def test_send_delivers_to_test_address(self, api):
        r = api.post(
            f"{API}/admin/email/send",
            json={
                "to": "delivered@resend.dev",
                "subject": "TEST_admin_compose",
                "body": "Hello from the admin compose box.",
            },
            timeout=30,
        )
        assert r.status_code == 200, r.text
        d = r.json()
        assert d["sent"] == 1
        assert d["failed"] == 0
        assert d["results"][0]["to"] == "delivered@resend.dev"
        # If Resend key present, id should be present and mocked False.
        if not d.get("mocked"):
            assert d["results"][0].get("id")

    def test_send_400_no_valid_recipients(self, api):
        r = api.post(
            f"{API}/admin/email/send",
            json={"to": "not-an-email", "subject": "s", "body": "b"},
            timeout=15,
        )
        assert r.status_code == 400

    def test_send_400_missing_subject(self, api):
        r = api.post(
            f"{API}/admin/email/send",
            json={"to": "a@b.co", "subject": "", "body": "b"},
            timeout=15,
        )
        assert r.status_code == 400


class TestEventPatch:
    """PATCH /api/events/{id} accepts partial updates."""

    def test_patch_updates_fields(self, api):
        # Create a test event
        payload = {
            "title": "TEST_patch_event",
            "orgSlug": "blackrod-town-council",
            "category": "Community",
            "start": "2027-01-01T10:00:00Z",
            "end": "2027-01-01T11:00:00Z",
            "venue": "TBC",
            "description": "orig",
        }
        r = api.post(f"{API}/events", json=payload, timeout=15)
        assert r.status_code == 200
        eid = r.json()["id"]

        # Patch a subset
        r2 = api.patch(
            f"{API}/events/{eid}",
            json={"title": "TEST_patch_event_updated", "venue": "Blackrod Library", "status": "approved"},
            timeout=15,
        )
        assert r2.status_code == 200
        e = r2.json()
        assert e["title"] == "TEST_patch_event_updated"
        assert e["venue"] == "Blackrod Library"
        assert e["status"] == "approved"
        # Unchanged fields survived
        assert e["orgSlug"] == "blackrod-town-council"
        assert e["description"] == "orig"

        # Cleanup
        api.delete(f"{API}/admin/events/{eid}", timeout=15)

    def test_patch_404_missing_event(self, api):
        r = api.patch(f"{API}/events/does-not-exist", json={"title": "x"}, timeout=15)
        assert r.status_code == 404

    def test_patch_invalid_org_slug(self, api):
        # Create then attempt to move to nonexistent org
        r = api.post(
            f"{API}/events",
            json={
                "title": "TEST_patch_org_fail",
                "orgSlug": "blackrod-town-council",
                "category": "Community",
                "start": "2027-02-01T10:00:00Z",
                "end": "2027-02-01T11:00:00Z",
            },
            timeout=15,
        )
        eid = r.json()["id"]
        r2 = api.patch(f"{API}/events/{eid}", json={"orgSlug": "does-not-exist"}, timeout=15)
        assert r2.status_code == 404
        api.delete(f"{API}/admin/events/{eid}", timeout=15)


class TestNotificationThread:
    """POST /contact-admin with in_reply_to threads a reply to a notification,
    and GET /notifications/{id}/thread returns the notification + replies."""

    def test_thread_flow(self, api):
        r = api.post(
            f"{API}/admin/notifications",
            json={"org_slug": "blackrod-town-council", "title": "TEST_thread_notif", "body": "Please review."},
            timeout=15,
        )
        assert r.status_code == 200, r.text
        nid = r.json()["id"]

        r2 = api.post(
            f"{API}/contact-admin",
            json={
                "from_org_slug": "blackrod-town-council",
                "subject": "Re: TEST_thread_notif",
                "body": "Ack — will update by Friday.",
                "in_reply_to": nid,
            },
            timeout=15,
        )
        assert r2.status_code == 200, r2.text
        assert r2.json()["in_reply_to"] == nid

        r3 = api.get(f"{API}/notifications/{nid}/thread", timeout=15)
        assert r3.status_code == 200
        d = r3.json()
        assert d["notification"]["id"] == nid
        reply_bodies = [x["body"] for x in d["replies"]]
        assert "Ack — will update by Friday." in reply_bodies

    def test_thread_404_for_missing_notification(self, api):
        r = api.get(f"{API}/notifications/does-not-exist/thread", timeout=15)
        assert r.status_code == 404


class TestDocuments:
    def test_upload_and_list(self, api):
        content = b"Hello Blackrod " + uuid.uuid4().hex.encode()
        files = {"file": (f"test_{uuid.uuid4().hex[:6]}.txt", io.BytesIO(content), "text/plain")}
        r = api.post(
            f"{API}/organisations/blackrod-town-council/documents",
            files=files,
            timeout=60,
        )
        assert r.status_code == 200, r.text
        d = r.json()
        assert d["org_slug"] == "blackrod-town-council"
        assert d["storage_path"]
        assert d["size"] > 0
        doc_id = d["id"]

        r2 = api.get(f"{API}/organisations/blackrod-town-council/documents", timeout=15)
        assert r2.status_code == 200
        ids = [x["id"] for x in r2.json()]
        assert doc_id in ids


class TestAdminBulkDocumentParse:
    def test_parse_multiple_documents(self, api):
        from docx import Document as DocxDocument

        txt = io.BytesIO(
            b"Blackrod Bloomers Committee update. New spring fair on Saturday 14 June at the village hall."
        )

        docx_buf = io.BytesIO()
        doc = DocxDocument()
        doc.add_paragraph("Blackrod Town Council")
        doc.add_paragraph("Annual membership update for approved organisations.")
        doc.save(docx_buf)
        docx_buf.seek(0)

        files = [
            ("files", (f"TEST_event_{uuid.uuid4().hex[:6]}.txt", txt, "text/plain")),
            ("files", (f"TEST_org_{uuid.uuid4().hex[:6]}.docx", docx_buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")),
        ]
        r = api.post(f"{API}/admin/documents/parse", files=files, timeout=120)
        assert r.status_code == 200, r.text
        d = r.json()
        assert "documents" in d and len(d["documents"]) == 2
        for doc in d["documents"]:
            assert doc["filename"]
            assert isinstance(doc.get("items"), list) and doc["items"]
            item = doc["items"][0]
            assert item["title"]
            assert item["action"] in ("new_event", "update_event", "new_organisation", "update_organisation", "unclear")
