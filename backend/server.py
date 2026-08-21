"""Blackrod Now backend — MongoDB-persisted community platform."""
from fastapi import FastAPI, APIRouter, HTTPException, UploadFile, File, Form, Response, Header, Query, Request
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
from pymongo import ReturnDocument
from pathlib import Path
from pydantic import BaseModel, EmailStr, Field, ConfigDict
from typing import List, Optional, Literal, Any, Dict
from datetime import datetime, timezone, timedelta, date as date_cls, time as time_cls
from collections import Counter, defaultdict
from difflib import SequenceMatcher
import os
import re
import json
import uuid
import asyncio
import logging
import hashlib
import hmac
import secrets
import csv
import html
import base64
import zipfile
import xml.etree.ElementTree as ET
import threading
from functools import lru_cache
import requests
from io import BytesIO
from PIL import Image, ImageOps
from urllib.parse import urlparse

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / ".env")

# ─────────── Config ───────────
MONGO_URL = os.environ["MONGO_URL"]
DB_NAME = os.environ["DB_NAME"]
EMERGENT_LLM_KEY = os.environ.get("EMERGENT_LLM_KEY")

# ─────────── LLM library preload (event-loop protection) ───────────
_llm_ready = threading.Event()
_llm_preload_error: Optional[str] = None


def _preload_llm_libs():
    global _llm_preload_error
    try:
        import emergentintegrations.llm.chat  # noqa: F401  — heavy import (litellm)
        logging.getLogger(__name__).info("LLM libraries preloaded")
    except Exception as exc:
        _llm_preload_error = str(exc)
        logging.getLogger(__name__).warning("LLM library preload failed: %s", exc)
    finally:
        _llm_ready.set()


async def _ensure_llm_loaded(timeout: float = 240) -> bool:
    """Await LLM library availability WITHOUT blocking the event loop.
    Returns False when the libraries can't be used (callers fall back)."""
    if not EMERGENT_LLM_KEY:
        return False
    if not _llm_ready.is_set():
        await asyncio.to_thread(_llm_ready.wait, timeout)
    return _llm_ready.is_set() and _llm_preload_error is None
RESEND_API_KEY = os.environ.get("RESEND_API_KEY")
SENDER_EMAIL = os.environ.get("SENDER_EMAIL", "onboarding@resend.dev")
SENDER_NAME = os.environ.get("SENDER_NAME", "Blackrod Now")
ADMIN_SENDER_EMAILS = [
    s.strip() for s in os.environ.get("ADMIN_SENDER_EMAILS", SENDER_EMAIL).split(",") if s.strip()
]
ADMIN_INBOX_EMAILS = [
    s.strip() for s in os.environ.get("ADMIN_INBOX_EMAILS", ",".join(ADMIN_SENDER_EMAILS)).split(",") if s.strip()
]
RESEND_WEBHOOK_SIGNING_SECRET = os.environ.get("RESEND_WEBHOOK_SIGNING_SECRET", "")
RESEND_WEBHOOK_TOLERANCE_SECONDS = int(os.environ.get("RESEND_WEBHOOK_TOLERANCE_SECONDS", "300"))
WEB_WIZARD_TO_EMAIL = os.environ.get("WEB_WIZARD_TO_EMAIL", SENDER_EMAIL)
WEB_WIZARD_BCC_EMAIL = os.environ.get("WEB_WIZARD_BCC_EMAIL", "benwordsworth@aol.com")
APP_NAME = os.environ.get("APP_NAME", "blackrodnow")
STORAGE_URL = "https://integrations.emergentagent.com/objstore/api/v1/storage"
PUBLIC_URL = os.environ.get("PUBLIC_URL", "https://blackrodnow.local")
ADMIN_LAUNCH_CODE = os.environ.get("ADMIN_LAUNCH_CODE", "Blackr0dN0w!&")
ORG_DEFAULT_PASSWORD = os.environ.get("ORG_DEFAULT_PASSWORD", "Organisat10n!&")
ORG_AUTH_SECRET = os.environ.get("ORG_AUTH_SECRET", f"{APP_NAME}:{ADMIN_LAUNCH_CODE}")
ORG_AUTH_TOKEN_TTL_SECONDS = int(os.environ.get("ORG_AUTH_TOKEN_TTL_SECONDS", "43200"))
ORG_CLAIM_VERIFY_TTL_MINUTES = int(os.environ.get("ORG_CLAIM_VERIFY_TTL_MINUTES", "20"))
ORG_CLAIM_VERIFY_MAX_ATTEMPTS = int(os.environ.get("ORG_CLAIM_VERIFY_MAX_ATTEMPTS", "5"))
ORG_CLAIM_MAX_STARTS_PER_HOUR = int(os.environ.get("ORG_CLAIM_MAX_STARTS_PER_HOUR", "5"))

# ─────────── DB ───────────
client = AsyncIOMotorClient(MONGO_URL)
db = client[DB_NAME]

# ─────────── App ───────────
app = FastAPI(title="Blackrod Now API")
api = APIRouter(prefix="/api")
logger = logging.getLogger("blackrodnow")


def new_id() -> str:
    return str(uuid.uuid4())


def new_token() -> str:
    return uuid.uuid4().hex


def _require_org_write_access(slug: str, org_auth: Optional[str] = None, admin_code: Optional[str] = None) -> str:
    """Auth guard. Accepts (a) X-Admin-Code matching ADMIN_LAUNCH_CODE, (b) an
    admin/org JWT via X-Org-Auth `Bearer <jwt>` header, or (c) legacy org tokens
    (current permissive rollout). Returns the effective role ('admin' if the
    caller proved admin credentials, else 'org').
    Real JWT flow now uses `/api/auth/admin/login`; this shim keeps existing
    role-switcher UI working while we migrate."""
    # Admin: launch code header
    if admin_code and hmac.compare_digest(admin_code, ADMIN_LAUNCH_CODE):
        return "admin"
    # Admin/Org: bearer JWT via X-Org-Auth (frontend already sends this)
    if org_auth and org_auth.lower().startswith("bearer "):
        token = org_auth[7:].strip()
        try:
            from auth import decode_token
            payload = decode_token(token)
            if payload.get("role") == "admin":
                return "admin"
            if payload.get("role") == "org_member":
                token_slug = str(payload.get("org_slug") or "")
                member_status = str(payload.get("member_status") or "active")
                if token_slug == slug and member_status == "active":
                    return "org"
        except Exception:
            pass
    return "org"


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()
logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")


def _decode_resend_signing_secret(secret: str) -> bytes:
    token = (secret or "").strip()
    if token.startswith("whsec_"):
        token = token[6:]
    padded = token + ("=" * (-len(token) % 4))
    try:
        return base64.b64decode(padded)
    except Exception:
        return token.encode("utf-8")


def _parse_svix_v1_signatures(header_value: str) -> List[str]:
    parts = [part.strip() for part in (header_value or "").split() if part.strip()]
    signatures: List[str] = []
    for part in parts:
        prefix, sep, value = part.partition(",")
        if sep and prefix == "v1" and value:
            signatures.append(value.strip())
    return signatures


def _verify_resend_webhook_signature(
    body: bytes,
    svix_id: str,
    svix_timestamp: str,
    svix_signature: str,
) -> tuple[bool, str]:
    if not RESEND_WEBHOOK_SIGNING_SECRET:
        return False, "webhook signing secret is not configured"
    if not svix_id or not svix_timestamp or not svix_signature:
        return False, "missing svix headers"

    try:
        ts = int(svix_timestamp)
    except Exception:
        return False, "invalid svix timestamp"

    now_ts = int(datetime.now(timezone.utc).timestamp())
    if abs(now_ts - ts) > RESEND_WEBHOOK_TOLERANCE_SECONDS:
        return False, "svix timestamp outside tolerance"

    signing_key = _decode_resend_signing_secret(RESEND_WEBHOOK_SIGNING_SECRET)
    signed_payload = f"{svix_id}.{svix_timestamp}.{body.decode('utf-8')}".encode("utf-8")
    expected = base64.b64encode(hmac.new(signing_key, signed_payload, hashlib.sha256).digest()).decode("utf-8")
    provided = _parse_svix_v1_signatures(svix_signature)
    if not provided:
        return False, "no v1 signature provided"

    for candidate in provided:
        if hmac.compare_digest(candidate, expected):
            return True, "ok"
    return False, "signature mismatch"


def _extract_email_address(value: Any) -> str:
    if isinstance(value, dict):
        for key in ("email", "address"):
            if value.get(key):
                return str(value.get(key)).strip().lower()
        return ""
    if isinstance(value, list):
        for item in value:
            addr = _extract_email_address(item)
            if addr:
                return addr
        return ""
    text = str(value or "").strip()
    if not text:
        return ""
    match = re.search(r"<([^>]+)>", text)
    if match:
        return match.group(1).strip().lower()
    return text.strip().lower()


def _extract_sender_name(value: Any) -> str:
    if isinstance(value, dict):
        return str(value.get("name") or "").strip()
    text = str(value or "").strip()
    if not text:
        return ""
    match = re.match(r'"?([^"<]+)"?\s*<[^>]+>$', text)
    if match:
        return match.group(1).strip()
    return ""


def _extract_header_value(headers: Any, key: str) -> str:
    wanted = (key or "").strip().lower()
    if not wanted:
        return ""
    if isinstance(headers, dict):
        for k, v in headers.items():
            if str(k).strip().lower() == wanted:
                return str(v or "").strip()
    if isinstance(headers, list):
        for item in headers:
            if not isinstance(item, dict):
                continue
            name = str(item.get("name") or item.get("key") or "").strip().lower()
            if name == wanted:
                return str(item.get("value") or "").strip()
    return ""


def _html_to_text(source_html: str) -> str:
    html_text = str(source_html or "")
    if not html_text:
        return ""
    no_script = re.sub(r"<script[\s\S]*?</script>", "", html_text, flags=re.IGNORECASE)
    no_style = re.sub(r"<style[\s\S]*?</style>", "", no_script, flags=re.IGNORECASE)
    body = re.sub(r"<[^>]+>", " ", no_style)
    body = html.unescape(body)
    return re.sub(r"\s+", " ", body).strip()

# ─────────── Optional integrations ───────────
_storage_key: Optional[str] = None
def init_storage() -> Optional[str]:
    global _storage_key
    if _storage_key or not EMERGENT_LLM_KEY:
        return _storage_key
    try:
        r = requests.post(f"{STORAGE_URL}/init", json={"emergent_key": EMERGENT_LLM_KEY}, timeout=15)
        r.raise_for_status()
        _storage_key = r.json()["storage_key"]
        logger.info("Object storage initialized")
    except Exception as e:
        logger.warning("Storage init failed (uploads will fail gracefully): %s", e)
    return _storage_key


def put_object(path: str, data: bytes, content_type: str) -> dict:
    key = init_storage()
    if not key:
        raise HTTPException(500, "Storage not initialized")
    r = requests.put(
        f"{STORAGE_URL}/objects/{path}",
        headers={"X-Storage-Key": key, "Content-Type": content_type},
        data=data,
        timeout=120,
    )
    r.raise_for_status()
    return r.json()


def get_object(path: str):
    key = init_storage()
    if not key:
        raise HTTPException(500, "Storage not initialized")
    r = requests.get(f"{STORAGE_URL}/objects/{path}", headers={"X-Storage-Key": key}, timeout=60)
    r.raise_for_status()
    return r.content, r.headers.get("Content-Type", "application/octet-stream")


def resend_send(
    to_email: str,
    subject: str,
    html: str,
    from_email: Optional[str] = None,
    from_name: Optional[str] = None,
    bcc: Optional[List[str]] = None,
    attachments: Optional[List[Dict[str, Any]]] = None,
) -> Dict[str, Any]:
    """Send via Resend. Returns dict with ok flag + provider id or mock note.

    Optional `from_email` overrides the default sender (must be a
    domain-verified address, e.g. an entry in ADMIN_SENDER_EMAILS).
    """
    if not RESEND_API_KEY:
        logger.info(
            "[MOCK EMAIL] to=%s bcc=%s attachments=%s subject=%s (RESEND_API_KEY not set)",
            to_email,
            bcc or [],
            [a.get("filename") for a in attachments or []],
            subject,
        )
        return {
            "ok": True,
            "mocked": True,
            "to": to_email,
            "bcc": bcc or [],
            "subject": subject,
            "attachments": [{"filename": a.get("filename"), "content_type": a.get("content_type"), "size": a.get("size")} for a in attachments or []],
        }
    try:
        import resend  # lazy
        resend.api_key = RESEND_API_KEY
        sender_addr = from_email or SENDER_EMAIL
        sender_name = from_name if from_name is not None else SENDER_NAME
        # RFC 5322 style: `"Blackrod Now" <blackrodnow@…>` — most clients show
        # the display name instead of the raw address.
        from_field = f'"{sender_name}" <{sender_addr}>' if sender_name else sender_addr
        payload: Dict[str, Any] = {"from": from_field, "to": [to_email], "subject": subject, "html": html}
        if bcc:
            payload["bcc"] = bcc
        if attachments:
            payload["attachments"] = [
                {
                    "filename": a["filename"],
                    "content": a["content"],
                    "content_type": a.get("content_type", "application/octet-stream"),
                }
                for a in attachments
            ]
        result = resend.Emails.send(payload)
        return {
            "ok": True,
            "mocked": False,
            "id": result.get("id"),
            "to": to_email,
            "bcc": bcc or [],
            "attachments": [{"filename": a.get("filename"), "content_type": a.get("content_type"), "size": a.get("size")} for a in attachments or []],
        }
    except Exception as e:
        logger.exception("resend_send failed: %s", e)
        return {
            "ok": False,
            "mocked": False,
            "error": str(e),
            "to": to_email,
            "bcc": bcc or [],
            "attachments": [{"filename": a.get("filename"), "content_type": a.get("content_type"), "size": a.get("size")} for a in attachments or []],
        }


# ─────────── Models ───────────
AnalyticsKind = Literal[
    "org_view",
    "event_view",
    "share_click",
    "broadcast_preview",
    "share_pack_email",
    "newsletter_send",
    "broadcast_send",
    "admin_email_send",
    "volunteer_contact",
]


class AnalyticsEvent(BaseModel):
    id: str = Field(default_factory=new_id)
    kind: AnalyticsKind
    entity_type: Optional[Literal["org", "event", "site"]] = None
    entity_id: Optional[str] = None
    org_slug: Optional[str] = None
    platform: Optional[str] = None
    device_id: Optional[str] = None
    count: int = 1
    created_at: str = Field(default_factory=now_iso)


class AnalyticsTrackReq(BaseModel):
    kind: AnalyticsKind
    entity_type: Optional[Literal["org", "event"]] = None
    entity_id: Optional[str] = None
    org_slug: Optional[str] = None
    platform: Optional[str] = None
    device_id: Optional[str] = None


def _window_start(days: int) -> str:
    return (datetime.now(timezone.utc) - timedelta(days=days)).isoformat()


def _day_key(value: str) -> str:
    try:
        return datetime.fromisoformat(value.replace("Z", "+00:00")).date().isoformat()
    except Exception:
        return value[:10]


async def _record_analytics_event(
    kind: str,
    entity_type: Optional[str] = None,
    entity_id: Optional[str] = None,
    org_slug: Optional[str] = None,
    platform: Optional[str] = None,
    count: int = 1,
    device_id: Optional[str] = None,
) -> None:
    await db.analytics_events.insert_one(
        AnalyticsEvent(
            kind=kind,
            entity_type=entity_type,
            entity_id=entity_id,
            org_slug=org_slug,
            platform=platform,
            count=max(1, int(count or 1)),
            device_id=device_id,
        ).model_dump()
    )


async def _follow_counters() -> Dict[str, Any]:
    device_follows = await db.follows.find({}, {"_id": 0, "orgs": 1, "categories": 1}).to_list(10000)
    subscriber_follows = await db.subscribers.find(
        {"unsubscribed": False}, {"_id": 0, "followed_orgs": 1, "followed_categories": 1}
    ).to_list(10000)

    org_followers = Counter()
    category_followers = Counter()
    device_org_followers = Counter()
    subscriber_org_followers = Counter()
    active_follow_devices = 0
    active_follow_subscribers = 0

    for doc in device_follows:
        orgs = {slug for slug in (doc.get("orgs") or []) if slug}
        categories = {name for name in (doc.get("categories") or []) if name}
        if orgs or categories:
            active_follow_devices += 1
        for slug in orgs:
            org_followers[slug] += 1
            device_org_followers[slug] += 1
        for name in categories:
            category_followers[name] += 1

    for doc in subscriber_follows:
        orgs = {slug for slug in (doc.get("followed_orgs") or []) if slug}
        categories = {name for name in (doc.get("followed_categories") or []) if name}
        if orgs or categories:
            active_follow_subscribers += 1
        for slug in orgs:
            org_followers[slug] += 1
            subscriber_org_followers[slug] += 1
        for name in categories:
            category_followers[name] += 1

    return {
        "org_followers": org_followers,
        "category_followers": category_followers,
        "device_org_followers": device_org_followers,
        "subscriber_org_followers": subscriber_org_followers,
        "active_follow_devices": active_follow_devices,
        "active_follow_subscribers": active_follow_subscribers,
        "org_follow_links": sum(org_followers.values()),
        "category_follow_links": sum(category_followers.values()),
    }


def _summarize_analytics_events(rows: List[Dict[str, Any]], event_lookup: Dict[str, Dict[str, Any]]) -> Dict[str, Any]:
    page_views_by_org = Counter()
    event_views_by_org = Counter()
    share_clicks_by_org = Counter()
    event_views_by_event = Counter()
    shares_by_event = Counter()
    platform_counts = Counter()
    counts_by_kind = Counter()
    trend = defaultdict(lambda: {"org_views": 0, "event_views": 0, "shares": 0})

    for row in rows:
        kind = row.get("kind")
        count = int(row.get("count") or 1)
        org_slug = row.get("org_slug")
        entity_id = row.get("entity_id")
        day = _day_key(row.get("created_at") or now_iso())
        counts_by_kind[kind] += count

        if kind == "org_view":
            if org_slug:
                page_views_by_org[org_slug] += count
            trend[day]["org_views"] += count
        elif kind == "event_view":
            if entity_id:
                event_views_by_event[entity_id] += count
            if org_slug:
                event_views_by_org[org_slug] += count
            trend[day]["event_views"] += count
        elif kind == "share_click":
            if entity_id:
                shares_by_event[entity_id] += count
            if org_slug:
                share_clicks_by_org[org_slug] += count
            platform_counts[row.get("platform") or "unknown"] += count
            trend[day]["shares"] += count

    top_events = []
    ranked_event_ids = set(event_views_by_event.keys()) | set(shares_by_event.keys())
    for event_id in ranked_event_ids:
        event = event_lookup.get(event_id) or {}
        top_events.append(
            {
                "id": event_id,
                "title": event.get("title") or "Untitled event",
                "org_slug": event.get("orgSlug") or event.get("org_slug") or "",
                "start": event.get("start"),
                "views": event_views_by_event.get(event_id, 0),
                "shares": shares_by_event.get(event_id, 0),
                "score": event_views_by_event.get(event_id, 0) + shares_by_event.get(event_id, 0) * 2,
            }
        )
    top_events.sort(key=lambda item: (-item["score"], -item["views"], item["title"]))

    return {
        "counts_by_kind": counts_by_kind,
        "page_views_by_org": page_views_by_org,
        "event_views_by_org": event_views_by_org,
        "share_clicks_by_org": share_clicks_by_org,
        "event_views_by_event": event_views_by_event,
        "shares_by_event": shares_by_event,
        "platform_counts": platform_counts,
        "trend": [{"date": day, **trend[day]} for day in sorted(trend.keys())],
        "top_events": top_events,
    }


async def _build_site_analytics(days: int = 30) -> Dict[str, Any]:
    window_start = _window_start(days)
    approved_events = await db.events.find({"status": "approved"}, {"_id": 0}).to_list(5000)
    approved_orgs = await db.orgs.find({"status": {"$ne": "pending"}}, {"_id": 0}).to_list(3000)
    recent_events = await db.analytics_events.find({"created_at": {"$gte": window_start}}, {"_id": 0}).to_list(20000)
    follow_summary = await _follow_counters()
    digest_subscribers = await db.subscribers.count_documents({"unsubscribed": False, "digest": True})
    total_feed_posts = await db.feed.count_documents({})
    docs_total = await db.documents.count_documents({"deleted": {"$ne": True}})

    now = datetime.now(timezone.utc).isoformat()
    upcoming_events = [event for event in approved_events if (event.get("start") or "") >= now]
    upcoming_by_org = Counter(event.get("orgSlug") or "" for event in upcoming_events if event.get("orgSlug"))
    event_lookup = {event.get("id"): event for event in approved_events if event.get("id")}
    summary = _summarize_analytics_events(recent_events, event_lookup)

    top_orgs = []
    for org in approved_orgs:
        slug = org.get("slug") or ""
        top_orgs.append(
            {
                "slug": slug,
                "name": org.get("name") or slug,
                "page_views": summary["page_views_by_org"].get(slug, 0),
                "event_views": summary["event_views_by_org"].get(slug, 0),
                "share_clicks": summary["share_clicks_by_org"].get(slug, 0),
                "followers": follow_summary["org_followers"].get(slug, 0),
                "upcoming_events": upcoming_by_org.get(slug, 0),
                "score": summary["page_views_by_org"].get(slug, 0)
                + summary["event_views_by_org"].get(slug, 0)
                + summary["share_clicks_by_org"].get(slug, 0) * 2
                + follow_summary["org_followers"].get(slug, 0),
            }
        )
    top_orgs.sort(key=lambda item: (-item["score"], item["name"]))

    orgs_with_upcoming = sum(1 for count in upcoming_by_org.values() if count > 0)
    pending_content_total = await db.events.count_documents({"status": "pending"}) + await db.orgs.count_documents({"status": "pending"})
    approved_event_count = len(approved_events)
    total_views = summary["counts_by_kind"].get("event_view", 0)
    total_shares = summary["counts_by_kind"].get("share_click", 0)

    return {
        "window_days": days,
        "overview": {
            "approved_orgs": len(approved_orgs),
            "approved_events": approved_event_count,
            "upcoming_events": len(upcoming_events),
            "orgs_with_upcoming_events": orgs_with_upcoming,
            "digest_subscribers": digest_subscribers,
            "org_follow_links": follow_summary["org_follow_links"],
            "category_follow_links": follow_summary["category_follow_links"],
            "active_follow_devices": follow_summary["active_follow_devices"],
            "active_follow_subscribers": follow_summary["active_follow_subscribers"],
            "feed_posts": total_feed_posts,
            "documents": docs_total,
        },
        "engagement": {
            "org_views_30d": summary["counts_by_kind"].get("org_view", 0),
            "event_views_30d": total_views,
            "share_clicks_30d": total_shares,
            "share_pack_emails_30d": summary["counts_by_kind"].get("share_pack_email", 0),
            "newsletter_sends_30d": summary["counts_by_kind"].get("newsletter_send", 0),
            "broadcast_sends_30d": summary["counts_by_kind"].get("broadcast_send", 0),
            "admin_email_sends_30d": summary["counts_by_kind"].get("admin_email_send", 0),
        },
        "health": {
            "avg_event_views_30d": round(total_views / approved_event_count, 1) if approved_event_count else 0,
            "share_click_rate": round(total_shares / total_views, 2) if total_views else 0,
            "orgs_with_upcoming_rate": round(orgs_with_upcoming / len(approved_orgs), 2) if approved_orgs else 0,
            "pending_content_ratio": round(
                pending_content_total / max(1, approved_event_count + len(approved_orgs) + pending_content_total),
                2,
            ),
        },
        "share_platforms_30d": [
            {"platform": platform, "count": count}
            for platform, count in summary["platform_counts"].most_common()
        ],
        "top_orgs_30d": top_orgs[:5],
        "top_events_30d": summary["top_events"][:5],
        "trend_7d": summary["trend"][-7:],
    }


async def _build_org_analytics(slug: str, days: int = 30) -> Dict[str, Any]:
    org = await _find_org(slug)
    window_start = _window_start(days)
    org_events = await db.events.find({"orgSlug": slug}, {"_id": 0}).to_list(2000)
    org_event_ids = {event.get("id") for event in org_events if event.get("id")}
    recent_events = await db.analytics_events.find(
        {"created_at": {"$gte": window_start}, "$or": [{"org_slug": slug}, {"entity_id": {"$in": list(org_event_ids)}}]},
        {"_id": 0},
    ).to_list(10000)
    follow_summary = await _follow_counters()
    summary = _summarize_analytics_events(recent_events, {event.get("id"): event for event in org_events if event.get("id")})

    now = datetime.now(timezone.utc).isoformat()
    approved_events = [event for event in org_events if event.get("status") != "pending"]
    upcoming_events = [event for event in approved_events if (event.get("start") or "") >= now]
    pending_events = [event for event in org_events if event.get("status") == "pending"]
    featured_events = [event for event in approved_events if event.get("featured")]
    unread_notifications = await db.notifications.count_documents({"org_slug": slug, "read": False})
    docs_total = await db.documents.count_documents({"org_slug": slug, "deleted": {"$ne": True}})
    feed_posts = await db.feed.count_documents({"orgSlug": slug})

    top_events = []
    for event in org_events:
        event_id = event.get("id")
        if not event_id:
            continue
        top_events.append(
            {
                "id": event_id,
                "title": event.get("title") or "Untitled event",
                "start": event.get("start"),
                "status": event.get("status") or "approved",
                "views": summary["event_views_by_event"].get(event_id, 0),
                "shares": summary["shares_by_event"].get(event_id, 0),
                "featured": bool(event.get("featured")),
                "score": summary["event_views_by_event"].get(event_id, 0) + summary["shares_by_event"].get(event_id, 0) * 2,
            }
        )
    top_events.sort(key=lambda item: (-item["score"], -item["views"], item["title"]))

    return {
        "org_slug": slug,
        "org_name": org.get("name") or slug,
        "window_days": days,
        "overview": {
            "page_views_30d": summary["page_views_by_org"].get(slug, 0),
            "event_views_30d": summary["event_views_by_org"].get(slug, 0),
            "share_clicks_30d": summary["share_clicks_by_org"].get(slug, 0),
            "followers": follow_summary["org_followers"].get(slug, 0),
            "device_followers": follow_summary["device_org_followers"].get(slug, 0),
            "subscriber_followers": follow_summary["subscriber_org_followers"].get(slug, 0),
            "published_events": len(approved_events),
            "pending_events": len(pending_events),
            "upcoming_events": len(upcoming_events),
            "featured_events": len(featured_events),
            "feed_posts": feed_posts,
            "documents": docs_total,
            "notifications_unread": unread_notifications,
        },
        "share_platforms_30d": [
            {"platform": platform, "count": count}
            for platform, count in summary["platform_counts"].most_common()
        ],
        "top_events_30d": top_events[:5],
        "trend_7d": summary["trend"][-7:],
    }


class Socials(BaseModel):
    facebook: Optional[str] = ""
    instagram: Optional[str] = ""
    tiktok: Optional[str] = ""
    linkedin: Optional[str] = ""


class Organisation(BaseModel):
    model_config = ConfigDict(extra="ignore")
    slug: str
    name: str
    category: str
    tags: List[str] = []
    short: str = ""
    about: str = ""
    does: str = ""
    forWho: str = ""
    meeting: str = ""
    address: str = ""
    location: str = ""
    email: str = ""
    phone: str = ""
    website: str = ""
    socials: Socials = Field(default_factory=Socials)
    brandColor: str = "#0052FF"
    logo: str = "✨"
    cover: str = ""
    logo_path: Optional[str] = None
    logo_thumb_path: Optional[str] = None
    cover_path: Optional[str] = None
    upcoming: int = 0
    status: Literal["approved", "pending", "rejected", "suspended", "archived"] = "approved"
    verified: bool = False
    owner_email: Optional[EmailStr] = None
    admin_emails: List[EmailStr] = []
    trust_level: Literal["new", "unverified", "trusted"] = "new"
    archived_at: Optional[str] = None
    suspended_at: Optional[str] = None
    merged_into: Optional[str] = None
    updated_at: str = Field(default_factory=now_iso)


class EventRecurrence(BaseModel):
    """Recurrence with optional custom interval, nth-weekday-of-month mode and extra one-off dates."""
    freq: Literal["none", "daily", "weekly", "biweekly", "monthly", "monthly_weekday", "annually"] = "none"
    interval: Optional[int] = None  # every N units (1-12) for daily/weekly/monthly/monthly_weekday
    until: Optional[str] = None  # ISO date/datetime — inclusive upper bound
    count: Optional[int] = None  # OR a maximum number of instances (capped at 60)
    extra_dates: List[str] = []  # additional one-off YYYY-MM-DD dates the event also happens on


class Event(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=new_id)
    title: str
    orgSlug: str
    category: str
    start: str
    end: str
    venue: str = ""
    address: str = ""
    description: str = ""
    cost: str = ""
    age: str = ""
    accessibility: str = ""
    booking: str = ""
    contactEmail: str = ""
    contactPhone: str = ""
    image: str = ""
    featured: bool = False
    status: Literal["approved", "pending", "rejected", "draft", "cancelled"] = "pending"
    recurrence: Optional[EventRecurrence] = None


class FeedPost(BaseModel):
    id: str = Field(default_factory=new_id)
    orgSlug: str
    type: str
    title: str
    body: str
    image: str = ""
    time: str = Field(default_factory=now_iso)


class Subscriber(BaseModel):
    id: str = Field(default_factory=new_id)
    email: EmailStr
    device_id: Optional[str] = None
    unsub_token: str = Field(default_factory=new_token)
    pref_token: str = Field(default_factory=new_token)
    followed_orgs: List[str] = []
    followed_categories: List[str] = []
    saved_events: List[str] = []
    digest: bool = True
    unsubscribed: bool = False
    created_at: str = Field(default_factory=now_iso)


class Notification(BaseModel):
    id: str = Field(default_factory=new_id)
    org_slug: str
    title: str
    body: str
    read: bool = False
    created_at: str = Field(default_factory=now_iso)


class AdminMessage(BaseModel):
    id: str = Field(default_factory=new_id)
    from_org_slug: Optional[str] = None
    from_email: Optional[str] = None
    from_name: Optional[str] = None
    to_org_slug: Optional[str] = None
    to_email: Optional[str] = None
    subject: str
    body: str
    in_reply_to: Optional[str] = None  # notification id the org is replying to
    parent_message_id: Optional[str] = None
    direction: Literal["inbound_org", "outbound_admin"] = "inbound_org"
    delivery: Optional[Dict[str, Any]] = None
    read: bool = False
    created_at: str = Field(default_factory=now_iso)


class Document(BaseModel):
    id: str = Field(default_factory=new_id)
    org_slug: str
    name: str
    storage_path: str
    content_type: str
    size: int
    created_at: str = Field(default_factory=now_iso)
    deleted: bool = False


class Venue(BaseModel):
    id: str = Field(default_factory=new_id)
    name: str
    address: str
    facilities: List[str] = []
    accessibility: str = ""
    capacity: int = 0
    booking: str = ""
    image: str = ""


class VolunteerOpp(BaseModel):
    id: str
    title: str
    orgSlug: str
    description: str
    age: str = ""
    time: str = ""
    skills: str = ""


class NewsletterEdition(BaseModel):
    id: str = Field(default_factory=new_id)
    subject: str
    body_html: str
    body_intro: str = ""
    scheduled_for: Optional[str] = None
    sent_at: Optional[str] = None
    created_at: str = Field(default_factory=now_iso)


class OrgClaimReq(BaseModel):
    contact_name: str = ""
    contact_email: EmailStr
    contact_phone: str = ""
    message: str = ""
    verification_code: Optional[str] = None


class OrgSuggestReq(BaseModel):
    model_config = ConfigDict(extra="ignore")
    name: Optional[str] = None
    short: Optional[str] = None
    about: Optional[str] = None
    does: Optional[str] = None
    forWho: Optional[str] = None
    meeting: Optional[str] = None
    address: Optional[str] = None
    location: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    website: Optional[str] = None
    socials: Optional[Socials] = None
    brandColor: Optional[str] = None
    logo: Optional[str] = None
    cover: Optional[str] = None
    tags: Optional[List[str]] = None
    contact_name: str = ""
    contact_email: Optional[EmailStr] = None
    message: str = ""


class OrgEditRequest(BaseModel):
    id: str = Field(default_factory=new_id)
    request_type: Literal["claim", "suggest_edit"]
    org_slug: str
    org_name: str
    payload: Dict[str, Any] = Field(default_factory=dict)
    contact_name: str = ""
    contact_email: Optional[EmailStr] = None
    contact_phone: str = ""
    message: str = ""
    status: Literal["pending", "approved", "rejected"] = "pending"
    reviewer_notes: str = ""
    created_at: str = Field(default_factory=now_iso)
    reviewed_at: Optional[str] = None


class OrgPasswordDoc(BaseModel):
    slug: str
    password_salt: str
    password_hash: str
    updated_at: str = Field(default_factory=now_iso)
    updated_by: str = "org"


class AdminAuditEntry(BaseModel):
    id: str = Field(default_factory=new_id)
    actor: str = "admin"
    action: str
    entity_type: Literal["org", "event", "user", "site"] = "site"
    entity_id: str = ""
    summary: str
    meta: Dict[str, Any] = Field(default_factory=dict)
    created_at: str = Field(default_factory=now_iso)


class OrgOwnershipTransferReq(BaseModel):
    owner_email: EmailStr
    add_to_admins: bool = True


class OrgAssignAdminsReq(BaseModel):
    admin_emails: List[EmailStr] = []


class OrgLifecycleReq(BaseModel):
    action: Literal["approve", "suspend", "archive", "restore", "verify", "unverify"]
    reason: str = ""


class OrgInviteClaimReq(BaseModel):
    email: EmailStr
    note: str = ""


class OrgMergeReq(BaseModel):
    primary_slug: str
    duplicate_slug: str
    archive_duplicate: bool = True


class TaxonomyPatch(BaseModel):
    event_categories: Optional[List[str]] = None
    organisation_categories: Optional[List[str]] = None


class OrgMember(BaseModel):
    id: str = Field(default_factory=new_id)
    org_slug: str
    email: EmailStr
    name: str = ""
    role: Literal["owner", "admin", "editor", "viewer"] = "editor"
    status: Literal["active", "suspended"] = "active"
    permissions: List[str] = []
    password_salt: str = ""
    password_hash: str = ""
    invited_at: Optional[str] = None
    accepted_at: Optional[str] = None
    last_login_at: Optional[str] = None
    created_at: str = Field(default_factory=now_iso)
    updated_at: str = Field(default_factory=now_iso)


class OrgMemberInvite(BaseModel):
    id: str = Field(default_factory=new_id)
    org_slug: str
    org_name: str
    email: EmailStr
    role: Literal["owner", "admin", "editor", "viewer"] = "editor"
    token: str = Field(default_factory=new_token)
    status: Literal["pending", "accepted", "revoked", "expired"] = "pending"
    invited_by: str = "admin"
    note: str = ""
    created_at: str = Field(default_factory=now_iso)
    sent_at: Optional[str] = None
    accepted_at: Optional[str] = None


class OrgMemberInviteReq(BaseModel):
    email: EmailStr
    role: Literal["owner", "admin", "editor", "viewer"] = "editor"
    note: str = ""


class OrgMemberRedeemReq(BaseModel):
    token: str
    name: str = ""
    password: str


class OrgMemberLoginReq(BaseModel):
    org_slug: str
    email: EmailStr
    password: str


class OrgMemberRoleReq(BaseModel):
    role: Literal["owner", "admin", "editor", "viewer"]


class OrgMemberSuspendReq(BaseModel):
    suspended: bool = True


# ─────────── Helpers ───────────
async def _find_org(slug: str) -> dict:
    org = await db.orgs.find_one({"slug": slug}, {"_id": 0})
    if not org:
        raise HTTPException(404, "Organisation not found")
    return org


def _strip_id(doc: dict) -> dict:
    doc.pop("_id", None)
    return doc


def _admin_from_request(request: Request, authorization: Optional[str], admin_code: Optional[str]) -> bool:
    from auth import read_admin_from_request as _read_admin
    if _read_admin(request, authorization):
        return True
    return bool(admin_code and hmac.compare_digest(admin_code, ADMIN_LAUNCH_CODE))


def _require_admin_from_request(request: Request, authorization: Optional[str], admin_code: Optional[str]) -> None:
    if not _admin_from_request(request, authorization, admin_code):
        raise HTTPException(403, "Admin authentication required")


async def _audit(
    action: str,
    entity_type: str,
    entity_id: str,
    summary: str,
    meta: Optional[Dict[str, Any]] = None,
    actor: str = "admin",
) -> None:
    await db.admin_audit.insert_one(
        AdminAuditEntry(
            actor=actor,
            action=action,
            entity_type=entity_type,
            entity_id=entity_id,
            summary=summary,
            meta=meta or {},
        ).model_dump()
    )


DEFAULT_EVENT_CATEGORIES = [
    "Community",
    "Family",
    "Children & Young People",
    "Sports & Fitness",
    "Arts & Culture",
    "Heritage",
    "Faith",
    "Health & Wellbeing",
    "Education",
    "Social",
    "Fundraising",
    "Markets & Fairs",
    "Music & Entertainment",
]

DEFAULT_EVENT_CATEGORY_IMAGES = {
    "family": "/familyevent.png",
    "youth": "/youth.png",
    "children & young people": "/youth.png",
    "children and young people": "/youth.png",
    "sport": "/sport.png",
    "sports": "/sport.png",
    "sports & fitness": "/sport.png",
    "school": "/school.png",
    "education": "/school.png",
    "charity": "/charity.png",
    "fundraising": "/charity.png",
    "business": "/business.png",
    "community": "/communityevent.png",
    "social": "/communityevent.png",
    "markets & fairs": "/foodanddrink.png",
    "markets and fairs": "/foodanddrink.png",
    "music": "/music.png",
    "music & entertainment": "/music.png",
    "arts & culture": "/music.png",
    "arts and culture": "/music.png",
    "food & drink": "/foodanddrink.png",
    "food and drink": "/foodanddrink.png",
    "volunteering": "/volunteering.png",
    "faith": "/faith.png",
    "heritage": "/heritage.png",
    "health & wellbeing": "/healthandwellbeing.png",
}

LEGACY_EVENT_IMAGE_PREFIX = "https://images.unsplash.com/photo-1533174072545-7a4b6ad7a6c3"


def _event_category_image(category: str) -> str:
    key = (category or "community").strip().lower()
    return DEFAULT_EVENT_CATEGORY_IMAGES.get(key, "/communityevent.png")


def _is_blank_or_legacy_event_image(image: str) -> bool:
    value = (image or "").strip()
    if not value:
        return True
    return value.startswith(LEGACY_EVENT_IMAGE_PREFIX)

DEFAULT_ORGANISATION_CATEGORIES = [
    "Community groups",
    "Sports & fitness",
    "Arts & culture",
    "Faith groups",
    "Health & wellbeing",
    "Education",
    "Charities & social",
    "Venues & spaces",
    "Local business",
    "Other",
]


# ─────────── Root ───────────
@api.get("/")
async def root():
    return {"message": "Blackrod Now API", "ok": True}


@api.get("/admin/stats")
async def admin_stats():
    return {
        "events_total": await db.events.count_documents({}),
        "events_pending": await db.events.count_documents({"status": "pending"}),
        "orgs_total": await db.orgs.count_documents({}),
        "orgs_pending": await db.orgs.count_documents({"status": "pending"}),
        "org_edit_requests_pending": await db.org_edit_requests.count_documents({"status": "pending"}),
        "subscribers": await db.subscribers.count_documents({"unsubscribed": False}),
        "messages_unread": await db.messages.count_documents({"read": False}),
        "analytics": await _build_site_analytics(),
    }


@api.post("/analytics/track")
async def track_analytics(req: AnalyticsTrackReq):
    await _record_analytics_event(
        req.kind,
        entity_type=req.entity_type,
        entity_id=req.entity_id,
        org_slug=req.org_slug,
        platform=req.platform,
        device_id=req.device_id,
    )
    return {"ok": True}


@api.get("/organisations/{slug}/analytics")
async def org_analytics(slug: str):
    return await _build_org_analytics(slug)


# ─────────── Site settings (public: coming-soon flag + launch date) ───────────
_SITE_SETTINGS_ID = "site"


async def _get_site_settings() -> Dict[str, Any]:
    doc = await db.site_settings.find_one({"_id": _SITE_SETTINGS_ID}) or {}
    return {
        "coming_soon": bool(doc.get("coming_soon", False)),
        "launch_at": doc.get("launch_at", "2026-09-12T09:00:00+00:00"),
        "teaser": doc.get(
            "teaser",
            "A new community hub for what's on, what's new, and what's next in Blackrod.",
        ),
        "updated_at": doc.get("updated_at"),
    }


@api.get("/site/settings")
async def get_site_settings():
    return await _get_site_settings()


class SiteSettingsPatch(BaseModel):
    coming_soon: Optional[bool] = None
    launch_at: Optional[str] = None
    teaser: Optional[str] = None


@api.post("/admin/site/settings")
async def patch_site_settings(
    patch: SiteSettingsPatch,
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    # Accept either admin JWT or legacy launch code.
    ok = False
    from auth import read_admin_from_request as _read_admin
    if _read_admin(request, authorization):
        ok = True
    elif admin_code and hmac.compare_digest(admin_code, ADMIN_LAUNCH_CODE):
        ok = True
    if not ok:
        raise HTTPException(403, "Admin authentication required")
    updates = {k: v for k, v in patch.model_dump(exclude_none=True).items()}
    if not updates:
        return await _get_site_settings()
    updates["updated_at"] = now_iso()
    await db.site_settings.update_one(
        {"_id": _SITE_SETTINGS_ID},
        {"$set": updates},
        upsert=True,
    )
    return await _get_site_settings()


async def _get_taxonomy_settings() -> Dict[str, Any]:
    doc = await db.site_settings.find_one({"_id": _SITE_SETTINGS_ID}) or {}
    event_categories = doc.get("event_categories") or DEFAULT_EVENT_CATEGORIES
    organisation_categories = doc.get("organisation_categories") or DEFAULT_ORGANISATION_CATEGORIES
    return {
        "event_categories": [str(v).strip() for v in event_categories if str(v).strip()],
        "organisation_categories": [str(v).strip() for v in organisation_categories if str(v).strip()],
        "updated_at": doc.get("updated_at"),
    }


@api.get("/taxonomy")
async def get_taxonomy_settings():
    return await _get_taxonomy_settings()


@api.post("/admin/taxonomy")
async def patch_taxonomy_settings(
    patch: TaxonomyPatch,
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_admin_from_request(request, authorization, admin_code)
    updates = {k: v for k, v in patch.model_dump(exclude_none=True).items()}
    if not updates:
        return await _get_taxonomy_settings()
    updates["updated_at"] = now_iso()
    await db.site_settings.update_one(
        {"_id": _SITE_SETTINGS_ID},
        {"$set": updates},
        upsert=True,
    )
    await _audit(
        action="taxonomy_updated",
        entity_type="site",
        entity_id="taxonomy",
        summary="Updated category taxonomy",
        meta={"keys": sorted(list(updates.keys()))},
    )
    return await _get_taxonomy_settings()




# ─────────── Seed (frontend bootstraps once from mockData) ───────────
class SeedPayload(BaseModel):
    organisations: List[Dict[str, Any]] = []
    events: List[Dict[str, Any]] = []
    feed_posts: List[Dict[str, Any]] = []
    venues: List[Dict[str, Any]] = []
    volunteers: List[Dict[str, Any]] = []


@api.get("/admin/seeded")
async def is_seeded():
    return {"seeded": (await db.orgs.count_documents({})) > 0}


@api.post("/admin/seed")
async def seed_from_payload(payload: SeedPayload, force: bool = False):
    if not force and (await db.orgs.count_documents({})) > 0:
        return {"status": "already-seeded"}
    if force:
        await asyncio.gather(
            db.orgs.delete_many({}),
            db.events.delete_many({}),
            db.feed.delete_many({}),
            db.venues.delete_many({}),
            db.volunteers.delete_many({}),
        )
    if payload.organisations:
        await db.orgs.insert_many([{**o, "status": o.get("status", "approved")} for o in payload.organisations])
    if payload.events:
        await db.events.insert_many([{**e, "status": e.get("status", "approved")} for e in payload.events])
    if payload.feed_posts:
        await db.feed.insert_many(list(payload.feed_posts))
    if payload.venues:
        await db.venues.insert_many(list(payload.venues))
    if payload.volunteers:
        await db.volunteers.insert_many(list(payload.volunteers))
    return {"status": "seeded", "orgs": len(payload.organisations), "events": len(payload.events)}


# ─────────── Organisations ───────────
@api.get("/organisations")
async def list_organisations(include_pending: bool = False):
    q = {} if include_pending else {"status": "approved"}
    return await db.orgs.find(q, {"_id": 0}).to_list(1000)


@api.get("/organisations/{slug}")
async def get_organisation(slug: str):
    return await _find_org(slug)


@api.post("/organisations")
async def submit_organisation(org: Organisation):
    # Require a setup contact email so we can send dashboard credentials on approval.
    contact_email = str(org.email or org.owner_email or "").strip().lower()
    if not contact_email:
        raise HTTPException(400, "Contact email is required when submitting an organisation")
    org.email = contact_email
    if not org.owner_email:
        org.owner_email = contact_email
    org.status = "pending"
    if not org.trust_level:
        org.trust_level = "new"
    doc = org.model_dump()
    doc["updated_at"] = now_iso()
    await db.orgs.insert_one(dict(doc))
    await _audit(
        action="org_submitted",
        entity_type="org",
        entity_id=doc.get("slug", ""),
        summary=f"Organisation submitted: {doc.get('name', doc.get('slug', ''))}",
        actor="public",
    )
    return _strip_id(doc)


class OrgPatch(BaseModel):
    name: Optional[str] = None
    category: Optional[str] = None
    tags: Optional[List[str]] = None
    short: Optional[str] = None
    about: Optional[str] = None
    does: Optional[str] = None
    forWho: Optional[str] = None
    meeting: Optional[str] = None
    address: Optional[str] = None
    location: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    website: Optional[str] = None
    socials: Optional[Socials] = None
    brandColor: Optional[str] = None
    logo: Optional[str] = None
    cover: Optional[str] = None
    logo_path: Optional[str] = None
    logo_thumb_path: Optional[str] = None
    cover_path: Optional[str] = None


@api.patch("/organisations/{slug}")
async def patch_organisation(
    slug: str,
    patch: OrgPatch,
    org_auth: Optional[str] = Header(None, alias="X-Org-Auth"),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_org_write_access(slug, org_auth, admin_code)
    updates = {k: v for k, v in patch.model_dump(exclude_none=True).items()}
    if not updates:
        return await _find_org(slug)
    updates["updated_at"] = now_iso()
    res = await db.orgs.update_one({"slug": slug}, {"$set": updates})
    if not res.matched_count:
        raise HTTPException(404, "Organisation not found")
    return await _find_org(slug)


@api.post("/admin/organisations/{slug}/status")
async def admin_org_status(
    slug: str,
    body: Dict[str, str],
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_admin_from_request(request, authorization, admin_code or (body.get("admin_code") or ""))
    status = body.get("status", "approved")
    existing = await _find_org(slug)
    updates: Dict[str, Any] = {"status": status, "updated_at": now_iso()}
    if status == "suspended":
        updates["suspended_at"] = now_iso()
    if status == "archived":
        updates["archived_at"] = now_iso()
    await db.orgs.update_one({"slug": slug}, {"$set": updates})
    org = await _find_org(slug)
    if status == "approved" and existing.get("status") != "approved":
        await _provision_org_credentials_on_approval(org, approved_by="admin-status")
    await _audit(
        action="org_status_changed",
        entity_type="org",
        entity_id=slug,
        summary=f"Organisation status set to {status}",
        meta={"status": status},
    )
    return org


@api.delete("/admin/organisations/{slug}")
async def admin_delete_org(slug: str):
    await db.orgs.delete_one({"slug": slug})
    await _audit(
        action="org_deleted",
        entity_type="org",
        entity_id=slug,
        summary="Organisation deleted",
    )
    return {"ok": True}


@api.post("/admin/organisations")
async def admin_create_organisation(
    org: Organisation,
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_admin_from_request(request, authorization, admin_code)
    doc = org.model_dump()
    doc["status"] = doc.get("status") or "approved"
    doc["updated_at"] = now_iso()
    exists = await db.orgs.find_one({"slug": doc["slug"]}, {"_id": 1})
    if exists:
        raise HTTPException(409, "Organisation slug already exists")
    await db.orgs.insert_one(doc)
    await _audit(
        action="org_created",
        entity_type="org",
        entity_id=doc["slug"],
        summary=f"Organisation created: {doc.get('name', doc['slug'])}",
    )
    return _strip_id(doc)


@api.post("/admin/organisations/{slug}/assign-admins")
async def admin_assign_org_admins(
    slug: str,
    req: OrgAssignAdminsReq,
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_admin_from_request(request, authorization, admin_code)
    org = await _find_org(slug)
    existing = [str(v).strip().lower() for v in (org.get("admin_emails") or []) if str(v).strip()]
    incoming = [str(v).strip().lower() for v in (req.admin_emails or []) if str(v).strip()]
    merged = sorted(set(existing + incoming))
    await db.orgs.update_one({"slug": slug}, {"$set": {"admin_emails": merged, "updated_at": now_iso()}})
    await _audit(
        action="org_admins_assigned",
        entity_type="org",
        entity_id=slug,
        summary=f"Assigned {len(incoming)} organisation admin email(s)",
        meta={"emails": incoming},
    )
    return await _find_org(slug)


@api.post("/admin/organisations/{slug}/transfer-ownership")
async def admin_transfer_org_ownership(
    slug: str,
    req: OrgOwnershipTransferReq,
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_admin_from_request(request, authorization, admin_code)
    org = await _find_org(slug)
    owner = str(req.owner_email).strip().lower()
    admin_emails = [str(v).strip().lower() for v in (org.get("admin_emails") or []) if str(v).strip()]
    if req.add_to_admins and owner not in admin_emails:
        admin_emails.append(owner)
    await db.orgs.update_one(
        {"slug": slug},
        {"$set": {"owner_email": owner, "admin_emails": sorted(set(admin_emails)), "updated_at": now_iso()}},
    )
    await _audit(
        action="org_ownership_transferred",
        entity_type="org",
        entity_id=slug,
        summary=f"Transferred ownership to {owner}",
        meta={"owner_email": owner},
    )
    return await _find_org(slug)


@api.post("/admin/organisations/{slug}/lifecycle")
async def admin_org_lifecycle(
    slug: str,
    req: OrgLifecycleReq,
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_admin_from_request(request, authorization, admin_code)
    existing = await _find_org(slug)
    updates: Dict[str, Any] = {"updated_at": now_iso()}
    if req.action == "approve":
        updates.update({"status": "approved"})
    elif req.action == "suspend":
        updates.update({"status": "suspended", "suspended_at": now_iso()})
    elif req.action == "archive":
        updates.update({"status": "archived", "archived_at": now_iso()})
    elif req.action == "restore":
        updates.update({"status": "approved", "suspended_at": None})
    elif req.action == "verify":
        updates.update({"verified": True})
    elif req.action == "unverify":
        updates.update({"verified": False})
    await db.orgs.update_one({"slug": slug}, {"$set": updates})
    org = await _find_org(slug)
    if req.action == "approve" and existing.get("status") != "approved":
        await _provision_org_credentials_on_approval(org, approved_by="admin-lifecycle")
    await _audit(
        action="org_lifecycle",
        entity_type="org",
        entity_id=slug,
        summary=f"Organisation lifecycle action: {req.action}",
        meta={"action": req.action, "reason": req.reason},
    )
    return org


@api.get("/admin/organisations/without-admins")
async def admin_orgs_without_admins(
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_admin_from_request(request, authorization, admin_code)
    rows = await db.orgs.find({}, {"_id": 0}).to_list(3000)
    out = []
    for row in rows:
        owner = str(row.get("owner_email") or "").strip()
        admins = [str(v or "").strip() for v in (row.get("admin_emails") or []) if str(v or "").strip()]
        if not owner and not admins:
            out.append(row)
    return out


@api.post("/admin/organisations/{slug}/invite-claim")
async def admin_invite_claim_org(
    slug: str,
    req: OrgInviteClaimReq,
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_admin_from_request(request, authorization, admin_code)
    org = await _find_org(slug)
    invite_id = new_id()
    claim_url = f"{PUBLIC_URL}/organisations/{slug}"
    invite = {
        "id": invite_id,
        "org_slug": slug,
        "org_name": org.get("name") or slug,
        "email": str(req.email).strip().lower(),
        "note": req.note.strip(),
        "status": "pending",
        "created_at": now_iso(),
    }
    await db.org_claim_invites.insert_one(invite)
    html_body = (
        f"<p>Hello,</p>"
        f"<p>You have been invited to claim and manage <strong>{html.escape(invite['org_name'])}</strong> on Blackrod Now.</p>"
        f"<p>Open the organisation page and click <em>Claim this profile</em>:</p>"
        f"<p><a href='{claim_url}'>{claim_url}</a></p>"
        + (f"<p><em>{html.escape(invite['note'])}</em></p>" if invite["note"] else "")
        + "<p>Once approved by site admin, you will receive dashboard login details.</p>"
        + _EMAIL_SIGNATURE
    )
    asyncio.create_task(asyncio.to_thread(resend_send, invite["email"], f"Claim your Blackrod Now page: {invite['org_name']}", html_body))
    await _audit(
        action="org_claim_invited",
        entity_type="org",
        entity_id=slug,
        summary=f"Sent claim invitation to {invite['email']}",
        meta={"invite_id": invite_id},
    )
    return invite


@api.get("/admin/org-claim-invites")
async def admin_list_org_claim_invites(
    status: Optional[str] = None,
    request: Request = None,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_admin_from_request(request, authorization, admin_code)
    q: Dict[str, Any] = {}
    if status:
        q["status"] = status
    return await db.org_claim_invites.find(q, {"_id": 0}).sort("created_at", -1).to_list(400)


@api.get("/admin/organisations/{slug}/members")
async def admin_list_org_members(
    slug: str,
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_admin_from_request(request, authorization, admin_code)
    await _find_org(slug)
    members = await db.org_members.find({"org_slug": slug}, {"_id": 0}).sort("created_at", -1).to_list(500)
    invites = await db.org_member_invites.find({"org_slug": slug}, {"_id": 0}).sort("created_at", -1).to_list(500)
    return {"members": members, "invites": invites}


@api.post("/admin/organisations/{slug}/members/invite")
async def admin_invite_org_member(
    slug: str,
    req: OrgMemberInviteReq,
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_admin_from_request(request, authorization, admin_code)
    org = await _find_org(slug)
    invite = OrgMemberInvite(
        org_slug=slug,
        org_name=org.get("name") or slug,
        email=str(req.email).strip().lower(),
        role=req.role,
        invited_by="admin",
        note=req.note.strip(),
        sent_at=now_iso(),
    ).model_dump()
    await db.org_member_invites.insert_one(invite)
    asyncio.create_task(_send_org_member_invite(invite))
    await _audit(
        action="org_member_invited",
        entity_type="user",
        entity_id=invite["email"],
        summary=f"Invited {invite['email']} to {org.get('name') or slug} as {invite['role']}",
        meta={"org_slug": slug, "invite_id": invite["id"]},
    )
    return invite


@api.post("/admin/member-invites/{invite_id}/resend")
async def admin_resend_member_invite(
    invite_id: str,
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_admin_from_request(request, authorization, admin_code)
    invite = await db.org_member_invites.find_one({"id": invite_id}, {"_id": 0})
    if not invite:
        raise HTTPException(404, "Invitation not found")
    if invite.get("status") != "pending":
        raise HTTPException(400, "Only pending invitations can be resent")
    await db.org_member_invites.update_one({"id": invite_id}, {"$set": {"sent_at": now_iso()}})
    invite["sent_at"] = now_iso()
    asyncio.create_task(_send_org_member_invite(invite))
    await _audit(
        action="org_member_invite_resent",
        entity_type="user",
        entity_id=invite.get("email") or invite_id,
        summary=f"Resent invitation to {invite.get('email')}",
        meta={"invite_id": invite_id},
    )
    return {"ok": True, "invite_id": invite_id}


@api.post("/admin/member-invites/{invite_id}/reset")
async def admin_reset_member_invite(
    invite_id: str,
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_admin_from_request(request, authorization, admin_code)
    invite = await db.org_member_invites.find_one({"id": invite_id}, {"_id": 0})
    if not invite:
        raise HTTPException(404, "Invitation not found")
    if invite.get("status") != "pending":
        raise HTTPException(400, "Only pending invitations can be reset")
    new_token = new_token()
    await db.org_member_invites.update_one(
        {"id": invite_id},
        {"$set": {"token": new_token, "sent_at": now_iso()}},
    )
    invite["token"] = new_token
    invite["sent_at"] = now_iso()
    asyncio.create_task(_send_org_member_invite(invite))
    await _audit(
        action="org_member_invite_reset",
        entity_type="user",
        entity_id=invite.get("email") or invite_id,
        summary=f"Reset invitation token for {invite.get('email')}",
        meta={"invite_id": invite_id},
    )
    return {"ok": True, "invite_id": invite_id, "token": new_token}


@api.get("/admin/member-invites")
async def admin_list_member_invites(
    status: Optional[str] = None,
    request: Request = None,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_admin_from_request(request, authorization, admin_code)
    q: Dict[str, Any] = {}
    if status:
        q["status"] = status
    return await db.org_member_invites.find(q, {"_id": 0}).sort("created_at", -1).to_list(1000)


@api.post("/organisations/member-invites/redeem")
async def redeem_org_member_invite(req: OrgMemberRedeemReq):
    invite = await db.org_member_invites.find_one({"token": req.token}, {"_id": 0})
    if not invite:
        raise HTTPException(404, "Invitation token not found")
    if invite.get("status") != "pending":
        raise HTTPException(400, "Invitation is no longer valid")
    password = (req.password or "").strip()
    if len(password) < 8:
        raise HTTPException(400, "Password must be at least 8 characters")

    salt = secrets.token_hex(16)
    member = {
        "id": new_id(),
        "org_slug": invite.get("org_slug") or "",
        "email": str(invite.get("email") or "").strip().lower(),
        "name": (req.name or "").strip(),
        "role": invite.get("role") or "editor",
        "status": "active",
        "permissions": [],
        "password_salt": salt,
        "password_hash": _hash_member_password(password, salt),
        "invited_at": invite.get("created_at"),
        "accepted_at": now_iso(),
        "last_login_at": None,
        "created_at": now_iso(),
        "updated_at": now_iso(),
    }
    existing = await db.org_members.find_one(
        {"org_slug": member["org_slug"], "email": member["email"]},
        {"_id": 0},
    )
    if existing:
        await db.org_members.update_one(
            {"id": existing["id"]},
            {
                "$set": {
                    "name": member["name"] or existing.get("name") or "",
                    "role": member["role"],
                    "status": "active",
                    "password_salt": member["password_salt"],
                    "password_hash": member["password_hash"],
                    "accepted_at": now_iso(),
                    "updated_at": now_iso(),
                }
            },
        )
        member_id = existing["id"]
    else:
        await db.org_members.insert_one(member)
        member_id = member["id"]

    await db.org_member_invites.update_one(
        {"id": invite["id"]},
        {"$set": {"status": "accepted", "accepted_at": now_iso()}},
    )

    await _audit(
        action="org_member_invite_accepted",
        entity_type="user",
        entity_id=member["email"],
        summary=f"Member accepted invite for {invite.get('org_name') or invite.get('org_slug')}",
        actor="public",
        meta={"member_id": member_id, "org_slug": invite.get("org_slug")},
    )
    return {
        "ok": True,
        "member_id": member_id,
        "org_slug": invite.get("org_slug"),
        "email": member["email"],
    }


@api.post("/organisations/member/login")
async def org_member_login(req: OrgMemberLoginReq):
    org = await _find_org(req.org_slug)
    if org.get("status") in {"pending", "suspended", "archived", "rejected"}:
        raise HTTPException(403, "This organisation cannot sign in right now")
    member = await db.org_members.find_one(
        {"org_slug": req.org_slug, "email": str(req.email).strip().lower()},
        {"_id": 0},
    )
    if not member:
        raise HTTPException(404, "No member account found for this email")
    if member.get("status") != "active":
        raise HTTPException(403, "This member account is suspended")
    if not _member_password_matches(member, (req.password or "").strip()):
        raise HTTPException(401, "Invalid email or password")
    from auth import create_access_token
    token = create_access_token(
        sub=member.get("email") or "",
        role="org_member",
        extra={
            "org_slug": req.org_slug,
            "member_id": member.get("id"),
            "member_role": member.get("role") or "editor",
            "member_status": member.get("status") or "active",
        },
    )
    await db.org_members.update_one(
        {"id": member.get("id")},
        {"$set": {"last_login_at": now_iso(), "updated_at": now_iso()}},
    )
    await _audit(
        action="org_member_login",
        entity_type="user",
        entity_id=member.get("id") or "",
        summary=f"Member logged in: {member.get('email')}",
        actor="org",
        meta={"org_slug": req.org_slug},
    )
    return {
        "ok": True,
        "slug": req.org_slug,
        "org_name": org.get("name", ""),
        "token": f"Bearer {token}",
        "member": {
            "id": member.get("id"),
            "email": member.get("email"),
            "name": member.get("name") or "",
            "role": member.get("role") or "editor",
            "last_login_at": now_iso(),
        },
    }


@api.post("/admin/organisations/{slug}/members/{member_id}/role")
async def admin_set_member_role(
    slug: str,
    member_id: str,
    req: OrgMemberRoleReq,
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_admin_from_request(request, authorization, admin_code)
    await _find_org(slug)
    res = await db.org_members.update_one(
        {"id": member_id, "org_slug": slug},
        {"$set": {"role": req.role, "updated_at": now_iso()}},
    )
    if not res.matched_count:
        raise HTTPException(404, "Member not found")
    await _audit(
        action="org_member_role_changed",
        entity_type="user",
        entity_id=member_id,
        summary=f"Updated member role to {req.role}",
        meta={"org_slug": slug},
    )
    return await db.org_members.find_one({"id": member_id}, {"_id": 0})


@api.post("/admin/organisations/{slug}/members/{member_id}/suspend")
async def admin_suspend_member(
    slug: str,
    member_id: str,
    req: OrgMemberSuspendReq,
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_admin_from_request(request, authorization, admin_code)
    await _find_org(slug)
    status = "suspended" if req.suspended else "active"
    res = await db.org_members.update_one(
        {"id": member_id, "org_slug": slug},
        {"$set": {"status": status, "updated_at": now_iso()}},
    )
    if not res.matched_count:
        raise HTTPException(404, "Member not found")
    await _audit(
        action="org_member_status_changed",
        entity_type="user",
        entity_id=member_id,
        summary=f"Set member status to {status}",
        meta={"org_slug": slug},
    )
    return await db.org_members.find_one({"id": member_id}, {"_id": 0})


@api.delete("/admin/organisations/{slug}/members/{member_id}")
async def admin_remove_member(
    slug: str,
    member_id: str,
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_admin_from_request(request, authorization, admin_code)
    await _find_org(slug)
    await db.org_members.delete_one({"id": member_id, "org_slug": slug})
    await _audit(
        action="org_member_removed",
        entity_type="user",
        entity_id=member_id,
        summary="Removed organisation member",
        meta={"org_slug": slug},
    )
    return {"ok": True}


@api.post("/admin/organisations/merge")
async def admin_merge_organisations(
    req: OrgMergeReq,
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_admin_from_request(request, authorization, admin_code)
    if req.primary_slug == req.duplicate_slug:
        raise HTTPException(400, "Primary and duplicate slug cannot be the same")
    primary = await _find_org(req.primary_slug)
    duplicate = await _find_org(req.duplicate_slug)
    await db.events.update_many({"orgSlug": req.duplicate_slug}, {"$set": {"orgSlug": req.primary_slug}})
    await db.feed.update_many({"orgSlug": req.duplicate_slug}, {"$set": {"orgSlug": req.primary_slug}})
    await db.volunteers.update_many({"orgSlug": req.duplicate_slug}, {"$set": {"orgSlug": req.primary_slug}})
    await db.documents.update_many({"org_slug": req.duplicate_slug}, {"$set": {"org_slug": req.primary_slug}})
    updates = {
        "status": "archived" if req.archive_duplicate else duplicate.get("status", "approved"),
        "merged_into": req.primary_slug,
        "updated_at": now_iso(),
    }
    await db.orgs.update_one({"slug": req.duplicate_slug}, {"$set": updates})
    await _audit(
        action="org_merged",
        entity_type="org",
        entity_id=req.primary_slug,
        summary=f"Merged {duplicate.get('name', req.duplicate_slug)} into {primary.get('name', req.primary_slug)}",
        meta={"duplicate_slug": req.duplicate_slug},
    )
    return {
        "ok": True,
        "primary_slug": req.primary_slug,
        "duplicate_slug": req.duplicate_slug,
        "duplicate_status": updates["status"],
    }


@api.get("/admin/users")
async def admin_users_overview(
    q: str = "",
    request: Request = None,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_admin_from_request(request, authorization, admin_code)
    org_rows = await db.orgs.find({}, {"_id": 0, "slug": 1, "name": 1, "owner_email": 1, "admin_emails": 1}).to_list(4000)
    member_rows = await db.org_members.find({}, {"_id": 0}).to_list(10000)
    invite_rows = await db.org_member_invites.find({}, {"_id": 0}).to_list(10000)
    users: Dict[str, Dict[str, Any]] = {}
    for row in org_rows:
        slug = row.get("slug") or ""
        name = row.get("name") or slug
        emails = []
        owner = (row.get("owner_email") or "").strip().lower()
        if owner:
            emails.append((owner, "owner"))
        for mail in (row.get("admin_emails") or []):
            m = str(mail or "").strip().lower()
            if m:
                emails.append((m, "admin"))
        for email, role_label in emails:
            if email not in users:
                users[email] = {
                    "email": email,
                    "roles": set(),
                    "organisations": [],
                    "last_login": None,
                    "status": "active",
                    "pending_invitations": 0,
                }
            users[email]["roles"].add(role_label)
            users[email]["organisations"].append({"slug": slug, "name": name})

    for member in member_rows:
        email = str(member.get("email") or "").strip().lower()
        if not email:
            continue
        if email not in users:
            users[email] = {
                "email": email,
                "roles": set(),
                "organisations": [],
                "last_login": None,
                "status": "active",
                "pending_invitations": 0,
            }
        users[email]["roles"].add(str(member.get("role") or "editor"))
        users[email]["organisations"].append({
            "slug": member.get("org_slug") or "",
            "name": member.get("org_slug") or "",
        })
        users[email]["status"] = member.get("status") or users[email]["status"]
        if member.get("last_login_at"):
            current = users[email].get("last_login")
            if not current or str(member.get("last_login_at")) > str(current):
                users[email]["last_login"] = member.get("last_login_at")

    for inv in invite_rows:
        if inv.get("status") != "pending":
            continue
        email = str(inv.get("email") or "").strip().lower()
        if not email:
            continue
        if email not in users:
            users[email] = {
                "email": email,
                "roles": set(),
                "organisations": [],
                "last_login": None,
                "status": "invited",
                "pending_invitations": 0,
            }
        users[email]["pending_invitations"] = int(users[email].get("pending_invitations") or 0) + 1

    pending_claims = await db.org_edit_requests.count_documents({"request_type": "claim", "status": "pending"})
    pending_invites = await db.org_claim_invites.count_documents({"status": "pending"})
    pending_member_invites = await db.org_member_invites.count_documents({"status": "pending"})
    out = []
    needle = q.strip().lower()
    for email, payload in users.items():
        row = {
            **payload,
            "roles": sorted(payload["roles"]),
            "organisations": sorted(
                {f"{o.get('slug')}::{o.get('name')}": o for o in payload["organisations"]}.values(),
                key=lambda x: x.get("name") or "",
            ),
        }
        if needle and needle not in email and all(needle not in (o.get("name") or "").lower() for o in row["organisations"]):
            continue
        out.append(row)
    out.sort(key=lambda r: r["email"])
    return {
        "users": out,
        "pending_invites": pending_invites,
        "pending_member_invites": pending_member_invites,
        "pending_claims": pending_claims,
    }


@api.get("/admin/audit-log")
async def admin_audit_log(
    limit: int = 200,
    request: Request = None,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_admin_from_request(request, authorization, admin_code)
    capped = max(1, min(limit, 1000))
    return await db.admin_audit.find({}, {"_id": 0}).sort("created_at", -1).to_list(capped)


@api.get("/admin/events/attention")
async def admin_events_attention(
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_admin_from_request(request, authorization, admin_code)
    rows = await db.events.find({}, {"_id": 0}).to_list(10000)
    now_dt = datetime.now(timezone.utc)

    def _event_dt(value: Optional[str]) -> Optional[datetime]:
        try:
            if not value:
                return None
            return datetime.fromisoformat(value.replace("Z", "+00:00"))
        except Exception:
            return None

    def _normalized_title(v: str) -> str:
        return re.sub(r"[^a-z0-9]+", " ", (v or "").strip().lower()).strip()

    by_key = Counter()
    for ev in rows:
        key = (_normalized_title(ev.get("title") or ""), (ev.get("start") or "")[:10])
        if key[0] and key[1]:
            by_key[key] += 1

    possible_duplicate = sum(1 for _, count in by_key.items() if count > 1)
    missing_venue = sum(1 for ev in rows if not (ev.get("venue") or "").strip())
    missing_time = 0
    missing_image = sum(1 for ev in rows if not (ev.get("image") or "").strip())
    date_passed_published = 0
    status_counts = Counter((ev.get("status") or "pending") for ev in rows)
    for ev in rows:
        start = _event_dt(ev.get("start"))
        if not start:
            missing_time += 1
        else:
            if start.hour == 0 and start.minute == 0 and start.second == 0:
                missing_time += 1
        end_val = _event_dt(ev.get("end")) or start
        if (ev.get("status") or "pending") == "approved" and end_val and end_val < now_dt:
            date_passed_published += 1

    upcoming = sum(1 for ev in rows if (_event_dt(ev.get("end")) or _event_dt(ev.get("start")) or now_dt) >= now_dt)
    past = max(0, len(rows) - upcoming)

    return {
        "counts": {
            "total": len(rows),
            "upcoming": upcoming,
            "past": past,
            "draft": status_counts.get("draft", 0),
            "pending": status_counts.get("pending", 0),
            "cancelled": status_counts.get("cancelled", 0),
            "approved": status_counts.get("approved", 0),
            "rejected": status_counts.get("rejected", 0),
        },
        "attention": {
            "missing_venue": missing_venue,
            "missing_time": missing_time,
            "missing_image": missing_image,
            "possible_duplicate": possible_duplicate,
            "date_passed_but_published": date_passed_published,
        },
    }


def _clean_org_patch(data: Dict[str, Any]) -> Dict[str, Any]:
    updates = {k: v for k, v in data.items() if v not in (None, "", [], {})}
    if "socials" in updates and isinstance(updates["socials"], dict):
        updates["socials"] = Socials(**updates["socials"])
    return updates


def _org_claim_code_hash(slug: str, email: str, code: str) -> str:
    raw = f"{slug}|{email}|{code}|{ORG_AUTH_SECRET}"
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def _org_claim_code() -> str:
    return f"{secrets.randbelow(1000000):06d}"


async def _latest_claim_challenge(slug: str, email: str) -> Optional[Dict[str, Any]]:
    rows = await db.org_claim_challenges.find(
        {"org_slug": slug, "contact_email": email, "status": "pending"},
        {"_id": 0},
    ).sort("created_at", -1).to_list(1)
    return rows[0] if rows else None


def _ip_from_request(request: Request) -> str:
    if request.headers.get("x-forwarded-for"):
        return request.headers.get("x-forwarded-for", "").split(",")[0].strip()
    return (request.client.host if request and request.client else "") or ""


@api.post("/organisations/{slug}/claim")
async def claim_organisation(slug: str, req: OrgClaimReq, request: Request):
    org = await _find_org(slug)
    contact_email = str(req.contact_email).strip().lower()
    contact_name = req.contact_name.strip()
    contact_phone = req.contact_phone.strip()
    message = req.message.strip()
    ip = _ip_from_request(request)

    if len(message) < 20:
        raise HTTPException(400, "Please include more ownership detail (at least 20 characters)")

    # Step 1: request a verification code by email.
    if not (req.verification_code or "").strip():
        pending_existing = await db.org_edit_requests.find_one(
            {
                "request_type": "claim",
                "org_slug": org["slug"],
                "contact_email": contact_email,
                "status": "pending",
            },
            {"_id": 0, "id": 1},
        )
        if pending_existing:
            raise HTTPException(409, "A claim request is already pending for this email and organisation")

        hour_ago = (datetime.now(timezone.utc) - timedelta(hours=1)).isoformat()
        by_email = await db.org_claim_challenges.count_documents({
            "contact_email": contact_email,
            "created_at": {"$gte": hour_ago},
        })
        if by_email >= ORG_CLAIM_MAX_STARTS_PER_HOUR:
            raise HTTPException(429, "Too many claim attempts. Please try again later")
        if ip:
            by_ip = await db.org_claim_challenges.count_documents({
                "ip": ip,
                "created_at": {"$gte": hour_ago},
            })
            if by_ip >= (ORG_CLAIM_MAX_STARTS_PER_HOUR * 2):
                raise HTTPException(429, "Too many claim attempts from this network. Please try again later")

        code = _org_claim_code()
        challenge = {
            "id": new_id(),
            "org_slug": org["slug"],
            "org_name": org.get("name") or org["slug"],
            "contact_name": contact_name,
            "contact_email": contact_email,
            "contact_phone": contact_phone,
            "message": message,
            "code_hash": _org_claim_code_hash(org["slug"], contact_email, code),
            "attempts": 0,
            "status": "pending",
            "ip": ip,
            "created_at": now_iso(),
            "expires_at": (datetime.now(timezone.utc) + timedelta(minutes=ORG_CLAIM_VERIFY_TTL_MINUTES)).isoformat(),
        }
        await db.org_claim_challenges.insert_one(challenge)

        verify_html = (
            f"<p>Hello {html.escape(contact_name or 'there')},</p>"
            f"<p>Use this verification code to continue your claim for <strong>{html.escape(challenge['org_name'])}</strong> on Blackrod Now:</p>"
            f"<p style='font-size:28px;letter-spacing:4px;font-weight:700'>{code}</p>"
            f"<p>This code expires in {ORG_CLAIM_VERIFY_TTL_MINUTES} minutes.</p>"
            f"<p>If you did not request this, you can ignore this email.</p>"
            + _EMAIL_SIGNATURE
        )
        asyncio.create_task(asyncio.to_thread(
            resend_send,
            contact_email,
            f"Verify your claim for {challenge['org_name']} on Blackrod Now",
            verify_html,
        ))
        return {
            "ok": True,
            "requires_verification": True,
            "detail": "Verification code sent",
            "expires_in_minutes": ORG_CLAIM_VERIFY_TTL_MINUTES,
        }

    # Step 2: verify code and create the claim request.
    code_candidate = re.sub(r"\D", "", str(req.verification_code or ""))
    if len(code_candidate) != 6:
        raise HTTPException(400, "Enter a valid 6-digit verification code")

    challenge = await _latest_claim_challenge(org["slug"], contact_email)
    if not challenge:
        raise HTTPException(400, "Start your claim first so we can send a verification code")
    if (challenge.get("expires_at") or "") < now_iso():
        await db.org_claim_challenges.update_one({"id": challenge["id"]}, {"$set": {"status": "expired"}})
        raise HTTPException(400, "Verification code has expired. Request a new code")
    if int(challenge.get("attempts") or 0) >= ORG_CLAIM_VERIFY_MAX_ATTEMPTS:
        raise HTTPException(429, "Too many incorrect codes. Request a new verification code")

    expected = _org_claim_code_hash(org["slug"], contact_email, code_candidate)
    if not hmac.compare_digest(expected, challenge.get("code_hash") or ""):
        await db.org_claim_challenges.update_one(
            {"id": challenge["id"]},
            {"$inc": {"attempts": 1}, "$set": {"updated_at": now_iso()}},
        )
        raise HTTPException(400, "Verification code is incorrect")

    pending_existing = await db.org_edit_requests.find_one(
        {
            "request_type": "claim",
            "org_slug": org["slug"],
            "contact_email": contact_email,
            "status": "pending",
        },
        {"_id": 0, "id": 1},
    )
    if pending_existing:
        raise HTTPException(409, "A claim request is already pending for this email and organisation")

    edit_request = OrgEditRequest(
        request_type="claim",
        org_slug=org["slug"],
        org_name=org["name"],
        payload={
            "security": {
                "email_verified": True,
                "verification_channel": "email_code",
                "verified_at": now_iso(),
                "ip": ip,
            }
        },
        contact_name=challenge.get("contact_name", ""),
        contact_email=contact_email,
        contact_phone=challenge.get("contact_phone", ""),
        message=challenge.get("message", ""),
    )
    await db.org_edit_requests.insert_one(edit_request.model_dump())
    await db.org_claim_challenges.update_one(
        {"id": challenge["id"]},
        {"$set": {"status": "verified", "verified_at": now_iso()}},
    )
    await _audit(
        action="org_claim_submitted",
        entity_type="org",
        entity_id=org["slug"],
        summary=f"Claim submitted after email verification by {contact_email}",
        actor="public",
        meta={"channel": "email_code"},
    )
    return edit_request


@api.post("/organisations/{slug}/suggest-edits")
async def suggest_org_edits(slug: str, req: OrgSuggestReq):
    org = await _find_org(slug)
    payload = _clean_org_patch(req.model_dump(exclude_none=True))
    for key in ("contact_name", "contact_email", "message"):
        payload.pop(key, None)
    if not payload:
        raise HTTPException(400, "No suggested changes supplied")
    edit_request = OrgEditRequest(
        request_type="suggest_edit",
        org_slug=org["slug"],
        org_name=org["name"],
        payload=payload,
        contact_name=req.contact_name.strip(),
        contact_email=req.contact_email,
        message=req.message.strip(),
    )
    await db.org_edit_requests.insert_one(edit_request.model_dump())
    return edit_request


@api.get("/admin/org-edit-requests")
async def list_org_edit_requests(status: Optional[str] = None):
    q: Dict[str, Any] = {}
    if status:
        q["status"] = status
    return await db.org_edit_requests.find(q, {"_id": 0}).sort("created_at", -1).to_list(200)


class OrgEditRequestReview(BaseModel):
    status: Literal["approved", "rejected"]
    reviewer_notes: str = ""


class OrgPasswordVerifyReq(BaseModel):
    password: str


class OrgPasswordChangeReq(BaseModel):
    new_password: str
    current_password: Optional[str] = None
    admin_code: Optional[str] = None


class OrgAccessLoginReq(BaseModel):
    password: Optional[str] = None
    admin_code: Optional[str] = None


@api.post("/admin/org-edit-requests/{request_id}/status")
async def review_org_edit_request(request_id: str, req: OrgEditRequestReview):
    existing = await db.org_edit_requests.find_one({"id": request_id}, {"_id": 0})
    if not existing:
        raise HTTPException(404, "Request not found")
    updates = {
        "status": req.status,
        "reviewer_notes": req.reviewer_notes,
        "reviewed_at": now_iso(),
    }
    await db.org_edit_requests.update_one({"id": request_id}, {"$set": updates})
    if req.status == "approved" and existing.get("request_type") == "suggest_edit" and existing.get("payload"):
        org_updates = dict(existing.get("payload") or {})
        if org_updates:
            org_updates["updated_at"] = now_iso()
            if "socials" in org_updates and isinstance(org_updates["socials"], dict):
                org_updates["socials"] = Socials(**org_updates["socials"])
            await db.orgs.update_one({"slug": existing["org_slug"]}, {"$set": org_updates})
    if req.status == "approved" and existing.get("request_type") == "claim":
        # Set the default password (idempotent — won't overwrite a custom one already set).
        org_slug = existing["org_slug"]
        existing_pwd = await _get_org_password_doc(org_slug)
        if not existing_pwd:
            await _set_org_password(org_slug, ORG_DEFAULT_PASSWORD, updated_by="claim-approval")
        # Email credentials to the claimant.
        claimant_email = existing.get("contact_email")
        claimant_name = existing.get("contact_name") or "there"
        org_name = existing.get("org_name") or org_slug
        temp_password = ORG_DEFAULT_PASSWORD if not existing_pwd else "(your existing password)"
        dashboard_url = f"{PUBLIC_URL}/organisation-dashboard"
        email_html = (
            f"<p>Hi {html.escape(claimant_name)},</p>"
            f"<p>Your request to manage <strong>{html.escape(org_name)}</strong> on Blackrod Now has been approved.</p>"
            f"<p>You can log in to your organisation dashboard here:<br>"
            f"<a href='{dashboard_url}'>{dashboard_url}</a></p>"
            f"<p><strong>Organisation:</strong> {html.escape(org_name)}<br>"
            f"<strong>Temporary password:</strong> {html.escape(temp_password)}</p>"
            + (f"<p><em>{html.escape(req.reviewer_notes)}</em></p>" if req.reviewer_notes.strip() else "")
            + f"<p>Please change your password after your first login.</p>"
            f"<p>If you have any questions, just reply to this email.</p>"
            + _EMAIL_SIGNATURE
        )
        if claimant_email:
            asyncio.create_task(asyncio.to_thread(
                resend_send,
                claimant_email,
                f"Your {org_name} page on Blackrod Now \u2014 access granted",
                email_html,
                None,
                None,
                ADMIN_INBOX_EMAILS or None,
            ))
    if req.status == "rejected" and existing.get("request_type") == "claim":
        claimant_email = existing.get("contact_email")
        claimant_name = existing.get("contact_name") or "there"
        org_name = existing.get("org_name") or existing.get("org_slug", "")
        rejection_html = (
            f"<p>Hi {html.escape(claimant_name)},</p>"
            f"<p>Thank you for your request to manage <strong>{html.escape(org_name)}</strong> on Blackrod Now.</p>"
            f"<p>Unfortunately we were not able to verify your ownership at this time."
            + (f" {html.escape(req.reviewer_notes)}" if req.reviewer_notes.strip() else "")
            + "</p>"
            f"<p>If you believe this is a mistake, please reply to this email with further proof of your connection to the organisation.</p>"
            + _EMAIL_SIGNATURE
        )
        if claimant_email:
            asyncio.create_task(asyncio.to_thread(
                resend_send,
                claimant_email,
                f"Your claim request for {org_name} \u2014 update",
                rejection_html,
            ))
    return await db.org_edit_requests.find_one({"id": request_id}, {"_id": 0})


def _hash_org_password(password: str, salt: str) -> str:
    return hashlib.pbkdf2_hmac(
        "sha256",
        password.encode("utf-8"),
        salt.encode("utf-8"),
        200000,
    ).hex()


def _hash_member_password(password: str, salt: str) -> str:
    return hashlib.pbkdf2_hmac(
        "sha256",
        password.encode("utf-8"),
        salt.encode("utf-8"),
        200000,
    ).hex()


def _member_password_matches(doc: Optional[Dict[str, Any]], password: str) -> bool:
    if not doc or not password:
        return False
    expected = str(doc.get("password_hash") or "")
    salt = str(doc.get("password_salt") or "")
    if not expected or not salt:
        return False
    actual = _hash_member_password(password, salt)
    return hmac.compare_digest(actual, expected)


async def _send_org_member_invite(invite: Dict[str, Any]) -> None:
    org_slug = invite.get("org_slug") or ""
    token = invite.get("token") or ""
    redeem_url = f"{PUBLIC_URL}/member/redeem?token={token}"
    login_url = f"{PUBLIC_URL}/member/login?slug={org_slug}"
    html_body = (
        f"<p>Hello,</p>"
        f"<p>You have been invited to manage <strong>{html.escape(invite.get('org_name') or org_slug)}</strong> on Blackrod Now as <strong>{html.escape(invite.get('role') or 'editor')}</strong>.</p>"
        f"<p><strong>Activate your account</strong> by clicking the link below and setting a password:</p>"
        f"<p><a href='{redeem_url}' style='display:inline-block;padding:12px 24px;background:#000;color:#fff;border-radius:9999px;font-weight:bold;text-decoration:none'>Activate account</a></p>"
        f"<p>Or copy this link into your browser:<br>{redeem_url}</p>"
        f"<p>Once activated, sign in here: <a href='{login_url}'>{login_url}</a></p>"
        + (f"<p><em>{html.escape(invite.get('note') or '')}</em></p>" if invite.get("note") else "")
        + "<p>If you did not expect this invitation, you can ignore this email.</p>"
        + _EMAIL_SIGNATURE
    )
    await asyncio.to_thread(
        resend_send,
        invite.get("email") or "",
        f"You are invited to manage {invite.get('org_name') or org_slug} on Blackrod Now",
        html_body,
    )


def _new_org_password_doc(slug: str, password: str, updated_by: str) -> OrgPasswordDoc:
    salt = secrets.token_hex(16)
    return OrgPasswordDoc(
        slug=slug,
        password_salt=salt,
        password_hash=_hash_org_password(password, salt),
        updated_by=updated_by,
    )


def _password_matches(doc: Optional[Dict[str, Any]], password: str) -> bool:
    if not password:
        return False
    if not doc:
        return hmac.compare_digest(password, ORG_DEFAULT_PASSWORD)
    expected = doc.get("password_hash") or ""
    actual = _hash_org_password(password, doc.get("password_salt") or "")
    return hmac.compare_digest(actual, expected)


async def _get_org_password_doc(slug: str) -> Optional[Dict[str, Any]]:
    return await db.org_passwords.find_one({"slug": slug}, {"_id": 0})


async def _set_org_password(slug: str, new_password: str, updated_by: str) -> None:
    doc = _new_org_password_doc(slug, new_password, updated_by)
    await db.org_passwords.update_one(
        {"slug": slug},
        {"$set": doc.model_dump()},
        upsert=True,
    )


def _new_temporary_org_password(length: int = 12) -> str:
    alphabet = "ABCDEFGHJKLMNPQRSTUVWXYZabcdefghijkmnopqrstuvwxyz23456789!@#$%&*"
    return "".join(secrets.choice(alphabet) for _ in range(max(10, length)))


def _org_contact_email(org: Dict[str, Any]) -> str:
    return str(org.get("email") or org.get("owner_email") or "").strip().lower()


async def _provision_org_credentials_on_approval(org: Dict[str, Any], approved_by: str = "admin") -> bool:
    """Create first-time org credentials and email them to the setup contact.
    Returns True when credentials were created and email dispatch was queued."""
    slug = str(org.get("slug") or "").strip()
    if not slug:
        return False

    recipient = _org_contact_email(org)
    if not recipient:
        raise HTTPException(400, "Organisation contact email is required before approval")

    existing_pwd = await _get_org_password_doc(slug)
    if existing_pwd:
        return False

    temp_password = _new_temporary_org_password()
    await _set_org_password(slug, temp_password, updated_by=approved_by)

    org_name = org.get("name") or slug
    dashboard_url = f"{PUBLIC_URL}/organisation-dashboard"
    html_body = (
        f"<p>Hello,</p>"
        f"<p><strong>{html.escape(org_name)}</strong> has now been approved on Blackrod Now.</p>"
        f"<p>You can now access your organisation dashboard:</p>"
        f"<p><a href='{dashboard_url}'>{dashboard_url}</a></p>"
        f"<p><strong>Organisation:</strong> {html.escape(org_name)}<br>"
        f"<strong>Temporary password:</strong> {html.escape(temp_password)}</p>"
        f"<p>Please sign in and change this password after your first login.</p>"
        + _EMAIL_SIGNATURE
    )
    asyncio.create_task(asyncio.to_thread(
        resend_send,
        recipient,
        f"{org_name} has been approved on Blackrod Now",
        html_body,
        None,
        None,
        ADMIN_INBOX_EMAILS or None,
    ))
    await _audit(
        action="org_credentials_provisioned",
        entity_type="org",
        entity_id=slug,
        summary=f"Provisioned first-time credentials and emailed {recipient}",
        meta={"recipient": recipient},
    )
    return True


@api.post("/organisations/{slug}/password/verify")
async def verify_org_password(slug: str, req: OrgPasswordVerifyReq):
    await _find_org(slug)
    doc = await _get_org_password_doc(slug)
    ok = _password_matches(doc, (req.password or "").strip())
    return {"ok": ok}


@api.post("/organisations/{slug}/auth/login")
async def org_auth_login(slug: str, req: OrgPasswordVerifyReq):
    """Password → per-org access token. Frontend stores under `rn-org-tokens`
    and sends as X-Org-Auth on protected calls."""
    org = await _find_org(slug)
    if org.get("status") in {"pending", "suspended", "archived", "rejected"}:
        raise HTTPException(403, "This organisation cannot sign in right now. Contact site admin for access.")
    doc = await _get_org_password_doc(slug)
    if not _password_matches(doc, (req.password or "").strip()):
        raise HTTPException(401, "Invalid organisation password")
    # Token is opaque — the permissive middleware only checks presence today.
    token = f"org:{slug}:{new_token()}"
    return {
        "ok": True,
        "slug": slug,
        "org_name": org.get("name", ""),
        "token": token,
    }


@api.post("/organisations/{slug}/password/change")
async def change_org_password(slug: str, req: OrgPasswordChangeReq):
    await _find_org(slug)
    new_password = (req.new_password or "").strip()
    if len(new_password) < 8:
        raise HTTPException(400, "New password must be at least 8 characters")

    using_admin = False
    if req.admin_code:
        if not hmac.compare_digest((req.admin_code or "").strip(), ADMIN_LAUNCH_CODE):
            raise HTTPException(403, "Invalid admin code")
        using_admin = True
    else:
        current_password = (req.current_password or "").strip()
        doc = await _get_org_password_doc(slug)
        if not _password_matches(doc, current_password):
            raise HTTPException(403, "Current organisation password is incorrect")

    await _set_org_password(slug, new_password, updated_by=("admin" if using_admin else "org"))
    return {"ok": True}


@api.post("/admin/organisations/{slug}/password/reset")
async def reset_org_password(slug: str, body: Dict[str, str]):
    await _find_org(slug)
    admin_code = (body.get("admin_code") or "").strip()
    if not hmac.compare_digest(admin_code, ADMIN_LAUNCH_CODE):
        raise HTTPException(403, "Invalid admin code")
    password = (body.get("password") or ORG_DEFAULT_PASSWORD).strip()
    if len(password) < 8:
        raise HTTPException(400, "Password must be at least 8 characters")
    await _set_org_password(slug, password, updated_by="admin-reset")
    return {"ok": True, "slug": slug}


@api.post("/admin/organisations/{slug}/impersonate")
async def admin_impersonate_org(
    slug: str,
    body: Dict[str, str],
    request: Request,
    authorization: Optional[str] = Header(None),
):
    """Super admin exchanges credentials for an organisation access token so
    they can operate an organisation dashboard on the owner's behalf. Accepts
    EITHER an admin JWT (`Authorization: Bearer …` — new preferred flow) OR
    the legacy `admin_code` body field (backward compat during rollover).
    The returned token is a synthetic value the frontend passes via X-Org-Auth
    to signal an admin-impersonation session."""
    org = await _find_org(slug)
    admin_payload = read_admin_from_request(request, authorization)
    admin_code = (body.get("admin_code") or "").strip()
    if not admin_payload and not (admin_code and hmac.compare_digest(admin_code, ADMIN_LAUNCH_CODE)):
        raise HTTPException(403, "Admin authentication required")
    token = f"admin-impersonate:{slug}:{new_token()}"
    return {
        "ok": True,
        "slug": slug,
        "org_name": org.get("name", ""),
        "token": token,
        "mode": "impersonate",
    }


# ─────────── Events ───────────
def _nth_weekday_of_month(year: int, month: int, weekday: int, nth: int, template: datetime) -> datetime:
    """Date of the nth <weekday> in a month (falls back to the last occurrence),
    keeping time-of-day/tz from `template`."""
    first = template.replace(year=year, month=month, day=1)
    offset = (weekday - first.weekday()) % 7
    day = 1 + offset + (nth - 1) * 7
    days_in_month = ((first.replace(month=month % 12 + 1, year=year + (1 if month == 12 else 0), day=1)) - timedelta(days=1)).day
    while day > days_in_month:
        day -= 7
    return first.replace(day=day)


def _expand_recurring_event(ev: Dict[str, Any], horizon_days: int = 180) -> List[Dict[str, Any]]:
    """Return a list of virtual instances of a recurring event within
    `horizon_days` from now. The original event is included as instance #0.
    Non-recurring events are returned as-is."""
    rec = ev.get("recurrence")
    freq = (rec or {}).get("freq") or "none"
    extra_dates = (rec or {}).get("extra_dates") or []
    if not rec or (freq == "none" and not extra_dates):
        return [ev]
    try:
        start = datetime.fromisoformat((ev.get("start") or "").replace("Z", "+00:00"))
        end = datetime.fromisoformat((ev.get("end") or ev.get("start") or "").replace("Z", "+00:00"))
    except Exception:
        return [ev]
    duration = end - start
    horizon = datetime.now(timezone.utc) + timedelta(days=horizon_days)
    if start.tzinfo is None:
        horizon = horizon.replace(tzinfo=None)
    try:
        until = datetime.fromisoformat(rec.get("until").replace("Z", "+00:00")) if rec.get("until") else horizon
    except Exception:
        until = horizon
    upper = min(until, horizon)
    max_count_raw = rec.get("count")
    if max_count_raw is None:
        max_count = 240
    else:
        max_count = min(max(1, int(max_count_raw)), 240)
    try:
        interval = max(1, min(int(rec.get("interval") or 1), 12))
    except Exception:
        interval = 1

    occurrences: List[datetime] = []
    base_days = {"daily": 1, "weekly": 7, "biweekly": 14, "monthly": 30, "annually": 365}.get(freq)
    if base_days:
        step = timedelta(days=base_days * interval)
        cur, n = start, 0
        while cur <= upper and n < max_count:
            occurrences.append(cur)
            cur += step
            n += 1
    elif freq == "monthly_weekday":
        weekday = start.weekday()
        nth = (start.day - 1) // 7 + 1
        year, month = start.year, start.month
        cur, n = start, 0
        if cur < datetime.now(cur.tzinfo or timezone.utc) - timedelta(days=35):
            now_ref = datetime.now(cur.tzinfo or timezone.utc)
            diff_months = (now_ref.year - year) * 12 + (now_ref.month - month)
            if diff_months > interval:
                skip_cycles = max(0, (diff_months // interval) - 1)
                total = year * 12 + (month - 1) + (skip_cycles * interval)
                year, month = total // 12, total % 12 + 1
                try:
                    cur = _nth_weekday_of_month(year, month, weekday, nth, start)
                except Exception:
                    cur = start
        while cur <= upper and n < max_count:
            occurrences.append(cur)
            n += 1
            total = year * 12 + (month - 1) + interval
            year, month = total // 12, total % 12 + 1
            try:
                cur = _nth_weekday_of_month(year, month, weekday, nth, start)
            except Exception:
                break
    else:
        occurrences.append(start)

    seen = {o.date() for o in occurrences}
    for raw in extra_dates:
        try:
            d = datetime.fromisoformat(str(raw)[:10]).date()
            occ = start.replace(year=d.year, month=d.month, day=d.day)
        except Exception:
            continue
        if occ.date() in seen or occ > horizon:
            continue
        seen.add(occ.date())
        occurrences.append(occ)
    occurrences.sort()

    out: List[Dict[str, Any]] = []
    for occ in occurrences:
        instance = {**ev}
        instance["start"] = occ.isoformat()
        instance["end"] = (occ + duration).isoformat()
        if occ != start:
            instance["id"] = f"{ev['id']}__{occ.date().isoformat()}"
            instance["parent_id"] = ev["id"]
            instance["is_recurrence_instance"] = True
        out.append(instance)
    return out or [ev]


@api.get("/events")
async def list_events(upcoming_only: bool = False, include_pending: bool = False, expand_recurring: bool = True):
    q: Dict[str, Any] = {} if include_pending else {"status": {"$ne": "pending"}}
    events = await db.events.find(q, {"_id": 0}).to_list(2000)
    if expand_recurring:
        expanded: List[Dict[str, Any]] = []
        for e in events:
            expanded.extend(_expand_recurring_event(e))
        events = expanded
    if upcoming_only:
        nowi = now_iso()
        events = [e for e in events if (e.get("end") or e.get("start")) >= nowi]
    return events


@api.get("/events/{event_id}")
async def get_event(event_id: str):
    # Recurring virtual ids look like `<parent>__YYYY-MM-DD` — look up the parent
    # and return an on-the-fly instance stamped with the correct start/end.
    if "__" in event_id:
        parent_id, date_part = event_id.split("__", 1)
        parent = await db.events.find_one({"id": parent_id}, {"_id": 0})
        if not parent:
            raise HTTPException(404, "Event not found")
        for inst in _expand_recurring_event(parent):
            if inst.get("id") == event_id or (inst.get("start", "")[:10] == date_part and inst.get("parent_id") == parent_id):
                return inst
        raise HTTPException(404, "Event instance not found")
    e = await db.events.find_one({"id": event_id}, {"_id": 0})
    if not e:
        raise HTTPException(404, "Event not found")
    return e


@api.post("/events/{event_id}/duplicate")
async def duplicate_event(
    event_id: str,
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
    org_auth: Optional[str] = Header(None, alias="X-Org-Auth"),
):
    """One-click clone of an event. Same fields, blanked start/end (org fills
    the new date), status='pending' (or 'approved' if admin). Returns the
    saved new event."""
    src = await db.events.find_one({"id": event_id}, {"_id": 0})
    if not src:
        # Also allow duplicating a recurring instance — clone the parent.
        if "__" in event_id:
            src = await db.events.find_one({"id": event_id.split("__", 1)[0]}, {"_id": 0})
        if not src:
            raise HTTPException(404, "Event not found")
    role = _require_org_write_access(src["orgSlug"], org_auth=org_auth, admin_code=admin_code)
    new_ev = {**src}
    new_ev["id"] = new_id()
    new_ev["title"] = f"Copy of {src.get('title','')}".strip()
    new_ev["start"] = ""
    new_ev["end"] = ""
    new_ev["featured"] = False
    new_ev["recurrence"] = None
    new_ev["status"] = "approved" if role == "admin" else "pending"
    new_ev.pop("parent_id", None)
    new_ev.pop("is_recurrence_instance", None)
    await db.events.insert_one(new_ev)
    return {k: v for k, v in new_ev.items() if k != "_id"}


# ─────────── Event OG page (rich Facebook/LinkedIn/WhatsApp previews) ───────────
# Returns a tiny HTML page with per-event Open Graph tags for social crawlers,
# plus a meta-refresh + JS redirect so real humans land on the canonical
# React event page. Crawlers ignore the redirect; humans never linger here.
import html as _html_lib
from fastapi import Request as _Req
from fastapi.responses import HTMLResponse as _HTMLResp


def _abs_base_url(req: _Req) -> str:
    proto = req.headers.get("x-forwarded-proto") or req.url.scheme or "https"
    host = req.headers.get("x-forwarded-host") or req.headers.get("host") or req.url.netloc
    return f"{proto}://{host}"


@api.get("/events/{event_id}/og", response_class=_HTMLResp)
async def event_og_page(event_id: str, request: _Req):
    # Support recurring virtual instance ids (parent__YYYY-MM-DD) too.
    e = await _fetch_event_for_poster(event_id)

    org = await db.orgs.find_one({"slug": e.get("orgSlug")}, {"_id": 0}) or {}
    base = _abs_base_url(request)
    canonical = f"{base}/events/{event_id}"

    # Resolve a good preview image (event image → org cover → org logo → site logo)
    img = e.get("image") or ""
    if not img and org.get("cover_path"):
        img = f"{base}/api/organisations/{org['slug']}/cover"
    if not img and org.get("logo_path"):
        img = f"{base}/api/organisations/{org['slug']}/logo"
    if not img:
        img = f"{base}/logo.png"

    # Format a human date line for the description prefix
    try:
        start_iso = e.get("start") or ""
        dt = datetime.fromisoformat(start_iso.replace("Z", "+00:00")) if start_iso else None
        when = dt.strftime("%a %d %b %Y · %H:%M") if dt else ""
    except Exception:
        when = ""

    raw_title = e.get("title") or "Blackrod Event"
    raw_desc_parts = [p for p in [when, e.get("venue"), e.get("description")] if p]
    raw_desc = " · ".join(raw_desc_parts[:2])
    if e.get("description"):
        raw_desc = (raw_desc + " — " + e["description"]) if raw_desc else e["description"]
    raw_desc = raw_desc[:280]
    # Trim to nearest word boundary for prettier previews.
    if len(raw_desc) == 280 and " " in raw_desc:
        raw_desc = raw_desc.rsplit(" ", 1)[0].rstrip(",.:;-—") + "…"

    esc = _html_lib.escape
    title = esc(raw_title)
    desc = esc(raw_desc)
    site_name = "Blackrod Now"
    img_url = esc(img)
    canonical_esc = esc(canonical)

    # Sensible image dimensions for Facebook/LinkedIn/Twitter cards. If the
    # image is one of our org logos we already know it's 512×512; otherwise
    # assume the standard event flyer ratio 1200×630.
    if img.endswith("/logo") or img.endswith("/logo.png"):
        img_w, img_h = 512, 512
    elif "/organisations/" in img and img.endswith("/cover"):
        img_w, img_h = 1600, 500
    else:
        img_w, img_h = 1200, 630

    body = f"""<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8" />
<title>{title} — Blackrod Now</title>
<meta name="description" content="{desc}" />
<link rel="canonical" href="{canonical_esc}" />

<meta property="og:type" content="article" />
<meta property="og:site_name" content="{site_name}" />
<meta property="og:title" content="{title}" />
<meta property="og:description" content="{desc}" />
<meta property="og:url" content="{canonical_esc}" />
<meta property="og:image" content="{img_url}" />
<meta property="og:image:secure_url" content="{img_url}" />
<meta property="og:image:width" content="{img_w}" />
<meta property="og:image:height" content="{img_h}" />
<meta property="og:image:alt" content="{title}" />

<meta name="twitter:card" content="summary_large_image" />
<meta name="twitter:title" content="{title}" />
<meta name="twitter:description" content="{desc}" />
<meta name="twitter:image" content="{img_url}" />

<!-- IMPORTANT: no <meta http-equiv="refresh"> — Facebook's crawler follows
     meta-refresh and would then read the SPA's generic OG tags instead of
     these per-event tags. We rely on JS redirect for humans (crawlers do
     not execute JavaScript). -->
<script>window.location.replace({canonical_esc!r});</script>
<style>body{{font-family:system-ui,sans-serif;padding:2rem;color:#333;max-width:640px;margin:0 auto;text-align:center}}a{{color:#0052FF;text-decoration:none;font-weight:600}}</style>
</head>
<body>
<h1 style="font-size:20px">{title}</h1>
<p>Opening on Blackrod Now…</p>
<p><a href="{canonical_esc}">Continue &rarr;</a></p>
</body>
</html>"""

    return _HTMLResp(
        content=body,
        headers={"Cache-Control": "public, max-age=300"},
    )


# ─────────── Live calendar feed (.ics / webcal) ───────────
from fastapi.responses import PlainTextResponse as _PlainResp


def _ics_escape(txt: str) -> str:
    """Escape reserved iCalendar characters per RFC 5545."""
    if not txt:
        return ""
    return (
        txt.replace("\\", "\\\\")
        .replace(",", "\\,")
        .replace(";", "\\;")
        .replace("\n", "\\n")
        .replace("\r", "")
    )


def _ics_dt(iso: str) -> str:
    """Convert ISO datetime to iCal UTC form: 20260912T140000Z."""
    if not iso:
        return ""
    try:
        dt = datetime.fromisoformat(iso.replace("Z", "+00:00"))
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt.astimezone(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    except Exception:
        return ""


def _ics_fold(line: str) -> str:
    """Fold long lines to 75 octets per RFC 5545."""
    if len(line.encode("utf-8")) <= 75:
        return line
    out = []
    while line:
        chunk = line[:73]
        out.append(chunk)
        line = line[73:]
        if line:
            line = " " + line  # leading space marks continuation
    return "\r\n".join(out)


@api.get("/calendar.ics", response_class=_PlainResp)
async def calendar_feed(
    request: _Req,
    device: Optional[str] = None,
    orgs: Optional[str] = None,
    category: Optional[str] = None,
    upcoming_only: bool = True,
):
    """Live iCalendar feed of Blackrod events.

    Filters (all optional, apply in order of precedence):
      • ?device=<uuid>   – events for orgs/categories the anonymous user follows
      • ?orgs=slug1,slug2
      • ?category=Community
      • ?upcoming_only=false to include past events (default true)
    """
    base = _abs_base_url(request)
    q: Dict[str, Any] = {"status": {"$ne": "pending"}}

    org_slugs: Optional[List[str]] = None
    cats: Optional[List[str]] = None

    if device:
        sub = await db.subscribers.find_one({"deviceId": device}) or {}
        follows = sub.get("follows") or {}
        org_slugs = follows.get("orgs") or None
        cats = follows.get("categories") or None
    if orgs:
        org_slugs = [s.strip() for s in orgs.split(",") if s.strip()]
    if category:
        cats = [category]

    if org_slugs and cats:
        q["$or"] = [{"orgSlug": {"$in": org_slugs}}, {"category": {"$in": cats}}]
    elif org_slugs:
        q["orgSlug"] = {"$in": org_slugs}
    elif cats:
        q["category"] = {"$in": cats}

    events = await db.events.find(q, {"_id": 0}).sort("start", 1).to_list(2000)
    if upcoming_only:
        nowi = now_iso()
        events = [e for e in events if (e.get("end") or e.get("start") or "") >= nowi]

    # Calendar name reflects the applied filter
    name_parts = ["Blackrod Now"]
    if org_slugs:
        name_parts.append(f"{len(org_slugs)} orgs")
    if cats:
        name_parts.append(cats[0] if len(cats) == 1 else f"{len(cats)} categories")
    if device and not org_slugs and not cats:
        name_parts.append("Personal")
    cal_name = " · ".join(name_parts)
    cal_desc = "Live-updating calendar feed from Blackrod Now"

    lines = [
        "BEGIN:VCALENDAR",
        "VERSION:2.0",
        "PRODID:-//Blackrod Now//Events Feed//EN",
        "CALSCALE:GREGORIAN",
        "METHOD:PUBLISH",
        f"X-WR-CALNAME:{_ics_escape(cal_name)}",
        f"NAME:{_ics_escape(cal_name)}",
        f"X-WR-CALDESC:{_ics_escape(cal_desc)}",
        "REFRESH-INTERVAL;VALUE=DURATION:PT1H",
        "X-PUBLISHED-TTL:PT1H",
    ]

    dtstamp = _ics_dt(now_iso())
    for e in events:
        dtstart = _ics_dt(e.get("start") or "")
        dtend = _ics_dt(e.get("end") or e.get("start") or "")
        if not dtstart:
            continue
        canonical = f"{base}/events/{e.get('id')}"
        loc = ", ".join([p for p in [e.get("venue"), e.get("address")] if p])
        desc_parts = [e.get("description") or "", f"Details: {canonical}"]
        desc = "\n\n".join([p for p in desc_parts if p])
        lines += [
            "BEGIN:VEVENT",
            _ics_fold(f"UID:{e.get('id')}@blackrodnow"),
            f"DTSTAMP:{dtstamp}",
            f"DTSTART:{dtstart}",
            f"DTEND:{dtend}",
            _ics_fold(f"SUMMARY:{_ics_escape(e.get('title') or 'Blackrod event')}"),
            _ics_fold(f"DESCRIPTION:{_ics_escape(desc)}"),
        ]
        if loc:
            lines.append(_ics_fold(f"LOCATION:{_ics_escape(loc)}"))
        if e.get("category"):
            lines.append(_ics_fold(f"CATEGORIES:{_ics_escape(e['category'])}"))
        lines.append(f"URL:{canonical}")
        lines.append("END:VEVENT")
    lines.append("END:VCALENDAR")

    body = "\r\n".join(lines) + "\r\n"
    return _PlainResp(
        content=body,
        media_type="text/calendar; charset=utf-8",
        headers={"Cache-Control": "public, max-age=1800"},
    )


# ─────────── Org "share pack" (per-org email pack of upcoming events) ───────────
def _org_share_pack_data(org: Dict[str, Any], events: List[Dict[str, Any]], base: str, limit: int = 6) -> Dict[str, Any]:
    upcoming = [e for e in events if e.get("orgSlug") == org["slug"] and (e.get("end") or e.get("start") or "") >= now_iso()]
    upcoming.sort(key=lambda x: x.get("start", ""))
    upcoming = upcoming[:limit]
    items = []
    for e in upcoming:
        canonical = f"{base}/events/{e['id']}"
        og_url = f"{base}/api/events/{e['id']}/og"
        share_text = f"{e['title']} — {e.get('venue') or 'Blackrod'}"
        items.append({
            "id": e["id"],
            "title": e["title"],
            "start": e.get("start"),
            "venue": e.get("venue") or "",
            "description": (e.get("description") or "")[:220],
            "image": e.get("image") or "",
            "canonical_url": canonical,
            "og_url": og_url,
            "share_text": share_text,
            "share_links": {
                "facebook": f"https://www.facebook.com/sharer/sharer.php?u={requests.utils.quote(og_url, safe='')}",
                "linkedin": f"https://www.linkedin.com/sharing/share-offsite/?url={requests.utils.quote(og_url, safe='')}",
                "twitter": f"https://twitter.com/intent/tweet?text={requests.utils.quote(share_text, safe='')}&url={requests.utils.quote(og_url, safe='')}",
                "whatsapp": f"https://wa.me/?text={requests.utils.quote(share_text + ' ' + og_url, safe='')}",
            },
        })
    return {"org": {"slug": org["slug"], "name": org["name"], "brandColor": org.get("brandColor", "#0052FF")}, "events": items, "count": len(items)}


def _render_share_pack_html(pack: Dict[str, Any], base: str) -> str:
    org = pack["org"]
    brand = org.get("brandColor", "#0052FF")
    if not pack["events"]:
        return f"""<div style="font-family:system-ui,sans-serif;padding:24px;color:#333">
          <h1 style="color:{brand}">Hello {org['name']}</h1>
          <p>You have no upcoming events on Blackrod Now this week. Add one and we'll include it in your next share pack.</p>
          <p><a href="{base}/organisation-dashboard" style="color:{brand}">Open your dashboard →</a></p>
        </div>"""
    rows = []
    for e in pack["events"]:
        img = e["image"] or f"{base}/logo.png"
        try:
            dt = datetime.fromisoformat((e["start"] or "").replace("Z", "+00:00"))
            when = dt.strftime("%a %d %b · %H:%M")
        except Exception:
            when = ""
        rows.append(f"""
        <table role="presentation" cellspacing="0" cellpadding="0" width="100%" style="margin:0 0 24px 0;border:1px solid #E5E7EB;border-radius:16px;overflow:hidden">
          <tr><td>
            <img src="{img}" alt="" width="600" style="width:100%;max-width:600px;display:block;object-fit:cover;height:220px" />
          </td></tr>
          <tr><td style="padding:16px 20px">
            <div style="font-size:11px;font-weight:700;letter-spacing:0.1em;text-transform:uppercase;color:{brand}">{when}</div>
            <div style="font-size:20px;font-weight:800;margin-top:4px;color:#111827">{_html_lib.escape(e['title'])}</div>
            <div style="font-size:14px;color:#6B7280;margin-top:4px">{_html_lib.escape(e['venue'])}</div>
            <p style="font-size:14px;color:#374151;line-height:1.5;margin:12px 0 16px 0">{_html_lib.escape(e['description'])}</p>
            <div>
              <a href="{e['share_links']['facebook']}" style="display:inline-block;padding:8px 14px;margin:0 6px 6px 0;background:#1877F2;color:#fff;border-radius:999px;font-size:12px;font-weight:700;text-decoration:none">Facebook</a>
              <a href="{e['share_links']['linkedin']}" style="display:inline-block;padding:8px 14px;margin:0 6px 6px 0;background:#0A66C2;color:#fff;border-radius:999px;font-size:12px;font-weight:700;text-decoration:none">LinkedIn</a>
              <a href="{e['share_links']['twitter']}" style="display:inline-block;padding:8px 14px;margin:0 6px 6px 0;background:#111;color:#fff;border-radius:999px;font-size:12px;font-weight:700;text-decoration:none">X</a>
              <a href="{e['share_links']['whatsapp']}" style="display:inline-block;padding:8px 14px;margin:0 6px 6px 0;background:#25D366;color:#fff;border-radius:999px;font-size:12px;font-weight:700;text-decoration:none">WhatsApp</a>
              <a href="{e['canonical_url']}" style="display:inline-block;padding:8px 14px;margin:0 6px 6px 0;background:#F3F4F6;color:#111;border-radius:999px;font-size:12px;font-weight:700;text-decoration:none">Event page</a>
            </div>
          </td></tr>
        </table>
        """)
    body = "\n".join(rows)
    return f"""<!doctype html>
<html><body style="margin:0;padding:0;background:#F9FAFB;font-family:system-ui,-apple-system,Segoe UI,sans-serif">
<table role="presentation" cellspacing="0" cellpadding="0" width="100%" bgcolor="#F9FAFB">
<tr><td align="center">
<table role="presentation" cellspacing="0" cellpadding="0" width="640" style="max-width:640px;padding:24px">
  <tr><td style="text-align:center;padding:12px 0 24px 0">
    <div style="font-size:12px;letter-spacing:0.2em;text-transform:uppercase;color:{brand};font-weight:800">Blackrod Now · Share Pack</div>
    <h1 style="font-size:28px;color:#111827;margin:8px 0 4px 0">Your upcoming events</h1>
    <p style="font-size:15px;color:#6B7280;margin:0">Ready-to-share posts for {_html_lib.escape(org['name'])}. Tap a button below any event to share.</p>
  </td></tr>
  <tr><td>{body}</td></tr>
  <tr><td style="text-align:center;padding:24px 0;color:#9CA3AF;font-size:12px">
    Sent from Blackrod Now · <a href="{base}/organisation-dashboard" style="color:{brand}">Open dashboard</a>
  </td></tr>
</table>
</td></tr>
</table>
</body></html>"""


@api.get("/organisations/{slug}/share-pack")
async def get_share_pack(slug: str, request: _Req):
    org = await _find_org(slug)
    all_events = await db.events.find({"status": {"$ne": "pending"}}, {"_id": 0}).to_list(2000)
    base = _abs_base_url(request)
    return _org_share_pack_data(org, all_events, base)


class SharePackEmailReq(BaseModel):
    to: Optional[str] = None  # override; else uses org.email


@api.post("/organisations/{slug}/share-pack/email")
async def email_share_pack(
    slug: str,
    req: SharePackEmailReq,
    request: _Req,
    org_auth: Optional[str] = Header(None, alias="X-Org-Auth"),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_org_write_access(slug, org_auth, admin_code)
    org = await _find_org(slug)
    to = (req.to or org.get("email") or "").strip()
    if not to:
        raise HTTPException(400, "No recipient email — set an email on the org profile or pass `to`.")
    all_events = await db.events.find({"status": {"$ne": "pending"}}, {"_id": 0}).to_list(2000)
    base = _abs_base_url(request)
    pack = _org_share_pack_data(org, all_events, base)
    html = _render_share_pack_html(pack, base)
    subject = f"Your Blackrod Now share pack — {pack['count']} upcoming event{'s' if pack['count'] != 1 else ''}"
    result = await asyncio.to_thread(resend_send, to, subject, html)
    await _record_analytics_event("share_pack_email", entity_type="org", org_slug=slug, count=1)
    return {"ok": True, "to": to, "count": pack["count"], "email": result}


@api.post("/events")
async def create_event(
    evt: Event,
    org_auth: Optional[str] = Header(None, alias="X-Org-Auth"),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    auth_role = "org"
    authorized = True
    try:
        auth_role = _require_org_write_access(evt.orgSlug, org_auth, admin_code)
    except HTTPException:
        authorized = False

    if not authorized:
        # Public submissions stay allowed, but always as pending and never featured.
        await _find_org(evt.orgSlug)
        evt.status = "pending"
        evt.featured = False
    elif auth_role == "admin":
        if evt.status not in ("pending", "approved", "rejected", "draft", "cancelled"):
            evt.status = "approved"
    else:
        org = await _find_org(evt.orgSlug)
        trust = (org.get("trust_level") or "new").strip().lower()
        if trust == "trusted":
            evt.status = "approved"
        else:
            evt.status = "pending"
        evt.featured = False

    if evt.status not in ("pending", "approved", "rejected", "draft", "cancelled"):
        evt.status = "pending"
    if _is_blank_or_legacy_event_image(evt.image):
        evt.image = _event_category_image(evt.category)
    await db.events.insert_one(evt.model_dump())
    await _audit(
        action="event_created",
        entity_type="event",
        entity_id=evt.id,
        summary=f"Event created: {evt.title}",
        meta={"org_slug": evt.orgSlug, "status": evt.status},
        actor=("admin" if auth_role == "admin" and authorized else "org" if authorized else "public"),
    )
    return evt


class EventPatch(BaseModel):
    """Partial event update. Any subset of fields may be supplied."""
    model_config = ConfigDict(extra="ignore")
    title: Optional[str] = None
    orgSlug: Optional[str] = None
    category: Optional[str] = None
    start: Optional[str] = None
    end: Optional[str] = None
    venue: Optional[str] = None
    address: Optional[str] = None
    description: Optional[str] = None
    cost: Optional[str] = None
    age: Optional[str] = None
    accessibility: Optional[str] = None
    booking: Optional[str] = None
    contactEmail: Optional[str] = None
    contactPhone: Optional[str] = None
    image: Optional[str] = None
    featured: Optional[bool] = None
    status: Optional[Literal["approved", "pending", "rejected", "draft", "cancelled"]] = None
    recurrence: Optional[EventRecurrence] = None


@api.patch("/events/{event_id}")
async def update_event(
    event_id: str,
    patch: EventPatch,
    org_auth: Optional[str] = Header(None, alias="X-Org-Auth"),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    existing = await db.events.find_one({"id": event_id})
    if not existing:
        raise HTTPException(404, "Event not found")
    auth_role = _require_org_write_access(existing.get("orgSlug") or "", org_auth, admin_code)
    updates = patch.model_dump(exclude_none=True)
    if not updates:
        return await get_event(event_id)
    if updates.get("orgSlug"):
        # Ensure the target org exists
        await _find_org(updates["orgSlug"])
        if auth_role != "admin" and updates["orgSlug"] != existing.get("orgSlug"):
            raise HTTPException(403, "Organisation admins cannot move events between organisations")
    if "image" in updates and _is_blank_or_legacy_event_image(updates.get("image") or ""):
        updates["image"] = _event_category_image(updates.get("category") or existing.get("category") or "")
    elif "category" in updates and _is_blank_or_legacy_event_image(existing.get("image") or ""):
        updates["image"] = _event_category_image(updates.get("category") or "")
    await db.events.update_one({"id": event_id}, {"$set": updates})
    await _audit(
        action="event_updated",
        entity_type="event",
        entity_id=event_id,
        summary=f"Updated event: {updates.get('title') or existing.get('title') or event_id}",
        meta={"fields": sorted(list(updates.keys()))},
        actor=("admin" if auth_role == "admin" else "org"),
    )
    return await get_event(event_id)


@api.post("/admin/events/{event_id}/status")
async def admin_event_status(event_id: str, body: Dict[str, str]):
    status = body.get("status", "approved")
    await db.events.update_one({"id": event_id}, {"$set": {"status": status}})
    await _audit(
        action="event_status_changed",
        entity_type="event",
        entity_id=event_id,
        summary=f"Event status set to {status}",
        meta={"status": status},
    )
    return await get_event(event_id)


@api.post("/admin/events/{event_id}/feature")
async def admin_feature_event(event_id: str):
    e = await db.events.find_one({"id": event_id})
    if not e:
        raise HTTPException(404, "Event not found")
    featured = not e.get("featured", False)
    await db.events.update_one({"id": event_id}, {"$set": {"featured": featured}})
    await _audit(
        action="event_feature_toggled",
        entity_type="event",
        entity_id=event_id,
        summary=("Featured event" if featured else "Unfeatured event"),
    )
    return await get_event(event_id)


@api.delete("/admin/events/{event_id}")
async def admin_delete_event(event_id: str):
    await db.events.delete_one({"id": event_id})
    await _audit(
        action="event_deleted",
        entity_type="event",
        entity_id=event_id,
        summary="Deleted event",
    )
    return {"ok": True}


# ─────────── Feed ───────────
@api.get("/feed")
async def list_feed():
    return await db.feed.find({}, {"_id": 0}).sort("time", -1).to_list(500)


@api.post("/feed")
async def create_feed_post(
    post: FeedPost,
    org_auth: Optional[str] = Header(None, alias="X-Org-Auth"),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_org_write_access(post.orgSlug, org_auth, admin_code)
    await db.feed.insert_one(post.model_dump())
    return post


# ─────────── Venues & Volunteers ───────────
@api.get("/venues")
async def list_venues():
    return await db.venues.find({}, {"_id": 0}).to_list(200)


class VenuePatch(BaseModel):
    name: Optional[str] = None
    address: Optional[str] = None
    facilities: Optional[List[str]] = None
    accessibility: Optional[str] = None
    capacity: Optional[int] = None
    booking: Optional[str] = None
    image: Optional[str] = None


@api.post("/venues")
async def create_venue(venue: Venue):
    existing = await db.venues.find_one({"name": {"$regex": f"^{re.escape(venue.name)}$", "$options": "i"}}, {"_id": 0})
    if existing:
        raise HTTPException(409, "Venue with this name already exists")
    await db.venues.insert_one(venue.model_dump())
    return venue


@api.patch("/venues/{venue_id}")
async def update_venue(venue_id: str, patch: VenuePatch):
    existing = await db.venues.find_one({"id": venue_id}, {"_id": 0})
    if not existing:
        raise HTTPException(404, "Venue not found")
    updates = patch.model_dump(exclude_none=True)
    if not updates:
        return existing
    await db.venues.update_one({"id": venue_id}, {"$set": updates})
    updated = await db.venues.find_one({"id": venue_id}, {"_id": 0})
    if not updated:
        raise HTTPException(404, "Venue not found")
    return updated


@api.get("/volunteers")
async def list_volunteers():
    return await db.volunteers.find({}, {"_id": 0}).to_list(200)


@api.post("/volunteers")
async def create_volunteer(opp: VolunteerOpp):
    await _find_org(opp.orgSlug)
    await db.volunteers.insert_one(opp.model_dump())
    return opp


class VolunteerPatch(BaseModel):
    title: Optional[str] = None
    orgSlug: Optional[str] = None
    description: Optional[str] = None
    age: Optional[str] = None
    time: Optional[str] = None
    skills: Optional[str] = None


@api.patch("/volunteers/{vol_id}")
async def update_volunteer(vol_id: str, patch: VolunteerPatch):
    existing = await db.volunteers.find_one({"id": vol_id}, {"_id": 0})
    if not existing:
        raise HTTPException(404, "Volunteer opportunity not found")
    updates = patch.model_dump(exclude_none=True)
    if not updates:
        return existing
    if updates.get("orgSlug"):
        await _find_org(updates["orgSlug"])
    await db.volunteers.update_one({"id": vol_id}, {"$set": updates})
    updated = await db.volunteers.find_one({"id": vol_id}, {"_id": 0})
    if not updated:
        raise HTTPException(404, "Volunteer opportunity not found")
    return updated


# ─────────── Auth (JWT — admin + org) ───────────
from auth import (  # noqa: E402
    hash_password,
    verify_password,
    create_access_token,
    decode_token,
    read_admin_from_request,
)


class AdminLoginReq(BaseModel):
    email: EmailStr
    password: str


LOGIN_MAX_ATTEMPTS = int(os.environ.get("LOGIN_MAX_ATTEMPTS", "5"))
LOGIN_LOCKOUT_MINUTES = int(os.environ.get("LOGIN_LOCKOUT_MINUTES", "15"))


async def _record_login_attempt(identifier: str, success: bool) -> None:
    now = datetime.now(timezone.utc)
    if success:
        await db.login_attempts.delete_many({"identifier": identifier})
        return
    doc = await db.login_attempts.find_one({"identifier": identifier})
    if not doc:
        await db.login_attempts.insert_one({
            "identifier": identifier,
            "count": 1,
            "first_at": now.isoformat(),
            "last_at": now.isoformat(),
        })
    else:
        await db.login_attempts.update_one(
            {"identifier": identifier},
            {"$inc": {"count": 1}, "$set": {"last_at": now.isoformat()}},
        )


async def _login_is_locked(identifier: str) -> bool:
    doc = await db.login_attempts.find_one({"identifier": identifier})
    if not doc or (doc.get("count") or 0) < LOGIN_MAX_ATTEMPTS:
        return False
    try:
        last = datetime.fromisoformat(doc["last_at"])
    except Exception:
        return False
    if last.tzinfo is None:
        last = last.replace(tzinfo=timezone.utc)
    if datetime.now(timezone.utc) - last > timedelta(minutes=LOGIN_LOCKOUT_MINUTES):
        # Lockout expired — clear counter.
        await db.login_attempts.delete_many({"identifier": identifier})
        return False
    return True


@api.post("/auth/admin/login")
async def auth_admin_login(req: AdminLoginReq, request: Request):
    email = req.email.lower().strip()
    ip = request.client.host if request.client else "unknown"
    identifier = f"{ip}:{email}"
    if await _login_is_locked(identifier):
        raise HTTPException(429, "Too many failed attempts. Try again in a few minutes.")
    user = await db.users.find_one({"email": email, "role": "admin"})
    if not user or not verify_password(req.password, user.get("password_hash", "")):
        await _record_login_attempt(identifier, False)
        raise HTTPException(401, "Invalid email or password")
    await _record_login_attempt(identifier, True)
    token = create_access_token(sub=email, role="admin", extra={"email": email, "name": user.get("name") or "Admin"})
    return {
        "token": token,
        "user": {"email": email, "role": "admin", "name": user.get("name") or "Admin"},
    }


@api.get("/auth/me")
async def auth_me(request: Request, authorization: Optional[str] = Header(None)):
    payload = read_admin_from_request(request, authorization)
    if not payload:
        raise HTTPException(401, "Not authenticated")
    return {
        "email": payload.get("email") or payload.get("sub"),
        "role": payload.get("role"),
        "name": payload.get("name"),
        "exp": payload.get("exp"),
    }


async def _seed_admin_user() -> None:
    email = (os.environ.get("ADMIN_EMAIL") or "").lower().strip()
    password = os.environ.get("ADMIN_PASSWORD") or ""
    if not email or not password:
        logger.warning("Admin seed skipped: ADMIN_EMAIL / ADMIN_PASSWORD not set.")
        return
    existing = await db.users.find_one({"email": email})
    if not existing:
        await db.users.insert_one({
            "email": email,
            "password_hash": hash_password(password),
            "role": "admin",
            "name": "Site Admin",
            "created_at": datetime.now(timezone.utc).isoformat(),
        })
        logger.info("Seeded admin user %s", email)
    elif not verify_password(password, existing.get("password_hash", "")):
        await db.users.update_one(
            {"email": email},
            {"$set": {
                "password_hash": hash_password(password),
                "role": "admin",
                "updated_at": datetime.now(timezone.utc).isoformat(),
            }},
        )
        logger.info("Rotated admin password for %s", email)


# Update: impersonation endpoint now accepts either an admin JWT (preferred)
# or the legacy launch code (backward compat).


# ─────────── Funder Impact Dashboard ───────────
_UK_POSTCODE_RE = re.compile(
    r"\b([A-PR-UWYZ][A-HK-Y]?\d[A-Z0-9]?)\s*(\d[A-Z]{2})?\b", re.IGNORECASE
)


def _extract_postcode(text: str) -> Optional[str]:
    """Return the outward code (e.g. 'BL6') of the first UK postcode found."""
    if not text:
        return None
    m = _UK_POSTCODE_RE.search(text)
    if not m:
        return None
    outward = m.group(1)
    return outward.upper() if outward else None


async def _get_grant_config() -> Dict[str, Any]:
    doc = await db.site_settings.find_one({"_id": _SITE_SETTINGS_ID}) or {}
    return {
        "grant_amount": float(doc.get("grant_amount", 0) or 0),
        "grant_currency": doc.get("grant_currency", "GBP"),
        "grant_period_label": doc.get("grant_period_label", "annual"),
    }


async def _build_impact_snapshot(days: int = 90) -> Dict[str, Any]:
    now = datetime.now(timezone.utc)
    since = (now - timedelta(days=max(1, days))).isoformat()

    # Unique residents engaged = unique device_ids from analytics + unique subscriber emails.
    device_pipeline = [
        {"$match": {"created_at": {"$gte": since}, "device_id": {"$ne": None}}},
        {"$group": {"_id": "$device_id"}},
        {"$count": "n"},
    ]
    device_rows = await db.analytics_events.aggregate(device_pipeline).to_list(1)
    unique_devices = device_rows[0]["n"] if device_rows else 0

    sub_count = await db.subscribers.count_documents({"unsubscribed": {"$ne": True}})
    follow_count = await db.follows.count_documents({})
    unique_residents = max(unique_devices, sub_count, follow_count)

    # Volunteer conversions (contact-button clicks on volunteering cards).
    volunteer_clicks = await db.analytics_events.count_documents(
        {"kind": "volunteer_contact", "created_at": {"$gte": since}}
    )

    # Cross-org collaboration proxy: count of subscribers who follow >1 org.
    multi_follow = await db.subscribers.count_documents(
        {"unsubscribed": {"$ne": True}, "followed_orgs.1": {"$exists": True}}
    )

    # Retention proxy: subscribers with digest=true & last 60 days.
    active_subs = await db.subscribers.count_documents(
        {"unsubscribed": {"$ne": True}, "digest": {"$ne": False}}
    )

    # Total content served
    total_events_live = await db.events.count_documents({"status": "approved"})
    total_orgs_live = await db.orgs.count_documents({"status": "approved"})
    total_venues = await db.venues.count_documents({})
    total_volunteer_opps = await db.volunteers.count_documents({})

    # Reach — sum of analytics counts by kind
    reach_pipeline = [
        {"$match": {"created_at": {"$gte": since}}},
        {"$group": {"_id": "$kind", "n": {"$sum": "$count"}}},
    ]
    reach_rows = await db.analytics_events.aggregate(reach_pipeline).to_list(20)
    reach = {row["_id"]: row["n"] for row in reach_rows}

    # Geographic reach — pull outward postcodes from all approved orgs + events.
    postcode_counter: Counter = Counter()
    orgs_all = await db.orgs.find(
        {"status": "approved"}, {"_id": 0, "slug": 1, "name": 1, "address": 1, "location": 1}
    ).to_list(500)
    for o in orgs_all:
        pc = _extract_postcode(o.get("address") or "") or _extract_postcode(o.get("location") or "")
        if pc:
            postcode_counter[pc] += 1
    events_all = await db.events.find(
        {"status": "approved"}, {"_id": 0, "id": 1, "address": 1, "venue": 1}
    ).to_list(1000)
    for ev in events_all:
        pc = _extract_postcode(ev.get("address") or "") or _extract_postcode(ev.get("venue") or "")
        if pc:
            postcode_counter[pc] += 1

    geo = [
        {"postcode": pc, "count": count}
        for pc, count in postcode_counter.most_common(30)
    ]

    # Top orgs by follower + reach
    org_follow_pipeline = [
        {"$match": {"unsubscribed": {"$ne": True}}},
        {"$unwind": "$followed_orgs"},
        {"$group": {"_id": "$followed_orgs", "n": {"$sum": 1}}},
        {"$sort": {"n": -1}},
        {"$limit": 10},
    ]
    org_follow_rows = await db.subscribers.aggregate(org_follow_pipeline).to_list(10)
    org_lookup = {o.get("slug"): o.get("name") for o in orgs_all}
    top_orgs = [
        {"slug": row["_id"], "name": org_lookup.get(row["_id"], row["_id"]), "followers": row["n"]}
        for row in org_follow_rows
    ]

    # Cost per resident
    grant = await _get_grant_config()
    cost_per_resident = (
        (grant["grant_amount"] / unique_residents) if grant["grant_amount"] and unique_residents else 0
    )

    # SROI-ish proxy: assume 1 volunteer_contact = 4 hours saved, £15/hour standard multiplier.
    volunteer_hours_est = volunteer_clicks * 4
    volunteer_value_est = volunteer_hours_est * 15

    return {
        "window_days": days,
        "generated_at": now.isoformat(),
        "grant": grant,
        "headline": {
            "unique_residents": unique_residents,
            "subscribers": sub_count,
            "device_follows": follow_count,
            "cost_per_resident": round(cost_per_resident, 2),
            "orgs_live": total_orgs_live,
            "events_live": total_events_live,
            "venues": total_venues,
            "volunteer_opps": total_volunteer_opps,
            "volunteer_conversions": volunteer_clicks,
            "volunteer_hours_estimated": volunteer_hours_est,
            "volunteer_value_estimated": volunteer_value_est,
            "active_digest_subscribers": active_subs,
            "cross_org_engagement": multi_follow,
        },
        "reach": {
            "org_views": reach.get("org_view", 0),
            "event_views": reach.get("event_view", 0),
            "share_clicks": reach.get("share_click", 0),
            "volunteer_clicks": reach.get("volunteer_contact", 0),
        },
        "geography": geo,
        "top_orgs": top_orgs,
    }


@api.get("/admin/impact/summary")
async def admin_impact_summary(
    days: int = 90,
    request: Request = None,  # type: ignore
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    ok = False
    if request and read_admin_from_request(request, authorization):
        ok = True
    elif admin_code and hmac.compare_digest(admin_code, ADMIN_LAUNCH_CODE):
        ok = True
    if not ok:
        raise HTTPException(403, "Admin authentication required")
    return await _build_impact_snapshot(max(1, min(int(days), 730)))


class GrantConfigPatch(BaseModel):
    grant_amount: Optional[float] = None
    grant_currency: Optional[str] = None
    grant_period_label: Optional[str] = None


@api.post("/admin/impact/grant-config")
async def set_grant_config(
    patch: GrantConfigPatch,
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    ok = False
    if read_admin_from_request(request, authorization):
        ok = True
    elif admin_code and hmac.compare_digest(admin_code, ADMIN_LAUNCH_CODE):
        ok = True
    if not ok:
        raise HTTPException(403, "Admin authentication required")
    updates = {k: v for k, v in patch.model_dump(exclude_none=True).items()}
    if updates:
        updates["updated_at"] = now_iso()
        await db.site_settings.update_one({"_id": _SITE_SETTINGS_ID}, {"$set": updates}, upsert=True)
    return await _get_grant_config()


# PDF report ------------------------------------------------------------------
def _build_impact_pdf(snapshot: Dict[str, Any], variant: str) -> bytes:
    from io import BytesIO
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak,
    )

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=18 * mm, leftMargin=18 * mm,
        topMargin=18 * mm, bottomMargin=18 * mm,
        title="Blackrod Now — Impact Report",
    )
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], textColor=colors.HexColor("#0052FF"), fontSize=22, spaceAfter=8)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], textColor=colors.HexColor("#0F172A"), fontSize=14, spaceBefore=14, spaceAfter=6)
    body = ParagraphStyle("body", parent=styles["BodyText"], fontSize=10, leading=14, textColor=colors.HexColor("#334155"))
    small = ParagraphStyle("small", parent=styles["BodyText"], fontSize=8, textColor=colors.HexColor("#64748B"))
    tile_val = ParagraphStyle("tile_val", parent=styles["Heading1"], fontSize=20, textColor=colors.HexColor("#0052FF"), alignment=1)
    tile_label = ParagraphStyle("tile_label", parent=styles["BodyText"], fontSize=8, textColor=colors.HexColor("#64748B"), alignment=1)

    head = snapshot["headline"]
    reach = snapshot["reach"]
    grant = snapshot["grant"]
    cur = "£" if grant["grant_currency"] == "GBP" else grant["grant_currency"] + " "
    story = []

    story.append(Paragraph("Blackrod Now — Impact Report", h1))
    story.append(Paragraph(
        f"Reporting window: last {snapshot['window_days']} days · Generated {snapshot['generated_at'][:10]}",
        small,
    ))
    story.append(Spacer(1, 8))

    # Executive summary
    story.append(Paragraph("Executive summary", h2))
    exec_lines = [
        f"<b>{head['unique_residents']:,}</b> unique Blackrod residents engaged with the platform in the last {snapshot['window_days']} days.",
        f"<b>{head['orgs_live']}</b> community organisations are live on Blackrod Now, promoting <b>{head['events_live']}</b> events.",
        f"<b>{head['volunteer_conversions']}</b> volunteering contact clicks — an estimated <b>{head['volunteer_hours_estimated']}</b> volunteer hours (≈ {cur}{head['volunteer_value_estimated']:,.0f} in social value at £15/hour).",
    ]
    if grant["grant_amount"] and head["unique_residents"]:
        exec_lines.append(
            f"<b>{cur}{head['cost_per_resident']:.2f} per resident engaged</b>, against a grant of {cur}{grant['grant_amount']:,.0f} ({grant['grant_period_label']})."
        )
    for line in exec_lines:
        story.append(Paragraph("• " + line, body))
    story.append(Spacer(1, 8))

    # Big-number tiles
    def tile(v, label):
        cell = [
            Paragraph(str(v), tile_val),
            Paragraph(label, tile_label),
        ]
        return cell

    tiles_row1 = [
        tile(f"{head['unique_residents']:,}", "Unique residents"),
        tile(f"{head['orgs_live']}", "Organisations live"),
        tile(f"{head['events_live']}", "Events published"),
        tile(f"{head['volunteer_conversions']}", "Volunteer signups"),
    ]
    tiles_row2 = [
        tile(f"{reach['event_views']:,}", "Event views"),
        tile(f"{reach['org_views']:,}", "Org views"),
        tile(f"{reach['share_clicks']:,}", "Shares"),
        tile(f"{cur}{head['cost_per_resident']:.2f}", "Cost per resident"),
    ]
    for row in (tiles_row1, tiles_row2):
        table = Table([row], colWidths=[42 * mm] * 4, rowHeights=[24 * mm])
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F1F5F9")),
            ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0")),
            ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        story.append(table)
        story.append(Spacer(1, 6))

    # Short variant stops here
    if variant == "short":
        story.append(Spacer(1, 12))
        story.append(Paragraph(
            "This is a summary view. For the full report with geographic breakdown, top organisations and methodology notes, download the detailed version from the Impact dashboard.",
            small,
        ))
        doc.build(story)
        return buf.getvalue()

    # Full variant continues -----------------------------------------------
    story.append(PageBreak())

    story.append(Paragraph("Reach breakdown", h2))
    reach_rows = [
        ["Metric", "Count"],
        ["Event page views", f"{reach['event_views']:,}"],
        ["Organisation page views", f"{reach['org_views']:,}"],
        ["Share button clicks", f"{reach['share_clicks']:,}"],
        ["Volunteering contact clicks", f"{reach['volunteer_clicks']:,}"],
        ["Newsletter subscribers (active)", f"{head['active_digest_subscribers']:,}"],
        ["Cross-org engagement (follow >1 org)", f"{head['cross_org_engagement']:,}"],
    ]
    t = Table(reach_rows, colWidths=[110 * mm, 60 * mm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0052FF")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
        ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E1")),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(t)

    # Geography
    story.append(Paragraph("Geographic reach (by outward postcode)", h2))
    geo = snapshot["geography"]
    if geo:
        rows = [["Postcode area", "Listings"]] + [[g["postcode"], str(g["count"])] for g in geo[:20]]
        gt = Table(rows, colWidths=[80 * mm, 40 * mm])
        gt.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0052FF")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
            ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
            ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E1")),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
        ]))
        story.append(gt)
    else:
        story.append(Paragraph("No postcode data recorded yet. Encourage organisations to include a postcode in their venue / address fields.", body))

    # Top orgs
    story.append(Paragraph("Top organisations by follower count", h2))
    if snapshot["top_orgs"]:
        rows = [["Organisation", "Followers"]] + [[o["name"], str(o["followers"])] for o in snapshot["top_orgs"]]
        ot = Table(rows, colWidths=[130 * mm, 40 * mm])
        ot.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0052FF")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
            ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
            ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E1")),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
        ]))
        story.append(ot)
    else:
        story.append(Paragraph("No follower data recorded yet.", body))

    # Methodology
    story.append(Paragraph("Methodology &amp; notes", h2))
    methodology = [
        "<b>Unique residents</b> is the maximum of unique tracked devices, unique subscribed emails, and unique device-follow records — an upper bound guard against under-counting.",
        "<b>Volunteer conversions</b> counts clicks on volunteering opportunity contact buttons. Volunteer hours are estimated at 4 hours per contact (industry proxy) valued at £15/hour (national volunteer social value multiplier).",
        "<b>Cost per resident</b> = grant amount ÷ unique residents engaged in this window. Grant amount is configured by the site administrator.",
        "<b>Geographic reach</b> is derived from UK postcode outward codes found in organisation and event addresses. Fine-grained locations are never stored or displayed beyond the outward code.",
        "<b>Cross-org engagement</b> is the count of newsletter subscribers who follow more than one organisation — a proxy for cross-community discovery, a stated aim of the Community Alliance Fund.",
    ]
    for line in methodology:
        story.append(Paragraph("• " + line, body))

    doc.build(story)
    return buf.getvalue()


@api.get("/admin/impact/pdf")
async def admin_impact_pdf(
    days: int = 90,
    variant: str = "full",
    request: Request = None,  # type: ignore
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
    token: Optional[str] = None,
):
    """PDF impact report. Accepts admin JWT via header OR ?token= query param
    (needed because <a href download> can't set Authorization headers)."""
    ok = False
    if request and read_admin_from_request(request, authorization):
        ok = True
    elif token:
        try:
            payload = decode_token(token)
            if payload.get("role") == "admin":
                ok = True
        except Exception:
            pass
    elif admin_code and hmac.compare_digest(admin_code, ADMIN_LAUNCH_CODE):
        ok = True
    if not ok:
        raise HTTPException(403, "Admin authentication required")
    if variant not in ("short", "full"):
        variant = "full"
    snapshot = await _build_impact_snapshot(max(1, min(int(days), 730)))
    pdf_bytes = _build_impact_pdf(snapshot, variant)
    filename = f"blackrod-now-impact-{variant}-{snapshot['generated_at'][:10]}.pdf"
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ─────────── Auto-generated poster (per event) ───────────
def _draw_event_poster_png(event: dict, size: int = 1080, site_url: Optional[str] = None) -> bytes:
    from io import BytesIO
    from PIL import Image, ImageDraw, ImageFont
    import qrcode

    W = H = size
    BLUE = (0, 82, 255)
    LIME = (210, 255, 0)
    INK = (10, 10, 24)

    img = Image.new("RGB", (W, H), INK)
    draw = ImageDraw.Draw(img)

    # Background gradient (blue → dark)
    for y in range(H):
        t = y / H
        r = int(BLUE[0] * (1 - t) + 5 * t)
        g = int(BLUE[1] * (1 - t) + 8 * t)
        b = int(BLUE[2] * (1 - t) + 26 * t)
        draw.line([(0, y), (W, y)], fill=(r, g, b))

    def _font(sz):
        try:
            return ImageFont.truetype("/etc/alternatives/fonts-japanese-gothic.ttf", sz)
        except Exception:
            return ImageFont.load_default(size=sz) if hasattr(ImageFont, "load_default") else ImageFont.load_default()

    pad = 72
    # Brand strip
    draw.rectangle([pad, pad, pad + 220, pad + 44], fill=LIME)
    draw.text((pad + 16, pad + 8), "BLACKROD NOW", fill=INK, font=_font(22))

    # Category badge
    cat = (event.get("category") or "Community").upper()
    draw.text((pad, pad + 84), cat, fill=LIME, font=_font(22))

    # Title (wrap)
    title = (event.get("title") or "").strip()[:110]
    title_font = _font(76)
    words = title.split()
    lines: List[str] = []
    cur = ""
    max_w = W - 2 * pad
    for w in words:
        candidate = (cur + " " + w).strip()
        if draw.textlength(candidate, font=title_font) <= max_w:
            cur = candidate
        else:
            if cur:
                lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    y = pad + 140
    for ln in lines[:4]:
        draw.text((pad, y), ln, fill="white", font=title_font)
        y += 88

    # Date + time
    try:
        start = datetime.fromisoformat((event.get("start") or "").replace("Z", "+00:00"))
        when = start.strftime("%A %d %B · %H:%M")
    except Exception:
        when = event.get("start", "")
    y += 20
    draw.text((pad, y), when, fill=LIME, font=_font(36))
    y += 56

    # Venue
    venue = event.get("venue") or ""
    if venue:
        draw.text((pad, y), venue, fill="white", font=_font(32))
        y += 46
    address = event.get("address") or ""
    if address:
        draw.text((pad, y), address, fill=(200, 200, 220), font=_font(24))
        y += 44

    # Cost / Age chips
    chips = []
    if event.get("cost"): chips.append(event["cost"])
    if event.get("age"): chips.append(event["age"])
    if event.get("accessibility"): chips.append(event["accessibility"][:40])
    chip_x = pad
    chip_y = H - 220
    for chip in chips[:3]:
        w = int(draw.textlength(chip, font=_font(24))) + 40
        draw.rounded_rectangle([chip_x, chip_y, chip_x + w, chip_y + 44], radius=22, fill=(255, 255, 255, 40), outline=LIME, width=2)
        draw.text((chip_x + 20, chip_y + 8), chip, fill="white", font=_font(24))
        chip_x += w + 12

    # QR code linking to the event page
    site_url = (site_url or os.environ.get("PUBLIC_URL") or PUBLIC_URL).rstrip("/")
    event_url = f"{site_url}/events/{event.get('id')}" if event.get("id") else site_url
    qr = qrcode.QRCode(border=2, box_size=8)
    qr.add_data(event_url)
    qr.make(fit=True)
    qr_img = qr.make_image(fill_color="white", back_color=(0, 82, 255)).convert("RGB")
    qr_img = qr_img.resize((176, 176))
    img.paste(qr_img, (W - pad - 176, H - pad - 176))
    draw.text((W - pad - 176, H - pad - 200), "SCAN TO OPEN", fill=LIME, font=_font(18))

    # Footer
    draw.text((pad, H - pad - 40), "blackrodnow.co.uk · Made in Blackrod", fill=(220, 220, 240), font=_font(20))

    buf = BytesIO()
    img.save(buf, format="PNG", optimize=True)
    return buf.getvalue()


def _draw_event_poster_pdf(event: dict, site_url: Optional[str] = None) -> bytes:
    """A4 poster — reuse the PNG as a full-bleed image inside a PDF."""
    from io import BytesIO
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import ImageReader
    from PIL import Image
    png = _draw_event_poster_png(event, size=1600, site_url=site_url)
    img = Image.open(BytesIO(png))
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    W, H = A4
    # Fit the square poster into A4 (centre it).
    side = min(W, H) - 30
    x = (W - side) / 2
    y = (H - side) / 2
    c.drawImage(ImageReader(img), x, y, width=side, height=side)
    c.showPage()
    c.save()
    return buf.getvalue()


async def _fetch_event_for_poster(event_id: str) -> dict:
    if "__" in event_id:
        parent_id, _ = event_id.split("__", 1)
        parent = await db.events.find_one({"id": parent_id}, {"_id": 0})
        if not parent:
            raise HTTPException(404, "Event not found")
        for inst in _expand_recurring_event(parent):
            if inst.get("id") == event_id:
                return inst
        raise HTTPException(404, "Event instance not found")
    ev = await db.events.find_one({"id": event_id}, {"_id": 0})
    if not ev:
        raise HTTPException(404, "Event not found")
    return ev


def _safe_ascii_filename(title: Optional[str], ext: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", (title or "event").encode("ascii", "ignore").decode().lower()).strip("-")[:40]
    return f"blackrod-now-{slug or 'event'}.{ext}"


@api.get("/events/{event_id}/poster.png")
async def event_poster_png(event_id: str, request: Request):
    ev = await _fetch_event_for_poster(event_id)
    png = await asyncio.to_thread(_draw_event_poster_png, ev, 1080, _abs_base_url(request))
    fname = _safe_ascii_filename(ev.get("title"), "png")
    return Response(
        content=png,
        media_type="image/png",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


@api.get("/events/{event_id}/poster.pdf")
async def event_poster_pdf(event_id: str, request: Request):
    ev = await _fetch_event_for_poster(event_id)
    pdf = await asyncio.to_thread(_draw_event_poster_pdf, ev, _abs_base_url(request))
    fname = _safe_ascii_filename(ev.get("title"), "pdf")
    return Response(
        content=pdf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


# ─────────── "Post Now" social bundle (caption + hashtags) ───────────
_TONE_CHOICES = ("friendly", "punchy", "formal")


def _humanize_event_when(ev: dict) -> str:
    """Human-friendly date/time line for an event. Uses UK-style formatting."""
    start = ev.get("start") or ev.get("date")
    if not start:
        return ""
    try:
        # Support ISO strings with or without timezone
        s = start.replace("Z", "+00:00") if isinstance(start, str) else str(start)
        dt = datetime.fromisoformat(s)
    except Exception:
        return str(start)
    # UK style, e.g. "Sat 14 Jun · 11:00"
    return dt.strftime("%a %d %b · %H:%M").lstrip("0")


def _event_public_link(ev: dict, base: Optional[str] = None) -> str:
    return f"{(base or PUBLIC_URL).rstrip('/')}/events/{ev.get('id')}"


def _event_hashtags(ev: dict) -> List[str]:
    tags = ["#Blackrod", "#BlackrodNow"]
    cat = (ev.get("category") or "").strip()
    if cat:
        cat_tag = "#" + re.sub(r"[^A-Za-z0-9]+", "", cat.title())
        if cat_tag != "#":
            tags.append(cat_tag)
    venue = (ev.get("venue") or "").strip()
    if venue and len(venue) < 30:
        v_tag = "#" + re.sub(r"[^A-Za-z0-9]+", "", venue.title())
        if v_tag != "#" and v_tag.lower() not in [t.lower() for t in tags]:
            tags.append(v_tag)
    tags.append("#Community")
    # Dedupe (case-insensitive) while preserving order
    seen, out = set(), []
    for t in tags:
        k = t.lower()
        if k in seen:
            continue
        seen.add(k)
        out.append(t)
    return out[:6]


def _template_caption(ev: dict, tone: str = "friendly", base: Optional[str] = None) -> str:
    """Deterministic, no-LLM caption generator using event fields."""
    title = (ev.get("title") or "Our next event").strip()
    when = _humanize_event_when(ev)
    venue = (ev.get("venue") or "").strip()
    address = (ev.get("address") or "").strip()
    where = venue or address
    cost = (ev.get("cost") or "").strip()
    desc = (ev.get("description") or "").strip()
    # Short blurb: first sentence up to 180 chars
    blurb = re.split(r"(?<=[.!?])\s+", desc, 1)[0].strip() if desc else ""
    if len(blurb) > 180:
        blurb = blurb[:177].rstrip() + "…"
    link = _event_public_link(ev, base)
    hashtags = " ".join(_event_hashtags(ev))

    if tone == "punchy":
        lines = [f"🎉 {title}"]
        if when:
            lines.append(f"📅 {when}")
        if where:
            lines.append(f"📍 {where}")
        if cost:
            lines.append(f"💷 {cost}")
        if blurb:
            lines.append("")
            lines.append(blurb)
        lines.append("")
        lines.append(f"👉 Full details: {link}")
        lines.append("")
        lines.append(hashtags)
        return "\n".join(lines)

    if tone == "formal":
        parts = [f"You are warmly invited to {title}."]
        if when and where:
            parts.append(f"Taking place on {when} at {where}.")
        elif when:
            parts.append(f"Taking place on {when}.")
        elif where:
            parts.append(f"Held at {where}.")
        if cost:
            parts.append(f"Cost: {cost}.")
        if blurb:
            parts.append(blurb)
        parts.append(f"Full details and booking: {link}")
        parts.append("")
        parts.append(hashtags)
        return "\n\n".join(parts).strip()

    # friendly (default)
    header = f"✨ {title}"
    meta_bits = []
    if when:
        meta_bits.append(f"🗓 {when}")
    if where:
        meta_bits.append(f"📍 {where}")
    if cost:
        meta_bits.append(f"💷 {cost}")
    meta = "  ".join(meta_bits)
    body = blurb or "Come along and be part of it — everyone welcome."
    footer = f"🔗 {link}\n\n{hashtags}"
    return "\n\n".join(x for x in [header, meta, body, footer] if x)


async def _ai_caption(ev: dict, org: Optional[dict], tone: str, base: Optional[str] = None) -> str:
    """Optional AI polish using Emergent LLM Key (Claude). Falls back to
    template if key missing or the call fails."""
    if not EMERGENT_LLM_KEY:
        return _template_caption(ev, tone, base)
    if not await _ensure_llm_loaded():
        return _template_caption(ev, tone, base)
    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        tone_hint = {
            "friendly": "warm, welcoming, community-first — like a friend inviting neighbours along",
            "punchy": "short, energetic, high-impact with emojis and line breaks",
            "formal": "polite, respectful, no emojis, suitable for a church or civic group",
        }.get(tone, "warm and community-first")
        system = (
            "You write short social media captions for community events in Blackrod, Bolton. "
            f"Tone: {tone_hint}. "
            "Rules: 400 characters MAX for body, include the event link at the end, "
            "add 3-5 UK-focused hashtags on their own line, no hard-sell language, "
            "use British English spelling. Output the caption text ONLY — no preamble, no quotes."
        )
        facts = {
            "title": ev.get("title"),
            "when": _humanize_event_when(ev),
            "venue": ev.get("venue"),
            "address": ev.get("address"),
            "cost": ev.get("cost"),
            "age": ev.get("age"),
            "accessibility": ev.get("accessibility"),
            "description": (ev.get("description") or "")[:600],
            "organiser": (org or {}).get("name"),
            "category": ev.get("category"),
            "link": _event_public_link(ev, base),
            "hashtags_suggested": _event_hashtags(ev),
        }
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"caption-{ev.get('id')}-{tone}",
            system_message=system,
        ).with_model("anthropic", "claude-sonnet-4-5-20250929")
        prompt = "Write the caption using these event facts (JSON):\n\n" + json.dumps(facts, ensure_ascii=False)
        raw = await chat.send_message(UserMessage(text=prompt))
        text = (raw or "").strip()
        if text.startswith("```"):
            text = re.sub(r"^```(?:\w+)?", "", text).strip()
            text = re.sub(r"```$", "", text).strip()
        # Sanity check — ensure link is present, otherwise re-append
        link = _event_public_link(ev, base)
        if link not in text:
            text = text.rstrip() + f"\n\n🔗 {link}"
        return text
    except Exception as exc:
        logger.warning(f"AI caption fallback (template) — {exc}")
        return _template_caption(ev, tone, base)


@api.get("/events/{event_id}/social-bundle")
async def event_social_bundle(
    event_id: str,
    request: Request,
    tone: str = "friendly",
    ai: bool = False,
):
    """One-shot 'Post Now' bundle: caption + hashtags + poster/link URLs +
    ready-to-copy variants. Called by the Post Now dialog on the org
    dashboard and event detail pages.
    """
    tone = tone if tone in _TONE_CHOICES else "friendly"
    base = _abs_base_url(request)
    ev = await _fetch_event_for_poster(event_id)
    org = None
    if ev.get("orgSlug"):
        org = await db.orgs.find_one({"slug": ev["orgSlug"]}, {"_id": 0})
    if ai:
        caption = await _ai_caption(ev, org, tone, base)
    else:
        caption = _template_caption(ev, tone, base)
    hashtags = _event_hashtags(ev)
    link = _event_public_link(ev, base)
    return {
        "event_id": ev.get("id"),
        "title": ev.get("title"),
        "tone": tone,
        "ai": bool(ai and EMERGENT_LLM_KEY),
        "caption": caption,
        "caption_with_link": caption if link in caption else f"{caption}\n\n{link}",
        "hashtags": hashtags,
        "link": link,
        "og_url": f"{base}/api/events/{ev.get('id')}/og",
        "poster_png": f"{base}/api/events/{ev.get('id')}/poster.png",
        "poster_pdf": f"{base}/api/events/{ev.get('id')}/poster.pdf",
    }


# ─────────── Org dashboard analytics uplift ───────────
@api.get("/orgs/{slug}/analytics")
async def org_analytics_v2(slug: str, days: int = 30):
    """Per-day series for org's own reach + best-performing event over the
    window. Public (any visitor can see basic reach) — the org dashboard uses
    these numbers to celebrate what's working."""
    org = await _find_org(slug)
    days = max(1, min(int(days), 180))
    now = datetime.now(timezone.utc)
    since = (now - timedelta(days=days)).isoformat()

    # Bucket analytics events by day + kind for this org's content.
    pipeline = [
        {"$match": {"created_at": {"$gte": since}, "org_slug": slug}},
        {"$group": {
            "_id": {"day": {"$substr": ["$created_at", 0, 10]}, "kind": "$kind"},
            "n": {"$sum": "$count"},
        }},
        {"$sort": {"_id.day": 1}},
    ]
    rows = await db.analytics_events.aggregate(pipeline).to_list(4000)
    days_map: Dict[str, Dict[str, int]] = {}
    for r in rows:
        day = r["_id"]["day"]
        kind = r["_id"]["kind"] or "misc"
        days_map.setdefault(day, {})[kind] = r["n"]

    series: List[Dict[str, Any]] = []
    for i in range(days):
        d = (now - timedelta(days=days - 1 - i)).date().isoformat()
        entry = days_map.get(d, {})
        series.append({
            "day": d,
            "event_views": entry.get("event_view", 0),
            "org_views": entry.get("org_view", 0),
            "share_clicks": entry.get("share_click", 0),
            "volunteer_contacts": entry.get("volunteer_contact", 0),
        })

    # Totals + best-performing event
    total_event_views = sum(s["event_views"] for s in series)
    total_org_views = sum(s["org_views"] for s in series)
    total_shares = sum(s["share_clicks"] for s in series)

    best_pipeline = [
        {"$match": {
            "created_at": {"$gte": since},
            "org_slug": slug,
            "kind": "event_view",
            "entity_id": {"$ne": None},
        }},
        {"$group": {"_id": "$entity_id", "views": {"$sum": "$count"}}},
        {"$sort": {"views": -1}},
        {"$limit": 1},
    ]
    best_rows = await db.analytics_events.aggregate(best_pipeline).to_list(1)
    best: Optional[Dict[str, Any]] = None
    if best_rows:
        ev = await db.events.find_one({"id": best_rows[0]["_id"]}, {"_id": 0, "id": 1, "title": 1, "start": 1})
        if ev:
            best = {"id": ev["id"], "title": ev.get("title", ""), "start": ev.get("start", ""), "views": best_rows[0]["views"]}

    return {
        "slug": slug,
        "org_name": org.get("name", ""),
        "window_days": days,
        "series": series,
        "totals": {
            "event_views": total_event_views,
            "org_views": total_org_views,
            "share_clicks": total_shares,
        },
        "best_event": best,
    }






# ─────────── Chatbot context (public read-only feed for Charla) ───────────
from data.faqs import FAQS as _CHAT_FAQS  # noqa: E402


def _iso_date_only(value: Any) -> str:
    if not value:
        return ""
    try:
        if isinstance(value, str):
            return value[:10]
        return value.isoformat()[:10]
    except Exception:
        return ""


async def _build_chat_context(days: int = 30) -> Dict[str, Any]:
    """Compact, public snapshot for a chatbot knowledge base.
    - upcoming events (next `days` days, approved)
    - approved organisations (name, category, short, contact)
    - venues
    - volunteering opportunities
    - FAQs
    """
    now = datetime.now(timezone.utc)
    horizon = now + timedelta(days=max(1, days))

    orgs_cursor = db.orgs.find({"status": "approved"}, {"_id": 0}).sort("name", 1)
    orgs = await orgs_cursor.to_list(length=500)
    org_name_by_slug = {o.get("slug"): o.get("name") for o in orgs}

    events_cursor = db.events.find(
        {"status": "approved", "start": {"$lte": horizon.isoformat()}},
        {"_id": 0},
    ).sort("start", 1)
    events_all = await events_cursor.to_list(length=500)

    def _keep(ev: Dict[str, Any]) -> bool:
        start = ev.get("start") or ""
        if not start:
            return False
        try:
            when = datetime.fromisoformat(start.replace("Z", "+00:00"))
            if when.tzinfo is None:
                when = when.replace(tzinfo=timezone.utc)
        except Exception:
            return False
        return when >= now - timedelta(hours=6)  # allow "happening today"

    events = [ev for ev in events_all if _keep(ev)][:80]

    venues = await db.venues.find({}, {"_id": 0}).sort("name", 1).to_list(length=200)
    volunteers = await db.volunteers.find({}, {"_id": 0}).sort("title", 1).to_list(length=200)

    # Prefer explicit PUBLIC_URL, fall back to APP_URL supervisor env, else placeholder.
    site_url = (os.environ.get("PUBLIC_URL") or os.environ.get("APP_URL") or PUBLIC_URL).rstrip("/")

    def _e(ev: Dict[str, Any]) -> Dict[str, Any]:
        return {
            "id": ev.get("id"),
            "title": ev.get("title"),
            "date": _iso_date_only(ev.get("start")),
            "time": (ev.get("start") or "")[11:16],
            "end": (ev.get("end") or "")[11:16],
            "venue": ev.get("venue") or "",
            "address": ev.get("address") or "",
            "category": ev.get("category") or "",
            "cost": ev.get("cost") or "",
            "age": ev.get("age") or "",
            "organiser": org_name_by_slug.get(ev.get("orgSlug"), ev.get("orgSlug") or ""),
            "description": (ev.get("description") or "")[:400],
            "url": f"{site_url}/events/{ev.get('id')}" if ev.get("id") else "",
        }

    def _o(org: Dict[str, Any]) -> Dict[str, Any]:
        socials = org.get("socials") or {}
        return {
            "slug": org.get("slug"),
            "name": org.get("name"),
            "category": org.get("category") or "",
            "short": (org.get("short") or org.get("about") or "")[:280],
            "email": org.get("email") or "",
            "phone": org.get("phone") or "",
            "website": org.get("website") or "",
            "facebook": socials.get("facebook") or "",
            "instagram": socials.get("instagram") or "",
            "meeting": org.get("meeting") or "",
            "address": org.get("address") or "",
            "url": f"{site_url}/organisations/{org.get('slug')}" if org.get("slug") else "",
        }

    def _v(v: Dict[str, Any]) -> Dict[str, Any]:
        return {
            "id": v.get("id"),
            "name": v.get("name"),
            "address": v.get("address") or "",
            "capacity": v.get("capacity") or 0,
            "facilities": v.get("facilities") or [],
            "accessibility": v.get("accessibility") or "",
            "booking": v.get("booking") or "",
        }

    def _vol(v: Dict[str, Any]) -> Dict[str, Any]:
        return {
            "id": v.get("id"),
            "title": v.get("title"),
            "organiser": org_name_by_slug.get(v.get("orgSlug"), v.get("orgSlug") or ""),
            "description": (v.get("description") or "")[:400],
            "time": v.get("time") or "",
            "age": v.get("age") or "",
            "skills": v.get("skills") or "",
        }

    return {
        "site": {
            "name": "Blackrod Now",
            "url": site_url,
            "description": "Community hub for Blackrod, Bolton — events, groups, clubs, schools, businesses and volunteering.",
            "contact": ADMIN_SENDER_EMAILS[0] if 'ADMIN_SENDER_EMAILS' in globals() and ADMIN_SENDER_EMAILS else "hello@communityalliances.co.uk",
            "generated_at": now.isoformat(),
            "window_days": days,
        },
        "events": [_e(ev) for ev in events],
        "organisations": [_o(o) for o in orgs],
        "venues": [_v(v) for v in venues],
        "volunteering": [_vol(v) for v in volunteers],
        "faqs": _CHAT_FAQS,
    }


@api.get("/chat/context")
async def chat_context(days: int = 30):
    """Public JSON snapshot of Blackrod Now data — designed to be fed into an
    external chatbot's knowledge base (Charla, etc.). Refreshes on every hit."""
    try:
        days_clamped = max(1, min(int(days), 180))
    except Exception:
        days_clamped = 30
    return await _build_chat_context(days_clamped)


@api.get("/chat/context.md", response_class=_PlainResp)
async def chat_context_markdown(days: int = 30):
    """Human-readable Markdown mirror of the chat context. Many chatbot builders
    prefer to ingest a URL of Markdown rather than JSON."""
    try:
        days_clamped = max(1, min(int(days), 180))
    except Exception:
        days_clamped = 30
    ctx = await _build_chat_context(days_clamped)

    lines: List[str] = []
    site = ctx["site"]
    lines.append(f"# {site['name']}")
    lines.append("")
    lines.append(site["description"])
    lines.append("")
    lines.append(f"- Website: {site['url']}")
    lines.append(f"- Contact: {site['contact']}")
    lines.append(f"- Snapshot generated: {site['generated_at']}")
    lines.append("")

    lines.append(f"## Upcoming events (next {site['window_days']} days)")
    lines.append("")
    if not ctx["events"]:
        lines.append("_No upcoming events in the current window._")
    for ev in ctx["events"]:
        header = f"### {ev['title']}"
        lines.append(header)
        meta = [
            f"**When:** {ev['date']}" + (f" · {ev['time']}" if ev["time"] else ""),
            f"**Where:** {ev['venue']}" + (f", {ev['address']}" if ev["address"] else "") if ev["venue"] else None,
            f"**Organiser:** {ev['organiser']}" if ev["organiser"] else None,
            f"**Category:** {ev['category']}" if ev["category"] else None,
            f"**Cost:** {ev['cost']}" if ev["cost"] else None,
            f"**Age:** {ev['age']}" if ev["age"] else None,
        ]
        for m in meta:
            if m:
                lines.append(m)
        if ev["description"]:
            lines.append("")
            lines.append(ev["description"])
        if ev["url"]:
            lines.append("")
            lines.append(f"[Event page]({ev['url']})")
        lines.append("")

    lines.append("## Organisations")
    lines.append("")
    for org in ctx["organisations"]:
        lines.append(f"### {org['name']}")
        if org["category"]:
            lines.append(f"_{org['category']}_")
        if org["short"]:
            lines.append("")
            lines.append(org["short"])
        contacts = []
        if org["email"]:
            contacts.append(f"Email: {org['email']}")
        if org["phone"]:
            contacts.append(f"Phone: {org['phone']}")
        if org["website"]:
            contacts.append(f"Website: {org['website']}")
        if org["meeting"]:
            contacts.append(f"When they meet: {org['meeting']}")
        if org["address"]:
            contacts.append(f"Address: {org['address']}")
        if contacts:
            lines.append("")
            for c in contacts:
                lines.append(f"- {c}")
        if org["url"]:
            lines.append("")
            lines.append(f"[Organisation page]({org['url']})")
        lines.append("")

    lines.append("## Venues (spaces to hire)")
    lines.append("")
    for v in ctx["venues"]:
        lines.append(f"### {v['name']}")
        if v["address"]:
            lines.append(f"Address: {v['address']}")
        if v["capacity"]:
            lines.append(f"Capacity: {v['capacity']}")
        if v["facilities"]:
            lines.append(f"Facilities: {', '.join(v['facilities'])}")
        if v["accessibility"]:
            lines.append(f"Accessibility: {v['accessibility']}")
        if v["booking"]:
            lines.append(f"Booking: {v['booking']}")
        lines.append("")

    lines.append("## Volunteering opportunities")
    lines.append("")
    for v in ctx["volunteering"]:
        lines.append(f"### {v['title']}")
        if v["organiser"]:
            lines.append(f"_{v['organiser']}_")
        if v["description"]:
            lines.append("")
            lines.append(v["description"])
        if v["time"]:
            lines.append(f"- Time: {v['time']}")
        if v["age"]:
            lines.append(f"- Age: {v['age']}")
        if v["skills"]:
            lines.append(f"- Skills: {v['skills']}")
        lines.append("")

    lines.append("## Frequently asked questions")
    lines.append("")
    for f in ctx["faqs"]:
        lines.append(f"### {f['q']}")
        lines.append(f"_{f['cat']}_")
        lines.append("")
        lines.append(f["a"])
        lines.append("")

    return "\n".join(lines)




# ─────────── Subscribers ───────────
class SubscribeReq(BaseModel):
    email: EmailStr
    device_id: Optional[str] = None
    followed_orgs: Optional[List[str]] = None
    followed_categories: Optional[List[str]] = None
    saved_events: Optional[List[str]] = None


@api.post("/subscribe")
async def subscribe(req: SubscribeReq):
    existing = await db.subscribers.find_one({"email": req.email.lower()}, {"_id": 0})
    if existing:
        # merge follows + always reactivate (never insert a duplicate)
        followed_orgs = list(set((existing.get("followed_orgs") or []) + (req.followed_orgs or [])))
        followed_categories = list(set((existing.get("followed_categories") or []) + (req.followed_categories or [])))
        saved_events = list(set((existing.get("saved_events") or []) + (req.saved_events or [])))
        await db.subscribers.update_one(
            {"email": req.email.lower()},
            {"$set": {
                "followed_orgs": followed_orgs,
                "followed_categories": followed_categories,
                "saved_events": saved_events,
                "unsubscribed": False,
                "digest": True,
            }},
        )
        return {
            "ok": True,
            "already_subscribed": True,
            "reactivated": bool(existing.get("unsubscribed")),
            "unsub_token": existing["unsub_token"],
            "pref_token": existing["pref_token"],
        }
    sub = Subscriber(
        email=req.email.lower(),
        device_id=req.device_id,
        followed_orgs=req.followed_orgs or [],
        followed_categories=req.followed_categories or [],
        saved_events=req.saved_events or [],
    )
    await db.subscribers.insert_one(sub.model_dump())
    # welcome email
    unsub_link = f"{PUBLIC_URL}/unsubscribe/{sub.unsub_token}"
    pref_link = f"{PUBLIC_URL}/preferences/{sub.pref_token}"
    html = _render_welcome(sub.email, unsub_link, pref_link)
    asyncio.create_task(asyncio.to_thread(resend_send, sub.email, "Welcome to Blackrod Now 👋", html))
    return {"ok": True, "already_subscribed": False, "unsub_token": sub.unsub_token, "pref_token": sub.pref_token}


class SavedEventsSyncReq(BaseModel):
    email: Optional[EmailStr] = None
    device_id: Optional[str] = None
    saved_events: List[str] = []


@api.post("/subscribers/saved-events")
async def sync_saved_events(req: SavedEventsSyncReq):
    """Client pushes their locally-saved event ids to the backend so we can
    send reminder emails 24h + 2h before each event starts. Matches by email
    first, then device_id. No-ops for unsubscribed users."""
    query = None
    if req.email:
        query = {"email": req.email.lower(), "unsubscribed": {"$ne": True}}
    elif req.device_id:
        query = {"device_id": req.device_id, "unsubscribed": {"$ne": True}}
    if not query:
        return {"ok": False, "reason": "no email or device_id"}
    result = await db.subscribers.update_one(
        query, {"$set": {"saved_events": list(dict.fromkeys(req.saved_events or []))}}
    )
    return {"ok": True, "matched": result.matched_count, "count": len(req.saved_events or [])}


@api.post("/unsubscribe/{token}")
async def unsubscribe(token: str):
    res = await db.subscribers.update_one({"unsub_token": token}, {"$set": {"unsubscribed": True}})
    if not res.matched_count:
        raise HTTPException(404, "Unknown unsubscribe token")
    return {"ok": True}


@api.get("/preferences/{token}")
async def get_preferences(token: str):
    sub = await db.subscribers.find_one({"pref_token": token}, {"_id": 0})
    if not sub:
        raise HTTPException(404, "Unknown preferences token")
    return sub


class PrefPatch(BaseModel):
    followed_orgs: Optional[List[str]] = None
    followed_categories: Optional[List[str]] = None
    digest: Optional[bool] = None


@api.patch("/preferences/{token}")
async def patch_preferences(token: str, patch: PrefPatch):
    updates = {k: v for k, v in patch.model_dump(exclude_none=True).items()}
    if not updates:
        return await get_preferences(token)
    await db.subscribers.update_one({"pref_token": token}, {"$set": updates})
    return await get_preferences(token)


@api.get("/admin/subscribers")
async def admin_list_subscribers(
    q: str = "",
    include_unsubscribed: bool = False,
    digest_only: bool = False,
    limit: int = 500,
):
    safe_limit = max(1, min(limit, 2000))
    query: Dict[str, Any] = {}
    if not include_unsubscribed:
        query["unsubscribed"] = False
    if digest_only:
        query["digest"] = True
    if q.strip():
        query["email"] = {"$regex": re.escape(q.strip()), "$options": "i"}

    rows = await db.subscribers.find(query, {"_id": 0}).sort("created_at", -1).to_list(safe_limit)
    items = []
    for sub in rows:
        items.append(
            {
                **sub,
                "followed_orgs_count": len(sub.get("followed_orgs") or []),
                "followed_categories_count": len(sub.get("followed_categories") or []),
                "saved_events_count": len(sub.get("saved_events") or []),
            }
        )

    return {
        "items": items,
        "count": len(items),
        "total_active": await db.subscribers.count_documents({"unsubscribed": False}),
        "total_digest": await db.subscribers.count_documents({"unsubscribed": False, "digest": True}),
    }


# ─────────── Device follows (anonymous, no email) ───────────
class FollowReq(BaseModel):
    device_id: str
    kind: Literal["org", "category"]
    value: str
    action: Literal["add", "remove"]


@api.get("/follows/{device_id}")
async def get_follows(device_id: str):
    doc = await db.follows.find_one({"device_id": device_id}, {"_id": 0})
    return doc or {"device_id": device_id, "orgs": [], "categories": []}


@api.post("/follows")
async def toggle_follow(req: FollowReq):
    key = "orgs" if req.kind == "org" else "categories"
    op = "$addToSet" if req.action == "add" else "$pull"
    await db.follows.update_one({"device_id": req.device_id}, {op: {key: req.value}}, upsert=True)
    return await get_follows(req.device_id)


# ─────────── Notifications (admin → org) ───────────
class AdminNotifyReq(BaseModel):
    org_slug: str
    title: str
    body: str
    send_email: bool = True


@api.post("/admin/notifications")
async def admin_send_notification(req: AdminNotifyReq):
    n = Notification(org_slug=req.org_slug, title=req.title, body=req.body)
    await db.notifications.insert_one(n.model_dump())
    email_result: Optional[Dict[str, Any]] = None
    if req.send_email:
        org = await db.orgs.find_one({"slug": req.org_slug}, {"_id": 0, "email": 1, "name": 1})
        org_email = (org or {}).get("email", "")
        if org_email and EMAIL_RE.match(org_email):
            html_payload = _render_admin_email_html(req.title, req.body, SENDER_EMAIL)
            email_result = await asyncio.to_thread(
                resend_send,
                org_email,
                req.title,
                html_payload,
                SENDER_EMAIL,
                SENDER_NAME,
            )
    return {**n.model_dump(), "email_delivery": email_result}


@api.get("/organisations/{slug}/notifications")
async def get_org_notifications(slug: str):
    return await db.notifications.find({"org_slug": slug}, {"_id": 0}).sort("created_at", -1).to_list(200)


@api.patch("/notifications/{nid}/read")
async def mark_notification_read(nid: str):
    await db.notifications.update_one({"id": nid}, {"$set": {"read": True}})
    return {"ok": True}


# ─────────── Contact admin (org → super admin) ───────────
class ContactAdminReq(BaseModel):
    from_org_slug: Optional[str] = None
    from_email: Optional[str] = None
    from_name: Optional[str] = None
    subject: str
    body: str
    in_reply_to: Optional[str] = None


class WebWizardEnquiryReq(BaseModel):
    from_name: str
    from_email: EmailStr
    business: str = ""
    service: str
    budget: str = ""
    timeline: str = ""
    details: str


def _render_web_wizard_enquiry_html(req: WebWizardEnquiryReq) -> str:
    lines = [
        "<p><b>New Web Design Wizard enquiry</b></p>",
        f"<p><b>Name:</b> {html.escape(req.from_name)}</p>",
        f"<p><b>Email:</b> {html.escape(req.from_email)}</p>",
        f"<p><b>Business:</b> {html.escape(req.business or 'Not provided')}</p>",
        f"<p><b>Service needed:</b> {html.escape(req.service)}</p>",
        f"<p><b>Budget range:</b> {html.escape(req.budget or 'Not provided')}</p>",
        f"<p><b>Preferred timeline:</b> {html.escape(req.timeline or 'Not provided')}</p>",
        "<p><b>Project details:</b></p>",
        f"<p style='white-space:pre-wrap'>{html.escape(req.details)}</p>",
    ]
    return "".join(lines)


def _render_contact_admin_html(req: ContactAdminReq) -> str:
    sender = req.from_name or req.from_org_slug or req.from_email or "Unknown sender"
    sender_email = req.from_email or "Not provided"
    org_slug = req.from_org_slug or "Not provided"
    body_html = _auto_link(req.body or "")
    reply_hint = (
        f"<p><b>Related notification:</b> {html.escape(req.in_reply_to)}</p>"
        if req.in_reply_to
        else ""
    )
    return (
        "<p><b>New organisation message for site admin</b></p>"
        f"<p><b>From:</b> {html.escape(sender)}</p>"
        f"<p><b>Sender email:</b> {html.escape(sender_email)}</p>"
        f"<p><b>Organisation slug:</b> {html.escape(org_slug)}</p>"
        f"{reply_hint}"
        "<p><b>Message:</b></p>"
        f"{body_html}"
    )


@api.post("/contact-admin")
async def contact_admin(req: ContactAdminReq):
    m = AdminMessage(**req.model_dump(), direction="inbound_org")
    await db.messages.insert_one(m.model_dump())

    email_result: Optional[Dict[str, Any]] = None
    admin_inbox = [addr for addr in ADMIN_INBOX_EMAILS if EMAIL_RE.match(addr)]
    if admin_inbox:
        subject = req.subject.strip() or "Message from organisation"
        primary = admin_inbox[0]
        bcc = admin_inbox[1:] if len(admin_inbox) > 1 else None
        email_result = await asyncio.to_thread(
            resend_send,
            primary,
            f"[Org message] {subject}",
            _render_contact_admin_html(req),
            req.from_email if req.from_email and EMAIL_RE.match(req.from_email) else None,
            req.from_name or req.from_org_slug or "Organisation",
            bcc,
        )
        await db.messages.update_one({"id": m.id}, {"$set": {"delivery": email_result}})

    return {**m.model_dump(), "delivery": email_result}


@api.post("/web-wizard/enquiry")
async def web_wizard_enquiry(req: WebWizardEnquiryReq):
    result = resend_send(
        to_email=WEB_WIZARD_TO_EMAIL,
        subject=f"Web Design Wizard enquiry: {req.service}",
        html=_render_web_wizard_enquiry_html(req),
        bcc=[WEB_WIZARD_BCC_EMAIL],
    )
    if not result.get("ok"):
        raise HTTPException(502, "Could not send enquiry email")
    return {"ok": True, "mocked": bool(result.get("mocked", False))}


@api.get("/notifications/{nid}/thread")
async def notification_thread(nid: str):
    """Return the admin notification + any org replies linked to it (oldest first)."""
    notif = await db.notifications.find_one({"id": nid}, {"_id": 0})
    if not notif:
        raise HTTPException(404, "Notification not found")
    replies = await db.messages.find({"in_reply_to": nid}, {"_id": 0}).sort("created_at", 1).to_list(200)
    return {"notification": notif, "replies": replies}


@api.get("/admin/messages")
async def admin_get_messages():
    return await db.messages.find({}, {"_id": 0}).sort("created_at", -1).to_list(500)


@api.patch("/admin/messages/{mid}/read")
async def admin_mark_message_read(mid: str):
    await db.messages.update_one({"id": mid}, {"$set": {"read": True}})
    return {"ok": True}


@api.delete("/admin/messages/{mid}")
async def admin_delete_message(mid: str):
    res = await db.messages.delete_one({"id": mid})
    if not res.deleted_count:
        raise HTTPException(404, "Message not found")
    return {"ok": True}


class AdminMessageBulkReq(BaseModel):
    ids: List[str]
    action: Literal["delete", "archive"]


@api.post("/admin/messages/bulk")
async def admin_bulk_message_action(req: AdminMessageBulkReq):
    if not req.ids:
        raise HTTPException(400, "No message IDs supplied")
    if req.action == "delete":
        await db.messages.delete_many({"id": {"$in": req.ids}})
    else:
        await db.messages.update_many({"id": {"$in": req.ids}}, {"$set": {"read": True}})
    return {"ok": True, "affected": len(req.ids)}


@api.post("/webhooks/resend")
async def resend_webhook_email_received(request: Request):
    raw = await request.body()
    svix_id = request.headers.get("svix-id", "")
    svix_timestamp = request.headers.get("svix-timestamp", "")
    svix_signature = request.headers.get("svix-signature", "")
    verified, reason = _verify_resend_webhook_signature(raw, svix_id, svix_timestamp, svix_signature)
    if not verified:
        raise HTTPException(401, f"Invalid webhook signature: {reason}")

    try:
        payload = json.loads(raw.decode("utf-8") or "{}")
    except Exception:
        raise HTTPException(400, "Invalid JSON payload")

    event_type = str(payload.get("type") or payload.get("event") or "").strip()
    if event_type != "email.received":
        return {"ok": True, "ignored": True, "event_type": event_type or "unknown"}

    data = payload.get("data") if isinstance(payload.get("data"), dict) else payload
    from_field = data.get("from") or data.get("sender") or data.get("from_email") or ""
    from_email = _extract_email_address(from_field)
    from_name = _extract_sender_name(from_field)
    subject = str(data.get("subject") or "").strip() or "(No subject)"

    headers = data.get("headers")
    in_reply_to = str(
        data.get("in_reply_to")
        or data.get("inReplyTo")
        or _extract_header_value(headers, "in-reply-to")
        or ""
    ).strip() or None
    event_id = str(
        data.get("id")
        or data.get("message_id")
        or _extract_header_value(headers, "message-id")
        or svix_id
    ).strip()

    duplicate = await db.messages.find_one(
        {
            "delivery.provider": "resend-webhook",
            "delivery.event_id": event_id,
        },
        {"_id": 0, "id": 1},
    )
    if duplicate:
        return {"ok": True, "duplicate": True, "message_id": duplicate.get("id")}

    body_text = str(
        data.get("text")
        or data.get("text_body")
        or data.get("plain_text")
        or ""
    ).strip()
    if not body_text:
        body_text = _html_to_text(data.get("html") or data.get("html_body") or "")
    if not body_text:
        body_text = "(No message body provided by webhook)"

    org_slug = None
    if from_email:
        org = await db.orgs.find_one({"email": from_email}, {"_id": 0, "slug": 1, "name": 1})
        if org:
            org_slug = org.get("slug")
            if not from_name:
                from_name = org.get("name") or ""

    created = AdminMessage(
        from_org_slug=org_slug,
        from_email=from_email or None,
        from_name=from_name or None,
        subject=subject,
        body=body_text,
        in_reply_to=in_reply_to,
        direction="inbound_org",
        delivery={
            "provider": "resend-webhook",
            "event_type": event_type,
            "event_id": event_id,
            "svix_id": svix_id,
        },
    )
    await db.messages.insert_one(created.model_dump())
    return {
        "ok": True,
        "created": True,
        "message_id": created.id,
        "from_email": from_email,
        "matched_org_slug": org_slug,
    }


class AdminMessageReplyReq(BaseModel):
    body: str
    subject: Optional[str] = None
    from_email: Optional[str] = None
    from_name: Optional[str] = None


@api.post("/admin/messages/{mid}/reply")
async def admin_reply_to_message(mid: str, req: AdminMessageReplyReq):
    original = await db.messages.find_one({"id": mid}, {"_id": 0})
    if not original:
        raise HTTPException(404, "Message not found")

    recipient = (original.get("from_email") or "").strip().lower()
    if not recipient and original.get("from_org_slug"):
        org = await db.orgs.find_one({"slug": original["from_org_slug"]}, {"_id": 0, "email": 1})
        recipient = ((org or {}).get("email") or "").strip().lower()
    if not recipient or not EMAIL_RE.match(recipient):
        raise HTTPException(400, "No valid recipient email on this message")

    from_addr = (req.from_email or SENDER_EMAIL).strip().lower()
    if from_addr not in [s.lower() for s in ADMIN_SENDER_EMAILS]:
        raise HTTPException(400, f"Sender must be one of: {', '.join(ADMIN_SENDER_EMAILS)}")
    if not (req.body or "").strip():
        raise HTTPException(400, "Reply body is required")

    subject = (req.subject or f"Re: {original.get('subject', '').strip()}").strip()
    html_payload = _render_admin_email_html(subject, req.body, from_addr)
    result = await asyncio.to_thread(
        resend_send,
        recipient,
        subject,
        html_payload,
        from_addr,
        req.from_name or SENDER_NAME,
    )

    created = AdminMessage(
        from_name=req.from_name or "Admin",
        from_email=from_addr,
        to_org_slug=original.get("from_org_slug"),
        to_email=recipient,
        subject=subject,
        body=req.body,
        in_reply_to=original.get("in_reply_to") or original.get("id"),
        parent_message_id=original.get("id"),
        direction="outbound_admin",
        read=True,
        delivery=result,
    )
    await db.messages.insert_one(created.model_dump())
    await db.messages.update_one({"id": original.get("id")}, {"$set": {"read": True}})
    return {"ok": bool(result.get("ok")), "message": created, "delivery": result}


# ─────────── Documents (per org) ───────────
ALLOWED_EXT = {"pdf", "doc", "docx", "xls", "xlsx", "ppt", "pptx", "jpg", "jpeg", "png", "webp", "bmp", "gif", "tif", "tiff", "heic", "heif", "txt", "csv", "md"}
MAX_UPLOAD_BYTES = 10 * 1024 * 1024  # 10 MB
MAX_BULK_PARSE_FILES = int(os.environ.get("MAX_BULK_PARSE_FILES", "8"))
BULK_PARSE_TOTAL_TIMEOUT_SECONDS = int(os.environ.get("BULK_PARSE_TOTAL_TIMEOUT_SECONDS", "80"))
BULK_PARSE_EXTRACTION_TIMEOUT_SECONDS = int(os.environ.get("BULK_PARSE_EXTRACTION_TIMEOUT_SECONDS", "20"))
BULK_PARSE_CLASSIFICATION_TIMEOUT_SECONDS = int(os.environ.get("BULK_PARSE_CLASSIFICATION_TIMEOUT_SECONDS", "30"))
BULK_PARSE_TIMEOUT_SAFETY_SECONDS = 5


@api.post("/organisations/{slug}/documents")
async def upload_document(
    slug: str,
    file: UploadFile = File(...),
    org_auth: Optional[str] = Header(None, alias="X-Org-Auth"),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_org_write_access(slug, org_auth, admin_code)
    await _find_org(slug)
    filename = file.filename or "file"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "bin"
    if ext not in ALLOWED_EXT:
        raise HTTPException(400, f"Unsupported file type .{ext}")
    data = await file.read()
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(413, "File too large (max 10 MB)")
    path = f"{APP_NAME}/orgs/{slug}/{uuid.uuid4()}.{ext}"
    try:
        result = put_object(path, data, file.content_type or "application/octet-stream")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"Upload failed: {e}")
    d = Document(
        org_slug=slug,
        name=filename,
        storage_path=result["path"],
        content_type=file.content_type or "application/octet-stream",
        size=result.get("size", len(data)),
    )
    await db.documents.insert_one(d.model_dump())
    return d


@api.get("/organisations/{slug}/documents")
async def list_documents(slug: str):
    return await db.documents.find({"org_slug": slug, "deleted": False}, {"_id": 0}).sort("created_at", -1).to_list(200)


@api.delete("/organisations/{slug}/documents/{doc_id}")
async def delete_document(
    slug: str,
    doc_id: str,
    org_auth: Optional[str] = Header(None, alias="X-Org-Auth"),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_org_write_access(slug, org_auth, admin_code)
    await db.documents.update_one({"id": doc_id, "org_slug": slug}, {"$set": {"deleted": True}})
    return {"ok": True}


@api.get("/documents/{doc_id}/download")
async def download_document(doc_id: str):
    d = await db.documents.find_one({"id": doc_id, "deleted": False}, {"_id": 0})
    if not d:
        raise HTTPException(404, "Document not found")
    try:
        data, ctype = get_object(d["storage_path"])
    except Exception as e:
        raise HTTPException(500, f"Download failed: {e}")
    return Response(
        content=data,
        media_type=d.get("content_type", ctype),
        headers={"Content-Disposition": f'attachment; filename="{d["name"]}"'},
    )


# ─────────── Organisation Logo & Cover images ───────────
IMAGE_EXT = {"png", "jpg", "jpeg", "webp"}
LOGO_AVATAR_SIZE = 512
LOGO_THUMB_SIZE = 128
COVER_WIDTH = 1600
COVER_HEIGHT = 500
MAX_IMAGE_BYTES = 5 * 1024 * 1024  # 5 MB


def _open_image(data: bytes) -> Image.Image:
    try:
        img = Image.open(BytesIO(data))
        img = ImageOps.exif_transpose(img)
    except Exception as e:
        raise HTTPException(400, f"Invalid image: {e}")
    if img.mode not in ("RGB", "RGBA"):
        img = img.convert("RGBA")
    return img


def _process_and_upload_logo(slug: str, data: bytes) -> tuple[str, str]:
    """Center-crop to 512x512 + 128x128 thumb. Returns (avatar_path, thumb_path)."""
    img = _open_image(data)
    avatar = ImageOps.fit(img, (LOGO_AVATAR_SIZE, LOGO_AVATAR_SIZE), Image.Resampling.LANCZOS)
    thumb = ImageOps.fit(img, (LOGO_THUMB_SIZE, LOGO_THUMB_SIZE), Image.Resampling.LANCZOS)

    def _to_png_bytes(im: Image.Image) -> bytes:
        buf = BytesIO()
        im.save(buf, format="PNG", optimize=True)
        return buf.getvalue()

    avatar_bytes = _to_png_bytes(avatar)
    thumb_bytes = _to_png_bytes(thumb)
    avatar_path = f"{APP_NAME}/orgs/{slug}/logo-{uuid.uuid4()}.png"
    thumb_path = f"{APP_NAME}/orgs/{slug}/logo-thumb-{uuid.uuid4()}.png"
    put_object(avatar_path, avatar_bytes, "image/png")
    put_object(thumb_path, thumb_bytes, "image/png")
    return avatar_path, thumb_path


def _process_and_upload_cover(slug: str, data: bytes) -> str:
    """Fit-crop to 1600x500 cover banner. Returns storage path."""
    img = _open_image(data)
    if img.mode == "RGBA":
        # Cover is a background — flatten to RGB for smaller JPEG
        bg = Image.new("RGB", img.size, (255, 255, 255))
        bg.paste(img, mask=img.split()[-1])
        img = bg
    elif img.mode != "RGB":
        img = img.convert("RGB")
    cover = ImageOps.fit(img, (COVER_WIDTH, COVER_HEIGHT), Image.Resampling.LANCZOS)
    buf = BytesIO()
    cover.save(buf, format="JPEG", quality=85, optimize=True, progressive=True)
    cover_path = f"{APP_NAME}/orgs/{slug}/cover-{uuid.uuid4()}.jpg"
    put_object(cover_path, buf.getvalue(), "image/jpeg")
    return cover_path


def _validate_image_upload(file: UploadFile, data: bytes) -> None:
    filename = file.filename or "image"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in IMAGE_EXT:
        raise HTTPException(400, "Only PNG, JPG or WebP allowed")
    if len(data) > MAX_IMAGE_BYTES:
        raise HTTPException(413, "Image too large (max 5 MB)")


@api.post("/organisations/{slug}/logo")
async def upload_org_logo(
    slug: str,
    file: UploadFile = File(...),
    org_auth: Optional[str] = Header(None, alias="X-Org-Auth"),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_org_write_access(slug, org_auth, admin_code)
    org = await _find_org(slug)
    data = await file.read()
    _validate_image_upload(file, data)
    try:
        avatar_path, thumb_path = _process_and_upload_logo(slug, data)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"Upload failed: {e}")
    # Best-effort cleanup of previous objects
    for old in (org.get("logo_path"), org.get("logo_thumb_path")):
        if old:
            try:
                requests.delete(f"{STORAGE_URL}/objects/{old}", headers={"X-Storage-Key": init_storage() or ""}, timeout=15)
            except Exception:
                pass
    await db.orgs.update_one(
        {"slug": slug},
        {"$set": {"logo_path": avatar_path, "logo_thumb_path": thumb_path, "updated_at": now_iso()}},
    )
    return {"ok": True, "logo_url": f"/api/organisations/{slug}/logo", "thumb_url": f"/api/organisations/{slug}/logo/thumb"}


@api.post("/organisations/{slug}/cover")
async def upload_org_cover(
    slug: str,
    file: UploadFile = File(...),
    org_auth: Optional[str] = Header(None, alias="X-Org-Auth"),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_org_write_access(slug, org_auth, admin_code)
    org = await _find_org(slug)
    data = await file.read()
    _validate_image_upload(file, data)
    try:
        cover_path = _process_and_upload_cover(slug, data)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"Upload failed: {e}")
    old = org.get("cover_path")
    if old:
        try:
            requests.delete(f"{STORAGE_URL}/objects/{old}", headers={"X-Storage-Key": init_storage() or ""}, timeout=15)
        except Exception:
            pass
    await db.orgs.update_one(
        {"slug": slug},
        {"$set": {"cover_path": cover_path, "updated_at": now_iso()}},
    )
    return {"ok": True, "cover_url": f"/api/organisations/{slug}/cover"}


@api.delete("/organisations/{slug}/logo")
async def delete_org_logo(
    slug: str,
    org_auth: Optional[str] = Header(None, alias="X-Org-Auth"),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_org_write_access(slug, org_auth, admin_code)
    await _find_org(slug)
    await db.orgs.update_one(
        {"slug": slug},
        {"$set": {"logo_path": None, "logo_thumb_path": None, "updated_at": now_iso()}},
    )
    return {"ok": True}


@api.delete("/organisations/{slug}/cover")
async def delete_org_cover(
    slug: str,
    org_auth: Optional[str] = Header(None, alias="X-Org-Auth"),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    _require_org_write_access(slug, org_auth, admin_code)
    await _find_org(slug)
    await db.orgs.update_one(
        {"slug": slug},
        {"$set": {"cover_path": None, "updated_at": now_iso()}},
    )
    return {"ok": True}


def _serve_org_image(path: Optional[str], content_type_default: str = "image/png"):
    if not path:
        raise HTTPException(404, "Image not set")
    try:
        data, ctype = get_object(path)
    except Exception as e:
        raise HTTPException(500, f"Fetch failed: {e}")
    return Response(
        content=data,
        media_type=content_type_default if not ctype or ctype == "application/octet-stream" else ctype,
        headers={"Cache-Control": "public, max-age=86400"},
    )


@api.get("/organisations/{slug}/logo")
async def get_org_logo(slug: str):
    org = await _find_org(slug)
    return _serve_org_image(org.get("logo_path"), "image/png")


@api.get("/organisations/{slug}/logo/thumb")
async def get_org_logo_thumb(slug: str):
    org = await _find_org(slug)
    return _serve_org_image(org.get("logo_thumb_path") or org.get("logo_path"), "image/png")


@api.get("/organisations/{slug}/cover")
async def get_org_cover(slug: str):
    org = await _find_org(slug)
    return _serve_org_image(org.get("cover_path"), "image/jpeg")


# ─────────── Event image uploads ───────────
EVENT_IMAGE_MAX_WIDTH = 1600


def _process_and_upload_event_image(data: bytes) -> str:
    img = _open_image(data)
    if img.mode == "RGBA":
        bg = Image.new("RGB", img.size, (255, 255, 255))
        bg.paste(img, mask=img.split()[-1])
        img = bg
    elif img.mode != "RGB":
        img = img.convert("RGB")
    if img.width > EVENT_IMAGE_MAX_WIDTH:
        ratio = EVENT_IMAGE_MAX_WIDTH / img.width
        img = img.resize((EVENT_IMAGE_MAX_WIDTH, max(1, int(img.height * ratio))), Image.Resampling.LANCZOS)
    buf = BytesIO()
    img.save(buf, format="JPEG", quality=85, optimize=True, progressive=True)
    name = f"{uuid.uuid4()}.jpg"
    put_object(f"{APP_NAME}/event-images/{name}", buf.getvalue(), "image/jpeg")
    return name


@api.post("/uploads/event-image")
async def upload_event_image(file: UploadFile = File(...)):
    data = await file.read()
    _validate_image_upload(file, data)
    try:
        name = _process_and_upload_event_image(data)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"Upload failed: {e}")
    return {"ok": True, "url": f"/api/event-images/{name}"}


@api.get("/event-images/{name}")
async def get_event_image(name: str):
    if not re.fullmatch(r"[a-f0-9\-]{36}\.jpg", name):
        raise HTTPException(404, "Image not found")
    try:
        data, _ = get_object(f"{APP_NAME}/event-images/{name}")
    except Exception:
        raise HTTPException(404, "Image not found")
    return Response(content=data, media_type="image/jpeg", headers={"Cache-Control": "public, max-age=86400"})


# ─────────── AI parse (multi-event) ───────────
class ParseRequest(BaseModel):
    text: str
    hint: Optional[str] = None


class ParsedItem(BaseModel):
    suggested_type: Literal["event", "volunteer", "organisation", "venue", "update"]
    title: str
    date: Optional[str] = None
    end_date: Optional[str] = None  # multi-day events (YYYY-MM-DD)
    start_time: Optional[str] = None
    end_time: Optional[str] = None
    location: Optional[str] = None
    category: str = "Community"
    description: str
    social_caption: str = ""
    notification_text: str = ""
    action: Literal["new_event", "update_event", "new_volunteer", "update_volunteer", "new_organisation", "update_organisation", "new_venue", "update_venue", "unclear"] = "unclear"
    matched_org_slug: Optional[str] = None
    matched_org_name: Optional[str] = None
    matched_event_id: Optional[str] = None
    matched_event_title: Optional[str] = None
    matched_volunteer_id: Optional[str] = None
    matched_volunteer_title: Optional[str] = None
    matched_venue_id: Optional[str] = None
    matched_venue_name: Optional[str] = None
    confidence: Optional[float] = None
    entity_confidence: Optional[float] = None
    raw_date_text: Optional[str] = None
    raw_time_text: Optional[str] = None
    date_confidence: Optional[float] = None
    time_confidence: Optional[float] = None
    # Recurrence detected from the source text — e.g. "Mondays, 10am" ⇒ weekly.
    recurrence_freq: Optional[Literal["none", "daily", "weekly", "biweekly", "monthly", "monthly_weekday", "annually"]] = None
    recurrence_weekday: Optional[str] = None  # e.g. "Monday" (informational)
    recurrence_confidence: Optional[float] = None
    recurrence_raw_text: Optional[str] = None
    # Extra structured fields (filled by the spreadsheet row parser)
    cost: Optional[str] = None
    booking: Optional[str] = None
    url: Optional[str] = None
    image: Optional[str] = None
    # Whether the category came from an explicit labeled field (skips inference)
    category_explicit: bool = False
    address: Optional[str] = None        # street address, distinct from location/venue name
    age: Optional[str] = None            # audience/age suitability
    accessibility: Optional[str] = None  # accessibility notes
    contact_email: Optional[str] = None
    contact_phone: Optional[str] = None


class ParsedDocument(BaseModel):
    filename: str
    source_type: str
    text_excerpt: str
    warnings: List[str] = []
    items: List[ParsedItem]


class BulkParseResponse(BaseModel):
    documents: List[ParsedDocument]
    mocked: bool = False


class ParseResponse(BaseModel):
    items: List[ParsedItem]
    mocked: bool = False


def _fallback_parse(text: str) -> ParsedItem:
    title = (text.strip().split("\n")[0][:80] or "Community update")
    date_m = re.search(r"\b(\d{1,2})(?:st|nd|rd|th)?\s+(January|February|March|April|May|June|July|August|September|October|November|December)(?:\s+\d{4})?\b", text, re.I)
    time_m = re.search(r"\b(\d{1,2})(?::(\d{2}))?\s*(am|pm)\b", text, re.I)
    is_event = bool(date_m or time_m)
    base = ParsedItem(
        suggested_type="event" if is_event else "update",
        title=title,
        date=date_m.group(0) if date_m else None,
        start_time=time_m.group(0) if time_m else None,
        location=None,
        category="Community",
        description=text.strip()[:500],
        social_caption=f"📣 {title} — happening in Blackrod. More on Blackrod Now.",
        notification_text=f"New on Blackrod Now: {title}",
    )
    return _normalize_parsed_item(base, text)


MONTHS = {
    "january": 1,
    "february": 2,
    "march": 3,
    "april": 4,
    "may": 5,
    "june": 6,
    "july": 7,
    "august": 8,
    "september": 9,
    "october": 10,
    "november": 11,
    "december": 12,
}


def _looks_iso_date(value: Optional[str]) -> bool:
    return bool(value and re.fullmatch(r"\d{4}-\d{2}-\d{2}", value.strip()))


def _looks_hhmm(value: Optional[str]) -> bool:
    return bool(value and re.fullmatch(r"(?:[01]\d|2[0-3]):[0-5]\d", value.strip()))


def _format_hhmm(hour: int, minute: int, meridiem: Optional[str]) -> str:
    if meridiem:
        m = meridiem.lower()
        if m == "pm" and hour < 12:
            hour += 12
        if m == "am" and hour == 12:
            hour = 0
    return f"{hour:02d}:{minute:02d}"


def _extract_iso_date(text: str) -> tuple[Optional[str], Optional[str], Optional[float]]:
    now = datetime.now(timezone.utc)
    year_default = now.year

    # Already ISO: 2026-07-21
    m = re.search(r"\b(\d{4})-(\d{2})-(\d{2})\b", text)
    if m:
        try:
            return datetime(int(m.group(1)), int(m.group(2)), int(m.group(3)), tzinfo=timezone.utc).strftime("%Y-%m-%d"), m.group(0), 0.95
        except ValueError:
            pass

    # 21-23 July (multi-day range) — the first day is the start date
    m = re.search(
        r"\b(\d{1,2})(?:st|nd|rd|th)?\s*(?:-|–|—|to|until)\s*\d{1,2}(?:st|nd|rd|th)?\s+"
        r"(january|february|march|april|may|june|july|august|september|october|november|december)"
        r"(?:\s+(\d{4}))?\b",
        text,
        re.I,
    )
    if m:
        day = int(m.group(1))
        month = MONTHS.get(m.group(2).lower())
        year = int(m.group(3)) if m.group(3) else year_default
        if month:
            try:
                return datetime(year, month, day, tzinfo=timezone.utc).strftime("%Y-%m-%d"), m.group(0), 0.85 if m.group(3) else 0.75
            except ValueError:
                pass

    # 14 June 2026 / Saturday 14 June
    m = re.search(
        r"\b(?:monday|tuesday|wednesday|thursday|friday|saturday|sunday)?\s*(\d{1,2})(?:st|nd|rd|th)?\s+"
        r"(january|february|march|april|may|june|july|august|september|october|november|december)"
        r"(?:\s+(\d{4}))?\b",
        text,
        re.I,
    )
    if m:
        day = int(m.group(1))
        month = MONTHS.get(m.group(2).lower())
        year = int(m.group(3)) if m.group(3) else year_default
        if month:
            try:
                return datetime(year, month, day, tzinfo=timezone.utc).strftime("%Y-%m-%d"), m.group(0), 0.9 if m.group(3) else 0.8
            except ValueError:
                pass

    # June 14 2026 / June 14
    m = re.search(
        r"\b(january|february|march|april|may|june|july|august|september|october|november|december)\s+"
        r"(\d{1,2})(?:st|nd|rd|th)?(?:,?\s*(\d{4}))?\b",
        text,
        re.I,
    )
    if m:
        month = MONTHS.get(m.group(1).lower())
        day = int(m.group(2))
        year = int(m.group(3)) if m.group(3) else year_default
        if month:
            try:
                return datetime(year, month, day, tzinfo=timezone.utc).strftime("%Y-%m-%d"), m.group(0), 0.85 if m.group(3) else 0.75
            except ValueError:
                pass

    # 14/06/2026 or 14-06-26 (UK day-first)
    m = re.search(r"\b(\d{1,2})[/-](\d{1,2})(?:[/-](\d{2,4}))?\b", text)
    if m:
        day = int(m.group(1))
        month = int(m.group(2))
        year_raw = m.group(3)
        year = year_default if not year_raw else int(year_raw)
        if year < 100:
            year += 2000
        try:
            return datetime(year, month, day, tzinfo=timezone.utc).strftime("%Y-%m-%d"), m.group(0), 0.7 if year_raw else 0.6
        except ValueError:
            pass

    return None, None, None


def _extract_hhmm_times(text: str) -> tuple[Optional[str], Optional[str], Optional[str], Optional[float]]:
    # 11am-4pm / 11:30am to 1pm / 10:00-12:30
    for m in re.finditer(
        r"\b(\d{1,2})(?::(\d{2}))?\s*(am|pm)?\s*(?:-|–|—|to)\s*"
        r"(\d{1,2})(?::(\d{2}))?\s*(am|pm)?\b",
        text,
        re.I,
    ):
        # Avoid treating date-like numeric spans (e.g. 16-07-26) as time ranges.
        if not any([m.group(2), m.group(3), m.group(5), m.group(6)]):
            continue
        h1 = int(m.group(1))
        mm1 = int(m.group(2) or "0")
        ampm1 = m.group(3)
        h2 = int(m.group(4))
        mm2 = int(m.group(5) or "0")
        ampm2 = m.group(6)
        if not ampm1 and ampm2:
            ampm1 = ampm2
        start = _format_hhmm(h1, mm1, ampm1)
        end = _format_hhmm(h2, mm2, ampm2)
        return start, end, m.group(0), 0.9 if (ampm1 or ampm2) else 0.8

    # 24h range: 10:00-12:30
    m = re.search(r"\b([01]?\d|2[0-3]):([0-5]\d)\s*(?:-|–|—|to)\s*([01]?\d|2[0-3]):([0-5]\d)\b", text)
    if m:
        return f"{int(m.group(1)):02d}:{m.group(2)}", f"{int(m.group(3)):02d}:{m.group(4)}", m.group(0), 0.8

    # Single 12h time
    m = re.search(r"\b(\d{1,2})(?::(\d{2}))?\s*(am|pm)\b", text, re.I)
    if m:
        h = int(m.group(1))
        mm = int(m.group(2) or "0")
        return _format_hhmm(h, mm, m.group(3)), None, m.group(0), 0.75

    # Single 24h time
    m = re.search(r"\b([01]?\d|2[0-3]):([0-5]\d)\b", text)
    if m:
        return f"{int(m.group(1)):02d}:{m.group(2)}", None, m.group(0), 0.7

    return None, None, None, None


def _detect_recurrence(text: str) -> Optional[dict]:
    """Detect recurring event patterns from free text.

    Handles the common Blackrod newsletter patterns:
      - "Mondays, 10am"           → weekly (Monday)
      - "Every Tuesday"           → weekly (Tuesday)
      - "Every other Wednesday"   → biweekly (Wednesday)
      - "Every fortnight"         → biweekly
      - "Weekly"/"weekly group"   → weekly
      - "Monthly"                 → monthly
      - "Daily"                   → daily
      - "Annual"/"Yearly"         → annually

    Returns a dict {freq, weekday?, confidence, raw_text} or None.
    """
    if not text:
        return None
    lower = text.lower()

    weekdays = {
        "monday": "Monday", "mondays": "Monday", "mon": "Monday",
        "tuesday": "Tuesday", "tuesdays": "Tuesday", "tue": "Tuesday", "tues": "Tuesday",
        "wednesday": "Wednesday", "wednesdays": "Wednesday", "wed": "Wednesday",
        "thursday": "Thursday", "thursdays": "Thursday", "thu": "Thursday", "thur": "Thursday", "thurs": "Thursday",
        "friday": "Friday", "fridays": "Friday", "fri": "Friday",
        "saturday": "Saturday", "saturdays": "Saturday", "sat": "Saturday",
        "sunday": "Sunday", "sundays": "Sunday", "sun": "Sunday",
    }

    # 1) "every other <weekday>" or "every second <weekday>" ⇒ biweekly
    m = re.search(r"every\s+(?:other|2nd|second)\s+(monday|tuesday|wednesday|thursday|friday|saturday|sunday)s?\b", lower)
    if m:
        wd = weekdays.get(m.group(1))
        return {"freq": "biweekly", "weekday": wd, "confidence": 0.9, "raw_text": m.group(0)}

    # 2) "every <weekday>" ⇒ weekly
    m = re.search(r"every\s+(monday|tuesday|wednesday|thursday|friday|saturday|sunday)s?\b", lower)
    if m:
        wd = weekdays.get(m.group(1))
        return {"freq": "weekly", "weekday": wd, "confidence": 0.95, "raw_text": m.group(0)}

    # 3) plural weekday at the *start of a line or bullet* ("Mondays, 10am")
    for token, canonical in weekdays.items():
        if not token.endswith("s"):
            continue
        # Match as a standalone word followed by a comma / space / punctuation.
        pattern = rf"(?m)(?:^|[•*·\-\|]\s*|\s){re.escape(token)}\b[\s,]"
        if re.search(pattern, lower):
            return {"freq": "weekly", "weekday": canonical, "confidence": 0.85, "raw_text": token}

    # 4) "every fortnight" / "fortnightly" / "bi-weekly" ⇒ biweekly
    if re.search(r"\b(fortnightly|every\s+fortnight|bi[- ]?weekly)\b", lower):
        return {"freq": "biweekly", "weekday": None, "confidence": 0.85, "raw_text": "fortnightly"}

    # 5) "every week" / "weekly" ⇒ weekly (only when clearly a recurring group)
    if re.search(r"\b(every\s+week|weekly\s+(group|session|meeting|club|meet\s*up|meetup|drop[- ]?in))\b", lower):
        return {"freq": "weekly", "weekday": None, "confidence": 0.75, "raw_text": "weekly"}

    # 6) "first/second/.../last Thursday of the month" ⇒ monthly_weekday
    m = re.search(
        r"\b(first|second|third|fourth|fifth|last)\s+(monday|tuesday|wednesday|thursday|friday|saturday|sunday)\s+of\s+(the\s+)?month\b",
        lower,
    )
    if m:
        wd = weekdays.get(m.group(2))
        return {"freq": "monthly_weekday", "weekday": wd, "confidence": 0.92, "raw_text": m.group(0)}

    # 7) "monthly" / "every month" ⇒ monthly
    if re.search(r"\b(monthly|every\s+month)\b", lower):
        return {"freq": "monthly", "weekday": None, "confidence": 0.8, "raw_text": "monthly"}

    # 8) "daily" / "every day" ⇒ daily
    if re.search(r"\b(daily|every\s+day)\b", lower):
        return {"freq": "daily", "weekday": None, "confidence": 0.85, "raw_text": "daily"}

    # 9) "annual" / "yearly" / "every year" ⇒ annually
    if re.search(r"\b(annual(?:ly)?|yearly|every\s+year)\b", lower):
        return {"freq": "annually", "weekday": None, "confidence": 0.85, "raw_text": "annual"}

    return None


# ─────────── Deterministic field extractors (fill event details from text) ───────────
UK_POSTCODE_RE = r"[A-Za-z]{1,2}\d[A-Za-z\d]?\s*\d[A-Za-z]{2}"

VENUE_KEYWORDS = (
    "hall", "church", "chapel", "library", "centre", "center", "club", "school",
    "park", "pavilion", "playing fields?", "ground", "hotel", "pub", "inn", "cafe",
    "café", "studio", "rooms?", "museum", "green", "institute", "academy", "surgery",
)


def _extract_contact_email(text: str) -> Optional[str]:
    for m in re.finditer(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", text or ""):
        email = m.group(0)
        if not re.match(r"(?i)^(no-?reply|donotreply)@", email):
            return email
    return None


def _extract_contact_phone(text: str) -> Optional[str]:
    # UK formats: 01204 696295, (01204) 696295, 07123 456789, +44 1204 696295
    m = re.search(r"(?:\+44\s?\(?0?\)?[\s-]?|\(?0)\d{2,4}\)?[\s-]?\d{3,4}[\s-]?\d{3,4}\b", text or "")
    if m:
        candidate = m.group(0).strip()
        digits = re.sub(r"\D", "", candidate)
        if 10 <= len(digits) <= 13:
            return candidate
    return None


def _extract_url(text: str) -> Optional[str]:
    m = re.search(r"https?://[^\s)\]>\"']+", text or "", re.I)
    if m:
        return m.group(0).rstrip(".,;!?")
    m = re.search(r"\bwww\.[a-z0-9.-]+\.[a-z]{2,}(?:/[^\s)\]>\"']*)?", text or "", re.I)
    if m:
        return f"https://{m.group(0).rstrip('.,;!?')}"
    return None


def _extract_cost(text: str) -> Optional[str]:
    if not text:
        return None
    # "adults £5, children £3" style phrases first (richest info)
    m = re.search(r"(?i)\b(?:adults?|children|kids|concessions?|members?|non-members?)\s*[:\-]?\s*£\s?\d+(?:\.\d{2})?(?:[^.\n]{0,60}£\s?\d+(?:\.\d{2})?)?", text)
    if m:
        return re.sub(r"\s+", " ", m.group(0)).strip()
    # "£3", "£3.50 per person", "entry £2", "tickets £5"
    m = re.search(r"(?i)£\s?\d+(?:\.\d{2})?(?:\s*(?:per|each|pp|a)\s*(?:person|adult|child|family|ticket|session|head)?)?", text)
    if m:
        return re.sub(r"\s+", " ", m.group(0)).strip()
    if re.search(r"(?i)\b(free\s+(?:entry|admission|event|to\s+attend)|entry\s+(?:is\s+)?free|admission\s+free|no\s+charge|free\s+of\s+charge)\b", text):
        return "Free"
    if re.search(r"(?i)(?:^|[\s|:])free(?:[\s|.,!]|$)", text):
        return "Free"
    if re.search(r"(?i)\bdonations?\s+(?:welcome|appreciated|only|invited)\b", text):
        return "Donations welcome"
    return None


def _extract_booking(text: str) -> Optional[str]:
    if not text:
        return None
    for line in re.split(r"[\n.]", text):
        line = line.strip()
        if not line or len(line) > 160:
            continue
        if re.search(r"(?i)\b(book(?:ing)?|tickets?|reserve|register|rsvp|sign\s*up|just\s+turn\s+up|no\s+booking)\b", line):
            return re.sub(r"\s+", " ", line)
    return None


def _extract_age(text: str) -> Optional[str]:
    if not text:
        return None
    patterns = [
        r"(?i)\ball\s+ages(?:\s+welcome)?\b",
        r"(?i)\bages?\s+\d{1,2}\s*(?:-|–|to)\s*\d{1,2}\b",
        r"(?i)\b(?:under|over)[\s-]?\d{1,2}s?\b",
        r"(?i)\b\d{1,2}\+\s*(?:only|years)?",
        r"(?i)\badults?\s+only\b",
        r"(?i)\bfamily[\s-]friendly\b",
        r"(?i)\bsuitable\s+for\s+[^.\n]{3,40}",
    ]
    for pattern in patterns:
        m = re.search(pattern, text)
        if m:
            return re.sub(r"\s+", " ", m.group(0)).strip().rstrip(",")
    return None


def _extract_accessibility(text: str) -> Optional[str]:
    if not text:
        return None
    notes: list[str] = []
    for pattern in [
        r"(?i)\bwheelchair[\s-]?(?:accessible|friendly|access)\b",
        r"(?i)\bstep[\s-]?free(?:\s+access)?\b",
        r"(?i)\baccessible\s+(?:toilets?|parking|entrance)\b",
        r"(?i)\bhearing\s+loop\b",
        r"(?i)\bdisabled\s+(?:parking|access|toilets?)\b",
        r"(?i)\b(?:autism|dementia)[\s-]?friendly\b",
        r"(?i)\bBSL(?:\s+interpret\w*)?\b",
    ]:
        m = re.search(pattern, text)
        if m:
            notes.append(re.sub(r"\s+", " ", m.group(0)).strip())
    return "; ".join(dict.fromkeys(notes)) or None


def _extract_address(text: str) -> Optional[str]:
    if not text:
        return None
    # Prefer a line containing a UK postcode.
    for line in text.splitlines():
        line = line.strip().strip("|•*·")
        if re.search(rf"\b{UK_POSTCODE_RE}\b", line) and len(line) <= 120:
            cleaned = re.sub(r"(?i)^(address|venue|where|location)\s*[:\-]\s*", "", line).strip()
            if cleaned:
                return cleaned
    # "15 Church Road" style street address.
    m = re.search(
        r"\b\d{1,4}[A-Za-z]?\s+[A-Z][\w']*(?:\s+[A-Z][\w']*)*\s+"
        r"(?:Road|Street|Lane|Avenue|Drive|Close|Way|Court|Crescent|Terrace|Place|Row|Grove|Gardens|Rd|Ave|Ln)\b"
        r"(?:,\s*[A-Z][\w' ]+)*",
        text,
    )
    if m:
        return re.sub(r"\s+", " ", m.group(0)).strip()
    return None


def _extract_venue(text: str) -> Optional[str]:
    if not text:
        return None
    keyword_alt = "|".join(VENUE_KEYWORDS)
    # Explicit "Venue: X" / "Where: X" lines.
    m = re.search(r"(?im)^(?:venue|where|location)\s*[:\-]\s*(.+)$", text)
    if m:
        value = m.group(1).strip().rstrip(".,;")
        if value and len(value) <= 100:
            return value
    # "at St Katharine's Church Hall" — capitalised phrase ending in a venue keyword.
    m = re.search(
        rf"(?:\bat|\bheld at|\btakes place at|@)\s+(?:the\s+)?"
        rf"((?:[\w'’&.-]+\s+){{0,5}}(?:{keyword_alt}))\b",
        text,
        re.I,
    )
    if m:
        candidate = re.sub(r"\s+", " ", m.group(1)).strip().rstrip(".,;")
        # Require at least one capitalised word so we skip "at the hall".
        if re.search(r"[A-Z]", candidate) and 3 <= len(candidate) <= 80:
            return candidate
    return None


CATEGORY_KEYWORDS: list[tuple[str, str]] = [
    (r"\b(volunteer|volunteering|helpers?\s+needed)\b", "Volunteering"),
    (r"\b(football|rugby|cricket|netball|tennis|running|fitness|yoga|pilates|zumba|swim|sports?|athletics|bowls|walking\s+group)\b", "Sport"),
    (r"\b(choir|concert|band|music|gig|singing|orchestra|karaoke|open\s+mic)\b", "Music"),
    (r"\b(market|fair|fete|food|drink|bake|coffee\s+morning|tea|lunch|supper|bbq|barbecue)\b", "Food & Drink"),
    # "church" alone is too common in addresses (Church Road/Street) — require worship context.
    (r"\b(mass|worship|prayer|faith|carol\s+service|parish|chapel|church\s+(service|hall|group|congregation))\b", "Faith"),
    (r"\b(heritage|history|historical|museum|archive)\b", "Heritage"),
    (r"\b(health|wellbeing|well-being|mental\s+health|mindfulness|dementia|carers?)\b", "Health & Wellbeing"),
    (r"\b(school|education|class|course|lesson|tuition|homework)\b", "School"),
    (r"\b(charity|fundrais\w+|donation|raffle|tombola|sponsored)\b", "Charity"),
    (r"\b(youth|teen(?:ager)?s?|young\s+people)\b", "Youth"),
    (r"\b(family|families|kids|children|toddlers?|under\s+5s?|craft)\b", "Family"),
    (r"\b(business|networking|traders?|enterprise)\b", "Business"),
]


def _infer_category(text: str) -> Optional[str]:
    lower = (text or "").lower()
    for pattern, category in CATEGORY_KEYWORDS:
        if re.search(pattern, lower):
            return category
    return None


def _extract_end_date(text: str, start_date: Optional[str]) -> Optional[str]:
    """Detect a multi-day range like '21-23 July' or '14 June - 16 June'."""
    if not text:
        return None
    m = re.search(
        r"\b(\d{1,2})(?:st|nd|rd|th)?\s*(?:-|–|—|to|until)\s*(\d{1,2})(?:st|nd|rd|th)?\s+"
        r"(january|february|march|april|may|june|july|august|september|october|november|december)"
        r"(?:\s+(\d{4}))?\b",
        text,
        re.I,
    )
    if m:
        month = MONTHS.get(m.group(3).lower())
        year = int(m.group(4)) if m.group(4) else datetime.now(timezone.utc).year
        try:
            end = datetime(year, month, int(m.group(2)), tzinfo=timezone.utc).strftime("%Y-%m-%d")
        except (ValueError, TypeError):
            return None
        if not start_date or end > start_date:
            return end
    # '14 June - 16 July' style: two full dates joined by a range separator.
    m = re.search(
        r"(?:-|–|—|\bto\b|\buntil\b)\s*((?:\d{1,2})(?:st|nd|rd|th)?\s+"
        r"(?:january|february|march|april|may|june|july|august|september|october|november|december)(?:\s+\d{4})?)",
        text,
        re.I,
    )
    if m and start_date:
        end_iso, _, _ = _extract_iso_date(m.group(1))
        if end_iso and end_iso > start_date:
            return end_iso
    return None


def _normalize_parsed_item(item: ParsedItem, source_text: str) -> ParsedItem:
    merged = "\n".join([item.title or "", item.description or "", source_text or ""]).strip()
    updates: dict[str, Any] = {}

    if item.confidence is not None and item.entity_confidence is None:
        updates["entity_confidence"] = item.confidence

    if not _looks_iso_date(item.date):
        extracted_date, raw_date_text, date_conf = _extract_iso_date(merged)
        if extracted_date:
            updates["date"] = extracted_date
            updates["raw_date_text"] = raw_date_text
            updates["date_confidence"] = date_conf
    elif item.date and item.date_confidence is None:
        updates["date_confidence"] = 0.95
        updates["raw_date_text"] = item.date

    if not _looks_hhmm(item.start_time) or (item.end_time and not _looks_hhmm(item.end_time)):
        start_time, end_time, raw_time_text, time_conf = _extract_hhmm_times(merged)
        if start_time and not _looks_hhmm(item.start_time):
            updates["start_time"] = start_time
        if end_time and not _looks_hhmm(item.end_time):
            updates["end_time"] = end_time
        if (start_time or end_time) and time_conf is not None:
            updates["time_confidence"] = time_conf
            updates["raw_time_text"] = raw_time_text
    elif item.start_time and item.time_confidence is None:
        updates["time_confidence"] = 0.95
        if item.end_time:
            updates["raw_time_text"] = f"{item.start_time}-{item.end_time}"
        else:
            updates["raw_time_text"] = item.start_time

    # Recurrence: only auto-detect for events, and only if the AI didn't already return one.
    if item.suggested_type == "event" and not item.recurrence_freq:
        rec = _detect_recurrence(merged)
        if rec:
            updates["recurrence_freq"] = rec["freq"]
            updates["recurrence_weekday"] = rec.get("weekday")
            updates["recurrence_confidence"] = rec.get("confidence")
            updates["recurrence_raw_text"] = rec.get("raw_text")

    # Fill any remaining event fields deterministically so events arrive fully populated
    # regardless of which pipeline (AI, fallback, OCR, spreadsheet, URL) produced the item.
    if item.suggested_type == "event":
        if not item.end_date:
            end_date = _extract_end_date(merged, updates.get("date") or item.date)
            if end_date:
                updates["end_date"] = end_date
        if not item.location:
            venue = _extract_venue(merged)
            if venue:
                updates["location"] = venue
        if not item.address:
            address = _extract_address(merged)
            if address:
                updates["address"] = address
        if not item.cost:
            cost = _extract_cost(merged)
            if cost:
                updates["cost"] = cost
        if not item.booking:
            booking = _extract_booking(merged)
            if booking:
                updates["booking"] = booking
        if not item.url:
            url = _extract_url(merged)
            if url:
                updates["url"] = url
        if not item.age:
            age = _extract_age(merged)
            if age:
                updates["age"] = age
        if not item.accessibility:
            accessibility = _extract_accessibility(merged)
            if accessibility:
                updates["accessibility"] = accessibility
        if not item.contact_email:
            email = _extract_contact_email(merged)
            if email:
                updates["contact_email"] = email
        if not item.contact_phone:
            phone = _extract_contact_phone(merged)
            if phone:
                updates["contact_phone"] = phone
        if (not item.category or item.category == "Community") and not item.category_explicit:
            inferred = _infer_category(merged)
            if inferred:
                updates["category"] = inferred

    if updates:
        return item.model_copy(update=updates)
    return item


def _clean_text_excerpt(text: str, limit: int = 240) -> str:
    return re.sub(r"\s+", " ", (text or "")).strip()[:limit]


def _extract_pptx_text(data: bytes) -> tuple[str, list[str]]:
    warnings: list[str] = []
    try:
        namespaces = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
        slides: list[str] = []
        with zipfile.ZipFile(BytesIO(data)) as archive:
            slide_names = sorted(
                (n for n in archive.namelist() if re.fullmatch(r"ppt/slides/slide\d+\.xml", n)),
                key=lambda n: int(re.search(r"(\d+)", n).group(1)),
            )
            for name in slide_names:
                root = ET.fromstring(archive.read(name))
                texts = [node.text.strip() for node in root.findall(".//a:t", namespaces) if node.text and node.text.strip()]
                if texts:
                    slides.append("\n".join(texts))
        return "\n\n".join(slides).strip(), warnings
    except Exception as exc:
        warnings.append(f"PPTX extract failed: {exc}")
        return "", warnings


def _extract_docx_text(data: bytes) -> tuple[str, list[str]]:
    warnings: list[str] = []
    try:
        with zipfile.ZipFile(BytesIO(data)) as archive:
            xml_bytes = archive.read("word/document.xml")
        root = ET.fromstring(xml_bytes)
        namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        paragraphs: list[str] = []
        for para in root.findall(".//w:p", namespaces):
            pieces: list[str] = []
            for node in para.iter():
                tag = node.tag.rsplit("}", 1)[-1]
                if tag == "t" and node.text:
                    pieces.append(node.text)
                elif tag == "br":
                    pieces.append("\n")  # preserve Word soft line breaks
                elif tag == "tab":
                    pieces.append("\t")
            line = "".join(pieces).strip()
            if line:
                paragraphs.append(line)
        return "\n".join(paragraphs).strip(), warnings
    except Exception as exc:
        warnings.append(f"DOCX extract failed: {exc}")
        return "", warnings


def _extract_pdf_text(data: bytes) -> tuple[str, list[str]]:
    warnings: list[str] = []
    try:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(data))
        chunks: list[str] = []
        qr_values: list[str] = []
        for page in reader.pages:
            try:
                chunk = page.extract_text() or ""
            except Exception:
                chunk = ""
            if chunk.strip():
                chunks.append(chunk.strip())
            # Also attempt QR decode from embedded raster images in the page.
            try:
                images = getattr(page, "images", []) or []
                for image in images:
                    image_bytes = getattr(image, "data", None)
                    if not image_bytes:
                        continue
                    values, _ = _decode_qr_values_from_image(image_bytes)
                    qr_values.extend(values)
            except Exception:
                pass

        if qr_values:
            qr_values = list(dict.fromkeys(v.strip() for v in qr_values if v and v.strip()))
            chunks.append("\n".join([f"QR code link: {v}" for v in qr_values]))
            warnings.append(f"Detected {len(qr_values)} QR code link{'s' if len(qr_values) != 1 else ''} in PDF: {' | '.join(qr_values)}")
        text = "\n\n".join(chunks).strip()
        if len(text) < 40:
            # Little or no embedded text — likely a scanned/exported flyer. OCR the pages.
            ocr_text, ocr_warnings = _ocr_pdf_pages(data)
            warnings.extend(ocr_warnings)
            if len(ocr_text.strip()) > len(text):
                text = "\n\n".join(filter(None, [ocr_text.strip(), "\n".join(f"QR code link: {v}" for v in qr_values)])).strip()
        if not text:
            warnings.append("No readable text found in PDF")
        return text, warnings
    except Exception as exc:
        warnings.append(f"PDF extract failed: {exc}")
        return "", warnings


MAX_PDF_OCR_PAGES = int(os.environ.get("MAX_PDF_OCR_PAGES", "4"))


def _ocr_pdf_pages(data: bytes) -> tuple[str, list[str]]:
    """Rasterize PDF pages and OCR them — handles scanned flyers and Canva-style image PDFs."""
    warnings: list[str] = []
    try:
        import pypdfium2 as pdfium
    except Exception:
        warnings.append("Scanned PDF detected but OCR rasterizer is unavailable")
        return "", warnings
    try:
        pdf = pdfium.PdfDocument(BytesIO(data))
        page_count = len(pdf)
        ocr = _get_ocr_engine()
        chunks: list[str] = []
        for page_index in range(min(page_count, MAX_PDF_OCR_PAGES)):
            page = pdf[page_index]
            bitmap = page.render(scale=2.0)
            pil_image = bitmap.to_pil().convert("RGB")
            buffer = BytesIO()
            pil_image.save(buffer, format="PNG")
            result, _ = ocr(buffer.getvalue())
            lines = [_normalize_ocr_line(str(item[1]).strip()) for item in (result or []) if len(item) >= 2 and item[1] and str(item[1]).strip()]
            if lines:
                chunks.append("\n".join(lines))
        if page_count > MAX_PDF_OCR_PAGES:
            warnings.append(f"OCR limited to the first {MAX_PDF_OCR_PAGES} of {page_count} pages")
        text = "\n\n".join(chunks).strip()
        if text:
            warnings.append("PDF had no embedded text — content recovered with OCR")
        return text, warnings
    except Exception as exc:
        warnings.append(f"PDF OCR failed: {exc}")
        return "", warnings


def _extract_xlsx_text(data: bytes) -> tuple[str, list[str]]:
    warnings: list[str] = []
    try:
        from openpyxl import load_workbook

        workbook = load_workbook(BytesIO(data), read_only=True, data_only=True)
        lines: list[str] = []
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows(values_only=True):
                cells = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
                if cells:
                    lines.append(" | ".join(cells))
        text = "\n".join(lines).strip()
        if not text:
            warnings.append("No readable text found in spreadsheet")
        return text, warnings
    except Exception as exc:
        warnings.append(f"Spreadsheet extract failed: {exc}")
        return "", warnings


@lru_cache(maxsize=1)
def _get_ocr_engine():
    from rapidocr_onnxruntime import RapidOCR

    return RapidOCR()


@lru_cache(maxsize=1)
def _register_heif_opener() -> bool:
    try:
        from pillow_heif import register_heif_opener

        register_heif_opener()
        return True
    except Exception:
        return False


def _normalize_ocr_line(line: str) -> str:
    """Re-insert word spacing OCR often drops (e.g. 'Saturday12September2026').
    Short letter runs before digits are left alone to protect postcodes ('BL6') and times ('11am')."""
    line = re.sub(r"([a-z])([A-Z])", r"\1 \2", line)
    line = re.sub(r"([A-Za-z]{3,})(\d)", r"\1 \2", line)
    line = re.sub(r"(\d)([A-Z][a-z]{2,})", r"\1 \2", line)
    return line


def _extract_image_text(data: bytes) -> tuple[str, list[str]]:
    warnings: list[str] = []
    _register_heif_opener()
    try:
        with Image.open(BytesIO(data)) as img:
            img = ImageOps.exif_transpose(img)
            if img.mode not in {"RGB", "L"}:
                img = img.convert("RGB")
            buffer = BytesIO()
            img.save(buffer, format="PNG")
            ocr = _get_ocr_engine()
            result, _ = ocr(buffer.getvalue())
    except Exception as exc:
        warnings.append(f"Image OCR failed: {exc}")
        return "", warnings

    lines: list[str] = []
    for item in result or []:
        if len(item) >= 2 and item[1]:
            text = str(item[1]).strip()
            if text:
                lines.append(_normalize_ocr_line(text))

    qr_values, qr_warnings = _decode_qr_values_from_image(data)
    if qr_values:
        deduped = list(dict.fromkeys(v.strip() for v in qr_values if v and v.strip()))
        lines.extend([f"QR code link: {v}" for v in deduped])
        warnings.append(f"Detected {len(deduped)} QR code link{'s' if len(deduped) != 1 else ''} in image: {' | '.join(deduped)}")
    warnings.extend(qr_warnings)

    text = "\n".join(lines).strip()
    if not text:
        warnings.append("No readable text found in image")
    return text, warnings


def _decode_qr_values_from_image(data: bytes) -> tuple[list[str], list[str]]:
    warnings: list[str] = []
    try:
        import cv2
        import numpy as np
    except Exception:
        # QR decode is optional; OCR-only path still works.
        return [], warnings

    try:
        arr = np.frombuffer(data, dtype=np.uint8)
        image = cv2.imdecode(arr, cv2.IMREAD_COLOR)
        if image is None:
            return [], warnings

        detector = cv2.QRCodeDetector()
        values: list[str] = []

        try:
            multi = detector.detectAndDecodeMulti(image)
            decoded_info: list[str] = []
            if isinstance(multi, tuple):
                if len(multi) == 4:
                    _, decoded_info, _, _ = multi
                elif len(multi) == 3:
                    decoded_info, _, _ = multi
            for value in decoded_info or []:
                if isinstance(value, str) and value.strip():
                    values.append(value.strip())
        except Exception:
            pass

        if not values:
            try:
                single, _, _ = detector.detectAndDecode(image)
                if isinstance(single, str) and single.strip():
                    values.append(single.strip())
            except Exception:
                pass

        return list(dict.fromkeys(values)), warnings
    except Exception as exc:
        warnings.append(f"QR decode failed: {exc}")
        return [], warnings


def _extract_document_text(filename: str, content_type: str, data: bytes) -> tuple[str, str, list[str]]:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    warnings: list[str] = []
    if ext in {"txt", "md", "rtf", "csv"} or content_type.startswith("text/"):
        try:
            if ext == "csv":
                sample = data.decode("utf-8", errors="ignore")
                rows = []
                for row in csv.reader(sample.splitlines()):
                    row = [cell.strip() for cell in row if cell and cell.strip()]
                    if row:
                        rows.append(" | ".join(row))
                return "\n".join(rows).strip(), "csv", warnings
            return data.decode("utf-8", errors="ignore").strip(), "text", warnings
        except Exception as exc:
            warnings.append(f"Text decode failed: {exc}")
            return "", "text", warnings
    if ext == "docx":
        text, extra = _extract_docx_text(data)
        return text, "docx", warnings + extra
    if ext == "pptx":
        text, extra = _extract_pptx_text(data)
        return text, "pptx", warnings + extra
    if ext == "pdf":
        text, extra = _extract_pdf_text(data)
        return text, "pdf", warnings + extra
    if ext == "xlsx":
        text, extra = _extract_xlsx_text(data)
        return text, "xlsx", warnings + extra
    if ext in {"png", "jpg", "jpeg", "webp", "bmp", "gif", "tif", "tiff", "heic", "heif"} or content_type.startswith("image/"):
        text, extra = _extract_image_text(data)
        return text, "image", warnings + extra
    if ext in {"doc", "xls", "ppt"}:
        modern = {"doc": "DOCX", "xls": "XLSX", "ppt": "PPTX"}[ext]
        warnings.append(f"Legacy .{ext} files can't be read — please re-save the file as {modern} and upload again")
        return "", ext, warnings
    warnings.append(f"Unsupported auto-parse format: .{ext or 'bin'}")
    return "", ext or "bin", warnings


def _extract_html_text(html_text: str) -> tuple[str, list[str]]:
    warnings: list[str] = []
    try:
        title_match = re.search(r"<title[^>]*>(.*?)</title>", html_text, re.I | re.S)
        og_title_match = re.search(r'<meta[^>]+property=["\']og:title["\'][^>]+content=["\']([^"\']+)["\']', html_text, re.I)
        og_desc_match = re.search(r'<meta[^>]+property=["\']og:description["\'][^>]+content=["\']([^"\']+)["\']', html_text, re.I)
        desc_match = re.search(r'<meta[^>]+name=["\']description["\'][^>]+content=["\']([^"\']+)["\']', html_text, re.I)
        text = re.sub(r"(?is)<(script|style|noscript)[^>]*>.*?</\1>", " ", html_text)
        text = re.sub(r"(?s)<[^>]+>", " ", text)
        text = html.unescape(text)
        chunks = [
            (og_title_match.group(1) if og_title_match else "").strip(),
            (title_match.group(1) if title_match else "").strip(),
            (og_desc_match.group(1) if og_desc_match else "").strip(),
            (desc_match.group(1) if desc_match else "").strip(),
            re.sub(r"\s+", " ", text).strip(),
        ]
        cleaned = "\n".join([chunk for chunk in chunks if chunk]).strip()
        if not cleaned:
            warnings.append("No readable text found in web page")
        return cleaned, warnings
    except Exception as exc:
        warnings.append(f"Web page extract failed: {exc}")
        return "", warnings


def _extract_url_text(url: str) -> tuple[str, str, list[str]]:
    warnings: list[str] = []
    try:
        response = requests.get(
            url,
            timeout=20,
            headers={"User-Agent": "BlackrodNowBot/1.0 (+https://blackrodnow.local)"},
            allow_redirects=True,
        )
        response.raise_for_status()
    except Exception as exc:
        warnings.append(f"URL fetch failed: {exc}")
        return "", "url", warnings

    content_type = (response.headers.get("Content-Type") or "").split(";")[0].strip().lower()
    final_url = response.url or url
    path = urlparse(final_url).path
    ext = path.rsplit(".", 1)[-1].lower() if "." in path else ""

    if content_type == "text/html" or ext in {"", "html", "htm", "php", "asp", "aspx"}:
        text, extra = _extract_html_text(response.text)
        return text, "url", warnings + extra

    filename = path.rsplit("/", 1)[-1] or "link"
    if "." not in filename:
        filename = f"{filename or 'link'}.bin"
    text, source_type, extra = _extract_document_text(filename, content_type or "application/octet-stream", response.content)
    return text, source_type, warnings + extra


def _best_match(query: str, choices: list[dict], key: str) -> tuple[Optional[dict], float]:
    query_norm = query.lower().strip()
    best_item: Optional[dict] = None
    best_score = 0.0
    for item in choices:
        candidate = (item.get(key) or "").lower().strip()
        if not candidate:
            continue
        score = SequenceMatcher(None, query_norm, candidate).ratio()
        if candidate in query_norm or query_norm in candidate:
            score = max(score, 0.92)
        if score > best_score:
            best_item = item
            best_score = score
    return best_item, best_score


def _fallback_document_classify(filename: str, text: str, orgs: list[dict], events: list[dict], volunteers: list[dict], venues: list[dict]) -> list[ParsedItem]:
    base = _fallback_parse(text)
    lower = f"{filename}\n{text}".lower()
    has_event_signal = bool(re.search(r"\b(event|fair|match|meet|open day|session|workshop|talk|concert|race|fundraiser|fund raising)\b", lower)) or bool(re.search(r"\b\d{1,2}(:\d{2})?\s*(am|pm)\b", lower))
    has_volunteer_signal = bool(re.search(r"\b(volunteer|volunteering|helper|helpers needed|join our team|steward|coach wanted|mentor|befriender|trustee|committee member)\b", lower))
    has_org_signal = bool(re.search(r"\b(organisation|organization|club|group|society|association|committee|charity|church|school)\b", lower))
    has_venue_signal = bool(re.search(r"\b(venue|hall|community centre|community center|sports centre|sports center|facility|facilities|room hire|booking office|address)\b", lower))
    matched_org, org_score = _best_match(base.title, orgs, "name")
    matched_event, event_score = _best_match(base.title, events, "title")
    matched_volunteer, volunteer_score = _best_match(base.title, volunteers, "title")
    matched_venue, venue_score = _best_match(base.title, venues, "name")
    if matched_venue and venue_score >= max(org_score, event_score, volunteer_score) and venue_score >= 0.65:
        payload = base.model_dump()
        payload.update({
            "suggested_type": "venue",
            "action": "update_venue" if has_venue_signal else "new_venue",
            "matched_venue_id": matched_venue.get("id"),
            "matched_venue_name": matched_venue.get("name"),
            "confidence": round(venue_score, 2),
            "entity_confidence": round(venue_score, 2),
        })
        return [ParsedItem(**payload)]

    if matched_volunteer and volunteer_score >= max(org_score, event_score) and volunteer_score >= 0.65:
        payload = base.model_dump()
        payload.update({
            "suggested_type": "volunteer",
            "action": "update_volunteer" if has_volunteer_signal else "new_volunteer",
            "matched_volunteer_id": matched_volunteer.get("id"),
            "matched_volunteer_title": matched_volunteer.get("title"),
            "matched_org_slug": matched_volunteer.get("orgSlug"),
            "confidence": round(volunteer_score, 2),
            "entity_confidence": round(volunteer_score, 2),
        })
        return [ParsedItem(**payload)]
    if matched_event and event_score >= org_score and event_score >= 0.65:
        payload = base.model_dump()
        payload.update({
            "suggested_type": "event",
            "action": "update_event" if has_event_signal else "new_event",
            "matched_event_id": matched_event.get("id"),
            "matched_event_title": matched_event.get("title"),
            "confidence": round(event_score, 2),
            "entity_confidence": round(event_score, 2),
        })
        return [ParsedItem(**payload)]
    if matched_org and org_score >= 0.65:
        payload = base.model_dump()
        payload.update({
            "suggested_type": "organisation",
            "action": "update_organisation" if has_org_signal else "new_organisation",
            "matched_org_slug": matched_org.get("slug"),
            "matched_org_name": matched_org.get("name"),
            "confidence": round(org_score, 2),
            "entity_confidence": round(org_score, 2),
        })
        return [ParsedItem(**payload)]
    if has_volunteer_signal:
        payload = base.model_dump()
        payload.update({"suggested_type": "volunteer", "action": "new_volunteer", "confidence": 0.55, "entity_confidence": 0.55})
        return [ParsedItem(**payload)]
    if has_event_signal:
        payload = base.model_dump()
        payload.update({"suggested_type": "event", "action": "new_event", "confidence": 0.55, "entity_confidence": 0.55})
        return [ParsedItem(**payload)]
    if has_org_signal:
        payload = base.model_dump()
        payload.update({"suggested_type": "organisation", "action": "new_organisation", "confidence": 0.55, "entity_confidence": 0.55})
        return [ParsedItem(**payload)]
    if has_venue_signal:
        payload = base.model_dump()
        payload.update({"suggested_type": "venue", "action": "new_venue", "confidence": 0.55, "entity_confidence": 0.55})
        return [ParsedItem(**payload)]
    return [base]


# ─────────── Structured spreadsheet (columns → events) ───────────
SITE_CATEGORIES = [
    "Family", "Youth", "Sport", "School", "Charity", "Business", "Community",
    "Music", "Food & Drink", "Volunteering", "Faith", "Heritage", "Health & Wellbeing",
]

SPREADSHEET_HEADERS = {
    "organisation": {"organisation", "organization", "org", "group", "club", "host", "provider", "organiser", "organizer", "hostedby", "runby"},
    "title": {"title", "event", "eventtitle", "eventname", "name", "activity", "whatson"},
    "date": {"date", "dates", "eventdate", "startdate", "day", "when", "datefrom"},
    "end_date": {"enddate", "finishdate", "dateto", "lastdate", "until", "finishes"},
    "time": {"time", "times", "timings", "starttime", "timeslot", "start", "doorsopen"},
    "end_time": {"endtime", "finishtime", "finish", "end", "closes"},
    "venue": {"venue", "location", "where", "place", "venuename"},
    "address": {"address", "fulladdress", "streetaddress", "venueaddress", "postcode"},
    "category": {"category", "categories", "type", "cat", "eventtype"},
    "fee": {"fee", "fees", "cost", "price", "charge", "entry", "entryfee", "admission", "ticketprice"},
    "url": {"url", "link", "website", "web", "webpage", "moreinfo", "infolink"},
    "booking": {"booking", "bookinginfo", "bookinginformation", "howtobook", "tickets", "bookinglink", "bookingurl"},
    "description": {"description", "details", "info", "information", "about", "summary", "notes", "eventdescription"},
    "age": {"age", "ages", "agerange", "agegroup", "agesuitability", "suitablefor", "audience", "whofor", "whoisitfor"},
    "accessibility": {"accessibility", "access", "accessnotes", "accessibilitynotes", "accessinfo"},
    "contact_email": {"email", "contactemail", "organiseremail", "emailaddress", "enquiries"},
    "contact_phone": {"phone", "telephone", "tel", "contactphone", "contactnumber", "phonenumber", "mobile"},
    "recurrence": {"recurrence", "recurring", "repeats", "repeat", "frequency", "howoften", "schedule"},
}


def _cell_to_str(cell) -> str:
    if cell is None:
        return ""
    if isinstance(cell, datetime):
        if (cell.hour, cell.minute) != (0, 0):
            return f"{cell.date().isoformat()} {cell.strftime('%H:%M')}"
        return cell.date().isoformat()
    if isinstance(cell, time_cls):
        return cell.strftime("%H:%M")
    if isinstance(cell, date_cls):
        return cell.isoformat()
    return str(cell).strip()


def _spreadsheet_rows(filename: str, data: bytes) -> list[list[str]]:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "csv":
        sample = data.decode("utf-8", errors="ignore")
        return [[(c or "").strip() for c in row] for row in csv.reader(sample.splitlines())]
    from openpyxl import load_workbook

    wb = load_workbook(BytesIO(data), read_only=True, data_only=True)
    rows: list[list[str]] = []
    for sheet in wb.worksheets:
        for row in sheet.iter_rows(values_only=True):
            rows.append([_cell_to_str(c) for c in row])
    return rows


def _detect_header_map(rows: list[list[str]]) -> tuple[Optional[int], Optional[dict]]:
    for idx, row in enumerate(rows[:10]):
        mapping: dict[str, int] = {}
        for col, cell in enumerate(row):
            key = re.sub(r"[^a-z]", "", (cell or "").lower())
            if not key:
                continue
            for field, synonyms in SPREADSHEET_HEADERS.items():
                if key in synonyms and field not in mapping:
                    mapping[field] = col
        if "title" in mapping and ("date" in mapping or "organisation" in mapping) and len(mapping) >= 3:
            return idx, mapping
    return None, None


def _try_structured_spreadsheet(filename: str, data: bytes, orgs: list[dict], events: list[dict]) -> Optional[ParsedDocument]:
    """Deterministic row-per-event parse for spreadsheets with recognised column
    headers (Organisation, Title, Date, Time/s, Venue, Category, Fee, URL,
    Booking info). Returns None when no header row is found so the caller can
    fall back to the AI text pipeline."""
    try:
        rows = _spreadsheet_rows(filename, data)
    except Exception:
        return None
    header_idx, mapping = _detect_header_map(rows)
    if mapping is None:
        return None

    items: list[ParsedItem] = []
    skipped = 0
    for row in rows[header_idx + 1:]:
        def get(field: str) -> str:
            col = mapping.get(field)
            if col is None or col >= len(row):
                return ""
            return (row[col] or "").strip()

        title = get("title")
        if not title:
            if any((c or "").strip() for c in row):
                skipped += 1
            continue
        org_name = get("organisation")
        date_raw = get("date")
        end_date_raw = get("end_date")
        time_raw = get("time")
        end_time_raw = get("end_time")
        venue = get("venue")
        address = get("address")
        category_raw = get("category")
        fee = get("fee")
        url = get("url")
        booking = get("booking")
        description_raw = get("description")
        age = get("age")
        accessibility = get("accessibility")
        contact_email = get("contact_email")
        contact_phone = get("contact_phone")
        recurrence_raw = get("recurrence")

        category = "Community"
        if category_raw:
            cat_match, cat_score = _best_match(category_raw, [{"name": c} for c in SITE_CATEGORIES], "name")
            if cat_match and cat_score >= 0.6:
                category = cat_match["name"]
        if category == "Community":
            inferred = _infer_category(" ".join([title, description_raw, category_raw]))
            if inferred:
                category = inferred

        end_date = None
        if end_date_raw:
            end_date, _, _ = _extract_iso_date(end_date_raw)

        end_time = None
        if end_time_raw:
            end_time, _, _, _ = _extract_hhmm_times(end_time_raw)

        recurrence = _detect_recurrence(" ".join([recurrence_raw, date_raw, time_raw])) if (recurrence_raw or date_raw or time_raw) else None

        if description_raw:
            description = description_raw
        else:
            desc_parts = [p for p in [
                f"{title} at {venue}." if venue else f"{title}.",
                f"Organised by {org_name}." if org_name else "",
                f"Cost: {fee}." if fee else "",
                f"Suitable for: {age}." if age else "",
                f"Accessibility: {accessibility}." if accessibility else "",
                f"Booking: {booking}" if booking else "",
                f"More info: {url}" if url else "",
            ] if p]
            description = " ".join(desc_parts)

        item = ParsedItem(
            suggested_type="event",
            action="new_event",
            title=title,
            date=date_raw or None,
            end_date=end_date,
            start_time=time_raw or None,
            end_time=end_time,
            location=venue or None,
            category=category,
            description=description,
            social_caption=f"📣 {title} — happening in Blackrod. More on Blackrod Now.",
            notification_text=f"New on Blackrod Now: {title}",
            confidence=0.9,
            cost=fee or None,
            booking=booking or None,
            url=url or None,
            category_explicit=bool(category_raw),
            address=address or None,
            age=age or None,
            accessibility=accessibility or None,
            contact_email=contact_email or None,
            contact_phone=contact_phone or None,
            recurrence_freq=recurrence["freq"] if recurrence else None,
            recurrence_weekday=recurrence.get("weekday") if recurrence else None,
            recurrence_confidence=recurrence.get("confidence") if recurrence else None,
            recurrence_raw_text=recurrence.get("raw_text") if recurrence else None,
        )
        if org_name:
            org_match, org_score = _best_match(org_name, orgs, "name")
            if org_match and org_score >= 0.8:
                item = item.model_copy(update={
                    "matched_org_slug": org_match.get("slug"),
                    "matched_org_name": org_match.get("name"),
                    "entity_confidence": round(org_score, 2),
                })
        ev_match, ev_score = _best_match(title, events, "title")
        if ev_match and ev_score >= 0.85:
            item = item.model_copy(update={
                "action": "update_event",
                "matched_event_id": ev_match.get("id"),
                "matched_event_title": ev_match.get("title"),
            })
        row_text = " | ".join([c for c in row if c])
        items.append(_normalize_parsed_item(item, row_text))
        if len(items) >= 100:
            break

    if not items:
        return None
    warnings = [f"Structured spreadsheet detected — parsed {len(items)} event row{'s' if len(items) != 1 else ''} from columns"]
    if skipped:
        warnings.append(f"Skipped {skipped} row{'s' if skipped != 1 else ''} with no title")
    return ParsedDocument(
        filename=filename,
        source_type="spreadsheet",
        text_excerpt=_clean_text_excerpt(" | ".join(rows[header_idx])),
        warnings=warnings,
        items=items,
    )


async def _parse_text_to_response(text: str, hint: Optional[str] = None) -> ParseResponse:
    if not text or not text.strip():
        raise HTTPException(400, "Text required")
    if not EMERGENT_LLM_KEY:
        return ParseResponse(items=[_fallback_parse(text)], mocked=True)
    system = (
        "You are a friendly community editor for Blackrod Now (Blackrod, Bolton, UK). "
        "Read the pasted text and return a JSON object with a single key `items` — an ARRAY of one or more "
        "structured suggestions. If the paste describes multiple events (e.g. a newsletter or multi-item flyer), "
        "return one item per event. Each item must have keys: "
        "suggested_type ('event', 'volunteer', 'organisation', 'venue' or 'update'), "
        "title, "
        "date (YYYY-MM-DD or null), end_date (YYYY-MM-DD when the event spans multiple days, else null), "
        "start_time (HH:MM 24-hour or null), end_time (HH:MM 24-hour or null), "
        "location (venue/place name only — e.g. 'St Catherine\'s Hall', NOT an address — or null), "
        "address (full street address if present — e.g. '15 Church Road, Blackrod BL6 5AQ' — or null), "
        "category (pick exactly one from: Community, Family, Children & Young People, Sports & Fitness, "
        "Arts & Culture, Heritage, Faith, Health & Wellbeing, Education, Social, Fundraising, Markets & Fairs, Music & Entertainment), "
        "description (1-3 sentences), "
        "cost (ticket price or 'Free' — null if not mentioned), "
        "booking (booking URL, phone, or instructions — null if not present), "
        "age (who the event is suitable for, e.g. 'All ages', 'Adults', 'Under 16s' — null if not clear), "
        "accessibility (any accessibility notes — null if not mentioned), "
        "contact_email (organiser email if present — null otherwise), "
        "contact_phone (organiser phone if present — null otherwise), "
        "social_caption (2-3 emojis, 2-4 hashtags), notification_text (max 90 chars). "
        "Also include action, matched_org_slug, matched_org_name, matched_event_id, matched_event_title, matched_volunteer_id, matched_volunteer_title, matched_venue_id, matched_venue_name and confidence when you can identify the target. "
        "If possible include raw_date_text, raw_time_text, date_confidence, time_confidence and entity_confidence. "
        "For any recurring events (e.g. 'Mondays, 10am', 'every Tuesday', 'first Sunday of the month', 'weekly drop-in', 'monthly meet-up', 'annual switch-on'), "
        "also set recurrence_freq to one of 'none', 'daily', 'weekly', 'biweekly', 'monthly', 'monthly_weekday' or 'annually', and recurrence_weekday to the canonical weekday name when clear. "
        "Use action values new_event, update_event, new_volunteer, update_volunteer, new_organisation, update_organisation, new_venue, update_venue or unclear. "
        "Warm, modern, youth-friendly tone. Return ONLY the JSON object — no markdown, no prose."
    )
    if hint:
        system += (
            "\n\nContext for matching: "
            + hint
            + "\nIf the document is clearly an existing organisation update, event update, volunteering update or venue update, prefer update_organisation, update_event, update_volunteer or update_venue and fill in the matching slug/id."
        )
    try:
        if not await _ensure_llm_loaded():
            logger.warning("Classifier: LLM libraries unavailable (ready=%s err=%s key=%s) — using fallback", _llm_ready.is_set(), _llm_preload_error, bool(EMERGENT_LLM_KEY))
            return ParseResponse(items=[_fallback_parse(text)])
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"parse-{new_id()}",
            system_message=system,
        ).with_model("anthropic", "claude-sonnet-4-5-20250929")
        raw = await chat.send_message(UserMessage(text=f"Raw paste:\n\n{text}"))
        cleaned = raw.strip()
        if cleaned.startswith("```"):
            cleaned = re.sub(r"^```(?:json)?", "", cleaned).strip()
            cleaned = re.sub(r"```$", "", cleaned).strip()
        try:
            data = json.loads(cleaned)
        except json.JSONDecodeError:
            m = re.search(r"\{.*\}", cleaned, re.S)
            if not m:
                return ParseResponse(items=[_fallback_parse(text)])
            data = json.loads(m.group(0))
        items_raw = data.get("items") if isinstance(data, dict) and "items" in data else [data]
        items = []
        for it in items_raw:
            try:
                parsed_item = ParsedItem(**{**{"description": ""}, **it})
                items.append(_normalize_parsed_item(parsed_item, text))
            except Exception as item_exc:
                logger.warning("Classifier item rejected: %s | raw keys: %s", item_exc, list(it.keys()) if isinstance(it, dict) else type(it))
                continue
        if not items:
            items = [_fallback_parse(text)]
        return ParseResponse(items=items)
    except Exception as e:
        logger.exception("AI parse failed: %s", e)
        return ParseResponse(items=[_fallback_parse(text)])


def _decode_bulk_sources(raw: Optional[str]) -> list[str]:
    if not raw:
        return []
    try:
        parsed = json.loads(raw)
        if isinstance(parsed, str):
            parsed = [parsed]
        if isinstance(parsed, list):
            return [str(item).strip() for item in parsed if str(item).strip()]
    except Exception:
        pass
    return [part.strip() for part in re.split(r"\n{2,}|\r\n{2,}|\n", raw) if part.strip()]


# ─────────── Labeled-block documents (Event title: … / Venue: …) ───────────
LABELED_FIELD_MAP = {
    "eventtitle": "title", "title": "title", "eventname": "title",
    "organisation": "org", "organization": "org", "org": "org", "organiser": "org", "organizer": "org",
    "category": "category",
    "date": "date", "dates": "date", "eventdate": "date",
    "start": "start_time", "starttime": "start_time", "time": "start_time", "times": "start_time",
    "end": "end_time", "endtime": "end_time",
    "venue": "venue", "location": "venue", "where": "venue", "place": "venue",
    "address": "address",
    "description": "description", "details": "description", "about": "description",
    "cost": "cost", "fee": "cost", "fees": "cost", "price": "cost", "admission": "cost",
    "agesuitability": "age", "age": "age", "ages": "age", "agerange": "age",
    "accessibility": "accessibility", "access": "accessibility",
    "bookinglink": "booking", "booking": "booking", "bookinginfo": "booking", "howtobook": "booking", "tickets": "booking",
    "email": "email", "contactemail": "email",
    "phone": "phone", "contactphone": "phone", "tel": "phone", "telephone": "phone",
    "url": "url", "link": "url", "website": "url", "web": "url", "moreinfo": "url",
    "imageurl": "image", "image": "image", "poster": "image",
    "repeats": "repeats", "recurrence": "repeats", "repeat": "repeats", "frequency": "repeats",
}

_LABEL_LINE_RE = re.compile(r"^\s*(?:\d+\.\s*)?([A-Za-z][A-Za-z /&]{1,24})\s*:\s*(.*)$")
_PLACEHOLDER_VALUE_RE = re.compile(
    r"^(not published|not stated|not specified|not safely verified|not applicable|none stated|none|n/?a|tbc|tba|unknown|no advance booking.*|no booking.*|-+)$",
    re.I,
)


def _labeled_value(raw: str) -> str:
    value = (raw or "").strip()
    return "" if not value or _PLACEHOLDER_VALUE_RE.match(value) else value


def _try_labeled_blocks(filename: str, text: str, orgs: list[dict], events: list[dict]) -> Optional[ParsedDocument]:
    """Deterministic parse for documents where each event is a block of
    'Label: value' lines (e.g. 'Event title: … / Organisation: … / Date: …').
    Returns None unless at least 2 titled blocks are found, so free-text
    flyers still go through the AI pipeline."""
    if not text:
        return None
    title_marker = re.compile(r"^\s*(?:\d+\.\s*)?event\s*title\s*:", re.I | re.M)
    markers = list(title_marker.finditer(text))
    if len(markers) < 2:
        return None

    items: list[ParsedItem] = []
    for i, marker in enumerate(markers):
        block_start = marker.start()
        block_end = markers[i + 1].start() if i + 1 < len(markers) else len(text)
        block = text[block_start:block_end]
        fields: dict[str, str] = {}
        for line in block.splitlines():
            m = _LABEL_LINE_RE.match(line)
            if not m:
                continue
            key = re.sub(r"[^a-z]", "", m.group(1).lower())
            field = LABELED_FIELD_MAP.get(key)
            if field and field not in fields:
                fields[field] = _labeled_value(m.group(2))
        title = fields.get("title")
        if not title:
            continue

        # Contact-ish fields: treat "No …"/"Not …" prose as empty; sanity-check shape.
        email_val = fields.get("email") or ""
        if email_val and ("@" not in email_val or re.match(r"^(no|not)\b", email_val, re.I)):
            email_val = ""
        phone_val = fields.get("phone") or ""
        if phone_val and (not re.search(r"\d{5,}", phone_val.replace(" ", "")) or re.match(r"^(no|not)\b", phone_val, re.I)):
            phone_val = ""
        booking_val = fields.get("booking") or ""
        if re.match(r"^(no|not)\b", booking_val, re.I):
            booking_val = ""

        category = "Community"
        if fields.get("category"):
            cat_match, cat_score = _best_match(fields["category"], [{"name": c} for c in SITE_CATEGORIES], "name")
            if cat_match and cat_score >= 0.55:
                category = cat_match["name"]

        recurrence_freq = None
        recurrence_weekday = None
        repeats_raw = fields.get("repeats") or ""
        if repeats_raw and not re.match(r"^no\b", repeats_raw, re.I):
            repeats_normalized = re.sub(r"\beach month\b", "the month", repeats_raw, flags=re.I)
            detected = _detect_recurrence(repeats_normalized) or _detect_recurrence(f"every {repeats_normalized}")
            if detected:
                recurrence_freq = detected.get("freq")
                recurrence_weekday = detected.get("weekday")
            elif re.match(r"^yes\b", repeats_raw, re.I):
                recurrence_freq = "weekly"

        url_value = fields.get("url") or ""
        image_value = fields.get("image") or ""
        item = ParsedItem(
            suggested_type="event",
            action="new_event",
            title=title,
            date=fields.get("date") or None,
            start_time=fields.get("start_time") or None,
            end_time=fields.get("end_time") or None,
            location=fields.get("venue") or None,
            address=fields.get("address") or None,
            category=category,
            description=fields.get("description") or title,
            cost=fields.get("cost") or None,
            age=fields.get("age") or None,
            accessibility=fields.get("accessibility") or None,
            booking=booking_val or None,
            contact_email=email_val or None,
            contact_phone=phone_val or None,
            category_explicit=bool(fields.get("category")),
            url=url_value if url_value.startswith("http") else None,
            image=image_value if image_value.startswith("http") else None,
            recurrence_freq=recurrence_freq,
            recurrence_weekday=recurrence_weekday,
            recurrence_raw_text=repeats_raw or None,
            social_caption=f"📣 {title} — happening in Blackrod. More on Blackrod Now.",
            notification_text=f"New on Blackrod Now: {title}",
            confidence=0.92,
        )
        if fields.get("org"):
            org_match, org_score = _best_match(fields["org"], orgs, "name")
            if org_match and org_score >= 0.8:
                item = item.model_copy(update={
                    "matched_org_slug": org_match.get("slug"),
                    "matched_org_name": org_match.get("name"),
                    "entity_confidence": round(org_score, 2),
                })
        ev_match, ev_score = _best_match(title, events, "title")
        if ev_match and ev_score >= 0.85:
            item = item.model_copy(update={
                "action": "update_event",
                "matched_event_id": ev_match.get("id"),
                "matched_event_title": ev_match.get("title"),
            })
        # Strip placeholder label lines so normalizer back-fill can't re-extract them.
        clean_lines = []
        for line in block.splitlines():
            lm = _LABEL_LINE_RE.match(line)
            if lm and lm.group(2).strip() and not _labeled_value(lm.group(2)):
                continue
            clean_lines.append(line)
        normalized = _normalize_parsed_item(item, "\n".join(clean_lines))
        if normalized.date and not re.fullmatch(r"\d{4}-\d{2}-\d{2}", normalized.date):
            iso = re.search(r"\d{4}-\d{2}-\d{2}", normalized.date)
            if iso:
                normalized = normalized.model_copy(update={"date": iso.group(0)})
        items.append(normalized)
        if len(items) >= 100:
            break

    if len(items) < 2:
        return None
    return ParsedDocument(
        filename=filename,
        source_type="labeled_document",
        text_excerpt=_clean_text_excerpt(text),
        warnings=[f"Structured document detected — parsed {len(items)} labeled event blocks without AI"],
        items=items,
    )


async def _admin_parse_documents(
    file_sources: Optional[list[tuple]] = None,
    source_org_slug: Optional[str] = None,
    urls_json: Optional[str] = None,
    texts_json: Optional[str] = None,
    progress_cb=None,
    total_timeout: Optional[int] = None,
    per_source_timeout: Optional[int] = None,
    precomputed_docs: Optional[list] = None,
    doc_cb=None,
    use_llm: bool = True,
) -> BulkParseResponse:
    file_list = list(file_sources or [])
    extraction_budget = per_source_timeout or BULK_PARSE_EXTRACTION_TIMEOUT_SECONDS
    classification_budget = max(BULK_PARSE_CLASSIFICATION_TIMEOUT_SECONDS, per_source_timeout or 0)

    url_list = _decode_bulk_sources(urls_json)
    text_list = _decode_bulk_sources(texts_json)
    if len(file_list) + len(url_list) + len(text_list) > MAX_BULK_PARSE_FILES:
        raise HTTPException(400, f"Upload at most {MAX_BULK_PARSE_FILES} total sources per bulk parse request")

    orgs = await db.orgs.find({}, {"_id": 0, "slug": 1, "name": 1}).sort("name", 1).to_list(200)
    events = await db.events.find({}, {"_id": 0, "id": 1, "title": 1}).sort("start", 1).to_list(200)
    volunteers = await db.volunteers.find({}, {"_id": 0, "id": 1, "title": 1, "orgSlug": 1}).sort("title", 1).to_list(300)
    venues = await db.venues.find({}, {"_id": 0, "id": 1, "name": 1, "address": 1}).sort("name", 1).to_list(300)
    source_org = None
    if source_org_slug:
        source_org = await _find_org(source_org_slug)
    documents: list[ParsedDocument] = list(precomputed_docs or [])
    skip_remaining = len(documents)  # checkpointed sources already done in a previous attempt
    flushed = len(documents)
    deadline = asyncio.get_running_loop().time() + (total_timeout or BULK_PARSE_TOTAL_TIMEOUT_SECONDS)

    async def _flush_new():
        # Checkpoint newly completed documents so retries resume, not restart.
        nonlocal flushed
        while flushed < len(documents):
            if doc_cb:
                try:
                    await doc_cb(flushed, documents[flushed])
                except Exception:
                    pass
            flushed += 1

    async def _notify(current: str = ""):
        if progress_cb:
            try:
                await progress_cb(len(documents), current)
            except Exception:
                pass

    def _bulk_timeout_document(filename: str, source_type: str, warnings: list[str]) -> ParsedDocument:
        return ParsedDocument(
            filename=filename,
            source_type=source_type,
            text_excerpt="",
            warnings=warnings,
            items=[_fallback_parse(filename)],
        )

    async def _finalize_source(filename: str, source_type: str, text: str, warnings: list[str]) -> ParsedDocument:
        if not text.strip():
            return ParsedDocument(
                filename=filename,
                source_type=source_type,
                text_excerpt="",
                warnings=warnings + ["No readable text extracted"],
                items=[_fallback_parse(filename)],
            )

        # Deterministic path for label-formatted documents (Event title: … blocks)
        labeled = _try_labeled_blocks(filename, text, orgs, events)
        if labeled:
            return labeled.model_copy(update={"warnings": warnings + labeled.warnings})

        hint = "\n".join([
            f"Filename: {filename}",
            f"Source organisation: {source_org.get('slug')} | {source_org.get('name')}" if source_org else "",
            "Existing organisations:",
            *[f"ORG:{o['slug']}|{o['name']}" for o in orgs[:80]],
            "Existing events:",
            *[f"EVENT:{e['id']}|{e['title']}" for e in events[:120]],
            "Existing volunteering opportunities:",
            *[f"VOL:{v['id']}|{v['title']}|{v.get('orgSlug','')}" for v in volunteers[:160]],
            "Existing venues:",
            *[f"VENUE:{v['id']}|{v['name']}|{v.get('address','')}" for v in venues[:160]],
        ]).strip()
        classification_timeout = max(1, min(classification_budget, int(max(1, deadline - asyncio.get_running_loop().time() - BULK_PARSE_TIMEOUT_SAFETY_SECONDS))))
        if not use_llm:
            return ParsedDocument(
                filename=filename,
                source_type=source_type,
                text_excerpt=_clean_text_excerpt(text),
                warnings=warnings + ["Parsed without AI (recovery mode) — review suggested fields carefully"],
                items=_fallback_document_classify(filename, text, orgs, events, volunteers, venues),
            )
        try:
            parsed = await asyncio.wait_for(_parse_text_to_response(text, hint=hint), timeout=classification_timeout)
        except asyncio.TimeoutError:
            return ParsedDocument(
                filename=filename,
                source_type=source_type,
                text_excerpt=_clean_text_excerpt(text),
                warnings=warnings + [f"Text classification timed out after {classification_timeout}s"],
                items=_fallback_document_classify(filename, text, orgs, events, volunteers, venues),
            )
        except Exception as exc:
            return ParsedDocument(
                filename=filename,
                source_type=source_type,
                text_excerpt=_clean_text_excerpt(text),
                warnings=warnings + [f"Text classification failed: {exc}"],
                items=[_fallback_parse(text)],
            )

        items = parsed.items or [_fallback_parse(text)]
        if parsed.mocked:
            items = _fallback_document_classify(filename, text, orgs, events, volunteers, venues)
        return ParsedDocument(
            filename=filename,
            source_type=source_type,
            text_excerpt=_clean_text_excerpt(text),
            warnings=warnings,
            items=items,
        )

    for index, (filename, content_type, data) in enumerate(file_list):
        await _flush_new()
        if skip_remaining > 0:
            skip_remaining -= 1
            continue
        await _notify(filename)
        remaining = deadline - asyncio.get_running_loop().time()
        if remaining <= BULK_PARSE_TIMEOUT_SAFETY_SECONDS:
            timeout_warning = [f"Bulk parse timed out after {total_timeout or BULK_PARSE_TOTAL_TIMEOUT_SECONDS}s"]
            for remaining_name, _, _ in file_list[index:]:
                documents.append(_bulk_timeout_document(remaining_name or "file", "timed_out", timeout_warning))
            break

        if not data:
            documents.append(ParsedDocument(filename=filename, source_type="empty", text_excerpt="", warnings=["Empty file"], items=[_fallback_parse("")]))
            continue
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if ext in {"xlsx", "csv"}:
            structured = await asyncio.to_thread(_try_structured_spreadsheet, filename, data, orgs, events)
            if structured:
                documents.append(structured)
                continue
        extraction_timeout = max(1, min(extraction_budget, int(max(1, remaining - BULK_PARSE_TIMEOUT_SAFETY_SECONDS))))
        try:
            text, source_type, warnings = await asyncio.wait_for(
                asyncio.to_thread(_extract_document_text, filename, content_type or "application/octet-stream", data),
                timeout=extraction_timeout,
            )
        except asyncio.TimeoutError:
            documents.append(_bulk_timeout_document(filename, "timed_out", [f"Document extraction timed out after {extraction_timeout}s"]))
            continue
        except Exception as exc:
            documents.append(_bulk_timeout_document(filename, "timed_out", [f"Document extraction failed: {exc}"]))
            continue

        documents.append(await _finalize_source(filename, source_type, text, warnings))

    for index, url in enumerate(url_list):
        await _flush_new()
        if skip_remaining > 0:
            skip_remaining -= 1
            continue
        remaining = deadline - asyncio.get_running_loop().time()
        filename = urlparse(url).netloc or "link"
        await _notify(filename)
        if remaining <= BULK_PARSE_TIMEOUT_SAFETY_SECONDS:
            timeout_warning = [f"Bulk parse timed out after {total_timeout or BULK_PARSE_TOTAL_TIMEOUT_SECONDS}s"]
            for remaining_url in url_list[index:]:
                remaining_name = urlparse(remaining_url).netloc or "link"
                documents.append(_bulk_timeout_document(remaining_name, "timed_out", timeout_warning))
            break

        extraction_timeout = max(1, min(extraction_budget, int(max(1, remaining - BULK_PARSE_TIMEOUT_SAFETY_SECONDS))))
        try:
            text, source_type, warnings = await asyncio.wait_for(
                asyncio.to_thread(_extract_url_text, url),
                timeout=extraction_timeout,
            )
        except asyncio.TimeoutError:
            documents.append(_bulk_timeout_document(filename, "url", [f"URL fetch timed out after {extraction_timeout}s"]))
            continue
        except Exception as exc:
            documents.append(_bulk_timeout_document(filename, "url", [f"URL fetch failed: {exc}"]))
            continue

        documents.append(await _finalize_source(filename, source_type or "url", text, warnings))

    for index, text_source in enumerate(text_list):
        await _flush_new()
        if skip_remaining > 0:
            skip_remaining -= 1
            continue
        remaining = deadline - asyncio.get_running_loop().time()
        filename = f"pasted-text-{index + 1}.txt"
        await _notify(filename)
        if remaining <= BULK_PARSE_TIMEOUT_SAFETY_SECONDS:
            timeout_warning = [f"Bulk parse timed out after {total_timeout or BULK_PARSE_TOTAL_TIMEOUT_SECONDS}s"]
            for remaining_offset, _ in enumerate(text_list[index:], start=index):
                remaining_name = f"pasted-text-{remaining_offset + 1}.txt"
                documents.append(_bulk_timeout_document(remaining_name, "timed_out", timeout_warning))
            break

        documents.append(await _finalize_source(filename, "text", text_source, []))
    await _flush_new()
    await _notify("")
    return BulkParseResponse(documents=documents, mocked=not EMERGENT_LLM_KEY)


async def _read_upload_sources(files: List[UploadFile]) -> list[tuple]:
    file_sources: list[tuple] = []
    for file in files or []:
        filename = file.filename or "file"
        data = await file.read()
        if len(data) > MAX_UPLOAD_BYTES:
            raise HTTPException(413, f"File too large (max 10 MB): {filename}")
        file_sources.append((filename, file.content_type or "application/octet-stream", data))
    return file_sources


PARSE_JOB_HEARTBEAT_SECONDS = int(os.environ.get("PARSE_JOB_HEARTBEAT_SECONDS", "10"))
PARSE_JOB_STALE_SECONDS = int(os.environ.get("PARSE_JOB_STALE_SECONDS", "300"))
PARSE_JOB_WORKER_POLL_SECONDS = float(os.environ.get("PARSE_JOB_WORKER_POLL_SECONDS", "1.5"))
PARSE_JOB_WORKER_CONCURRENCY = max(1, int(os.environ.get("PARSE_JOB_WORKER_CONCURRENCY", "1")))
PARSE_JOB_MAX_ATTEMPTS = max(1, int(os.environ.get("PARSE_JOB_MAX_ATTEMPTS", "5")))
PARSE_JOB_TOTAL_TIMEOUT_SECONDS = int(os.environ.get("PARSE_JOB_TOTAL_TIMEOUT_SECONDS", "3600"))
PARSE_JOB_SOURCE_TIMEOUT_SECONDS = int(os.environ.get("PARSE_JOB_SOURCE_TIMEOUT_SECONDS", "180"))
PARSE_JOB_SOURCE_RETENTION_DAYS = max(1, int(os.environ.get("PARSE_JOB_SOURCE_RETENTION_DAYS", "7")))
PARSE_JOB_INTERRUPTED_ERROR = "Parser worker was interrupted; the job has been safely re-queued."

# Durable parser workers. The HTTP request only persists the source material and
# returns a job id. MongoDB is the queue, so a browser refresh or server restart
# does not lose the import.
_parse_worker_tasks: List[asyncio.Task] = []
_parse_worker_shutdown = asyncio.Event()


def _parse_job_is_stale(job: Dict[str, Any]) -> bool:
    if job.get("status") != "processing":
        return False
    try:
        updated_raw = job.get("updated_at") or job.get("started_at") or job.get("created_at")
        updated = datetime.fromisoformat(str(updated_raw).replace("Z", "+00:00"))
        if updated.tzinfo is None:
            updated = updated.replace(tzinfo=timezone.utc)
        return (datetime.now(timezone.utc) - updated).total_seconds() > PARSE_JOB_STALE_SECONDS
    except Exception:
        return False


async def _persist_parse_job_sources(
    job_id: str,
    file_sources: list[tuple],
    url_list: List[str],
    text_list: List[str],
) -> None:
    """Persist each source as its own Mongo document.

    Keeping files separate avoids MongoDB's 16 MB per-document limit while
    still allowing the worker to reconstruct the exact request after a restart.
    Sources expire automatically after a short retention period once the TTL
    index is present.
    """
    expires_at = datetime.now(timezone.utc) + timedelta(days=PARSE_JOB_SOURCE_RETENTION_DAYS)
    rows: List[Dict[str, Any]] = []
    order = 0

    for filename, content_type, data in file_sources:
        rows.append({
            "id": new_id(),
            "job_id": job_id,
            "order": order,
            "kind": "file",
            "filename": filename or "file",
            "content_type": content_type or "application/octet-stream",
            "data": data,
            "size": len(data or b""),
            "created_at": now_iso(),
            "expires_at": expires_at,
        })
        order += 1

    for url in url_list:
        rows.append({
            "id": new_id(),
            "job_id": job_id,
            "order": order,
            "kind": "url",
            "value": url,
            "created_at": now_iso(),
            "expires_at": expires_at,
        })
        order += 1

    for source_text in text_list:
        rows.append({
            "id": new_id(),
            "job_id": job_id,
            "order": order,
            "kind": "text",
            "value": source_text,
            "created_at": now_iso(),
            "expires_at": expires_at,
        })
        order += 1

    if rows:
        await db.parse_job_sources.insert_many(rows)


async def _load_parse_job_sources(job_id: str) -> tuple[list[tuple], List[str], List[str]]:
    rows = await db.parse_job_sources.find(
        {"job_id": job_id},
        {"_id": 0},
    ).sort("order", 1).to_list(MAX_BULK_PARSE_FILES + 20)

    file_sources: list[tuple] = []
    url_list: List[str] = []
    text_list: List[str] = []

    for row in rows:
        kind = row.get("kind")
        if kind == "file":
            raw = row.get("data") or b""
            try:
                data = bytes(raw)
            except Exception:
                data = raw
            file_sources.append((
                row.get("filename") or "file",
                row.get("content_type") or "application/octet-stream",
                data,
            ))
        elif kind == "url":
            value = str(row.get("value") or "").strip()
            if value:
                url_list.append(value)
        elif kind == "text":
            value = str(row.get("value") or "").strip()
            if value:
                text_list.append(value)

    return file_sources, url_list, text_list


async def _recover_stale_parse_jobs() -> int:
    """Re-queue jobs whose worker disappeared.

    This is deliberately recovery, not failure. A deploy/restart must not force
    the administrator to upload the same Word document or pasted text again.
    """
    recovered = 0
    rows = await db.parse_jobs.find({"status": "processing"}, {"_id": 0}).to_list(500)
    for job in rows:
        if not _parse_job_is_stale(job):
            continue
        attempts = int(job.get("attempts") or 0)
        if attempts >= PARSE_JOB_MAX_ATTEMPTS:
            result = await db.parse_jobs.update_one(
                {
                    "id": job.get("id"),
                    "status": "processing",
                    "updated_at": job.get("updated_at"),
                },
                {"$set": {
                    "status": "failed",
                    "current": "",
                    "error": f"Parser stopped responding after {attempts} attempts.",
                    "updated_at": now_iso(),
                }},
            )
            recovered += int(result.modified_count or 0)
            continue

        result = await db.parse_jobs.update_one(
            {
                "id": job.get("id"),
                "status": "processing",
                "updated_at": job.get("updated_at"),
            },
            {"$set": {
                "status": "queued",
                "current": "Recovering interrupted import",
                "error": PARSE_JOB_INTERRUPTED_ERROR,
                "worker_id": None,
                "updated_at": now_iso(),
            }, "$inc": {"recoveries": 1}},
        )
        recovered += int(result.modified_count or 0)
    return recovered


async def _claim_parse_job(worker_id: str) -> Optional[Dict[str, Any]]:
    return await db.parse_jobs.find_one_and_update(
        {
            "status": "queued",
            "$or": [
                {"attempts": {"$lt": PARSE_JOB_MAX_ATTEMPTS}},
                {"attempts": {"$exists": False}},
            ],
        },
        {
            "$set": {
                "status": "processing",
                "worker_id": worker_id,
                "started_at": now_iso(),
                "updated_at": now_iso(),
                "current": "Preparing sources",
                "error": None,
            },
            "$inc": {"attempts": 1},
        },
        sort=[("created_at", 1)],
        return_document=ReturnDocument.AFTER,
    )


async def _run_claimed_parse_job(job: Dict[str, Any], worker_id: str) -> None:
    job_id = str(job.get("id") or "")
    if not job_id:
        return

    total = int(job.get("total") or 0)
    source_org_slug = job.get("source_org_slug") or None

    async def progress(done: int, current: str):
        await db.parse_jobs.update_one(
            {"id": job_id, "status": "processing", "worker_id": worker_id},
            {"$set": {
                "done": int(done or 0),
                "current": current or "",
                "updated_at": now_iso(),
            }},
        )

    stop_heartbeat = asyncio.Event()

    async def _heartbeat():
        while not stop_heartbeat.is_set():
            try:
                await asyncio.wait_for(stop_heartbeat.wait(), timeout=PARSE_JOB_HEARTBEAT_SECONDS)
            except asyncio.TimeoutError:
                try:
                    await db.parse_jobs.update_one(
                        {"id": job_id, "status": "processing", "worker_id": worker_id},
                        {"$set": {"updated_at": now_iso()}},
                    )
                except Exception:
                    logger.exception("Parse-job heartbeat failed for %s", job_id)

    heartbeat_task = asyncio.create_task(_heartbeat())
    try:
        file_sources, url_list, text_list = await _load_parse_job_sources(job_id)
        loaded_total = len(file_sources) + len(url_list) + len(text_list)
        if loaded_total == 0:
            raise RuntimeError("The parser job has no stored source material")

        if total <= 0:
            total = loaded_total
            await db.parse_jobs.update_one({"id": job_id}, {"$set": {"total": total}})

        # Resume from checkpointed per-source results left by a previous attempt
        # (crash/restart mid-job) instead of redoing completed sources.
        fresh_job = await db.parse_jobs.find_one({"id": job_id}, {"_id": 0, "partial_docs": 1, "attempts": 1}) or {}
        precomputed: list = []
        for raw_doc in (fresh_job.get("partial_docs") or [])[:loaded_total]:
            try:
                precomputed.append(ParsedDocument.model_validate(raw_doc))
            except Exception:
                precomputed = []
                break
        attempts = int(fresh_job.get("attempts") or job.get("attempts") or 1)
        if not precomputed and fresh_job.get("partial_docs"):
            await db.parse_jobs.update_one({"id": job_id}, {"$unset": {"partial_docs": ""}})
        # After repeated failures (likely infra/LLM trouble), finish the job in
        # recovery mode without the LLM so admins always get reviewable drafts.
        use_llm = attempts < max(2, PARSE_JOB_MAX_ATTEMPTS - 1)

        async def checkpoint_doc(index: int, doc: ParsedDocument):
            await db.parse_jobs.update_one(
                {"id": job_id},
                {"$push": {"partial_docs": doc.model_dump()}, "$set": {"updated_at": now_iso()}},
            )

        # Durable jobs are allowed enough time to finish properly. The per-source
        # guard still protects us from genuinely wedged extractors/LLM calls.
        total_timeout = max(PARSE_JOB_TOTAL_TIMEOUT_SECONDS, 240 * max(1, loaded_total))
        result = await _admin_parse_documents(
            file_sources,
            source_org_slug=source_org_slug,
            urls_json=json.dumps(url_list, ensure_ascii=False),
            texts_json=json.dumps(text_list, ensure_ascii=False),
            progress_cb=progress,
            total_timeout=total_timeout,
            per_source_timeout=PARSE_JOB_SOURCE_TIMEOUT_SECONDS,
            precomputed_docs=precomputed,
            doc_cb=checkpoint_doc,
            use_llm=use_llm,
        )

        await db.parse_jobs.update_one(
            {"id": job_id, "status": "processing", "worker_id": worker_id},
            {"$set": {
                "status": "done",
                "done": total,
                "current": "",
                "result": result.model_dump(),
                "error": None,
                "completed_at": now_iso(),
                "updated_at": now_iso(),
            }, "$unset": {"partial_docs": ""}},
        )
    except asyncio.CancelledError:
        # Graceful deploy/restart: put the job straight back on the durable
        # queue so another worker can resume it immediately. A hard process
        # crash is handled separately by stale-heartbeat recovery.
        try:
            await db.parse_jobs.update_one(
                {"id": job_id, "status": "processing", "worker_id": worker_id},
                {"$set": {
                    "status": "queued",
                    "current": "Server restarted — resuming import",
                    "error": PARSE_JOB_INTERRUPTED_ERROR,
                    "worker_id": None,
                    "updated_at": now_iso(),
                }, "$inc": {"recoveries": 1}},
            )
        except Exception:
            pass
        raise
    except Exception as exc:
        logger.exception("Parse job %s failed on worker %s: %s", job_id, worker_id, exc)
        fresh = await db.parse_jobs.find_one({"id": job_id}, {"_id": 0, "attempts": 1}) or {}
        attempts = int(fresh.get("attempts") or job.get("attempts") or 1)
        if attempts < PARSE_JOB_MAX_ATTEMPTS:
            await db.parse_jobs.update_one(
                {"id": job_id, "status": "processing", "worker_id": worker_id},
                {"$set": {
                    "status": "queued",
                    "current": "Retrying import",
                    "error": str(exc),
                    "worker_id": None,
                    "updated_at": now_iso(),
                }},
            )
        else:
            await db.parse_jobs.update_one(
                {"id": job_id},
                {"$set": {
                    "status": "failed",
                    "current": "",
                    "error": str(exc),
                    "failed_at": now_iso(),
                    "updated_at": now_iso(),
                }},
            )
    finally:
        stop_heartbeat.set()
        heartbeat_task.cancel()
        try:
            await heartbeat_task
        except (asyncio.CancelledError, Exception):
            pass


async def _parse_job_worker(worker_number: int) -> None:
    worker_id = f"{os.getpid()}:{worker_number}:{uuid.uuid4().hex[:8]}"
    logger.info("Parse worker started: %s", worker_id)
    recovery_counter = 0

    while not _parse_worker_shutdown.is_set():
        try:
            # Periodically recover work abandoned by a killed/restarted process.
            if recovery_counter <= 0:
                recovered = await _recover_stale_parse_jobs()
                if recovered:
                    logger.warning("Recovered %s stale parse job(s)", recovered)
                recovery_counter = 20
            recovery_counter -= 1

            job = await _claim_parse_job(worker_id)
            if job:
                await _run_claimed_parse_job(job, worker_id)
                continue

            try:
                await asyncio.wait_for(_parse_worker_shutdown.wait(), timeout=PARSE_JOB_WORKER_POLL_SECONDS)
            except asyncio.TimeoutError:
                pass
        except asyncio.CancelledError:
            break
        except Exception as exc:
            logger.exception("Parse worker %s loop error: %s", worker_id, exc)
            try:
                await asyncio.wait_for(_parse_worker_shutdown.wait(), timeout=2.0)
            except asyncio.TimeoutError:
                pass

    logger.info("Parse worker stopped: %s", worker_id)


@api.post("/parse-content", response_model=ParseResponse)
async def parse_content(req: ParseRequest):
    return await _parse_text_to_response(req.text, hint=req.hint)


class PublicListSubmission(BaseModel):
    submitter_name: str
    submitter_email: str
    org_name: str = ""
    notes: str = ""
    items: List[ParsedItem]


PUBLIC_LIST_FALLBACK_ORG = "blackrod-sports-community-centre"


@api.post("/public/event-list/parse")
async def public_event_list_parse(file: UploadFile = File(...)):
    """Deterministic-only parse for the public 'Submit your events list' page.
    Never uses the AI: only the strict template formats are accepted, so
    submitters (and admins) always see exactly what was written."""
    filename = file.filename or "file"
    data = await file.read()
    if not data:
        raise HTTPException(400, "The file is empty")
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(413, "File too large (max 10 MB)")
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    orgs = await db.orgs.find({}, {"_id": 0, "slug": 1, "name": 1}).to_list(500)
    events = await db.events.find({}, {"_id": 0, "id": 1, "title": 1}).to_list(2000)

    doc: Optional[ParsedDocument] = None
    if ext in {"xlsx", "csv"}:
        doc = await asyncio.to_thread(_try_structured_spreadsheet, filename, data, orgs, events)
    elif ext in {"docx", "txt"}:
        if ext == "docx":
            text, _warn = await asyncio.to_thread(_extract_docx_text, data)
        else:
            text = data.decode("utf-8", errors="ignore")
        doc = _try_labeled_blocks(filename, text, orgs, events)
    else:
        raise HTTPException(422, "Please upload a filled-in Word (.docx), Excel (.xlsx) or CSV template")
    if not doc or not doc.items:
        raise HTTPException(
            422,
            "We couldn't find the template structure in this file. Please download the Word or "
            "spreadsheet template, fill it in without changing the labels/columns, and upload again.",
        )
    return {"format": doc.source_type, "count": len(doc.items), "warnings": doc.warnings, "items": [it.model_dump() for it in doc.items]}


@api.post("/public/event-list/submit")
async def public_event_list_submit(req: PublicListSubmission):
    name = req.submitter_name.strip()
    email = req.submitter_email.strip()
    if not name or "@" not in email:
        raise HTTPException(400, "Please provide your name and a valid email address")
    if not req.items:
        raise HTTPException(400, "No events to submit")
    if len(req.items) > 100:
        raise HTTPException(400, "Maximum 100 events per submission")

    orgs = await db.orgs.find({}, {"_id": 0, "slug": 1, "name": 1}).to_list(500)
    created_ids: list[str] = []
    submitted_titles: list[str] = []
    skipped: list[str] = []
    for item in req.items:
        if not item.title or not item.date or not re.fullmatch(r"\d{4}-\d{2}-\d{2}", item.date or ""):
            skipped.append(item.title or "(untitled)")
            continue
        org_slug = item.matched_org_slug
        if not org_slug and (item.matched_org_name or req.org_name):
            org_match, org_score = _best_match(item.matched_org_name or req.org_name, orgs, "name")
            if org_match and org_score >= 0.85:
                org_slug = org_match.get("slug")
        org_slug = org_slug or PUBLIC_LIST_FALLBACK_ORG
        start_time = item.start_time if re.fullmatch(r"\d{2}:\d{2}", item.start_time or "") else "10:00"
        end_time = item.end_time if re.fullmatch(r"\d{2}:\d{2}", item.end_time or "") else None
        end_date = item.end_date if re.fullmatch(r"\d{4}-\d{2}-\d{2}", item.end_date or "") else item.date
        end_value = end_time or start_time
        recurrence = None
        if item.recurrence_freq and item.recurrence_freq != "none":
            recurrence = EventRecurrence(freq=item.recurrence_freq)
        evt = Event(
            title=item.title.strip(),
            orgSlug=org_slug,
            category=item.category or "Community",
            start=f"{item.date}T{start_time}:00",
            end=f"{end_date}T{end_value}:00",
            venue=item.location or "",
            address=item.address or item.location or "",
            description=item.description or item.title,
            cost=item.cost or "",
            age=item.age or "",
            accessibility=item.accessibility or "",
            booking=item.booking or item.url or "",
            contactEmail=item.contact_email or "",
            contactPhone=item.contact_phone or "",
            image=item.image or "",
            featured=False,
            status="pending",  # NEVER auto-published — admin must approve each one
            recurrence=recurrence,
        )
        if _is_blank_or_legacy_event_image(evt.image):
            evt.image = _event_category_image(evt.category)
        await db.events.insert_one(evt.model_dump())
        created_ids.append(evt.id)
        submitted_titles.append(evt.title)

    if not created_ids:
        raise HTTPException(400, "No valid events found — every event needs at least a title and a date")

    submission = {
        "id": new_id(),
        "submitter_name": name,
        "submitter_email": email,
        "org_name": req.org_name.strip(),
        "notes": req.notes.strip()[:2000],
        "event_ids": created_ids,
        "skipped": skipped,
        "created_at": now_iso(),
    }
    await db.bulk_submissions.insert_one(submission)
    await _audit(
        action="public_event_list_submitted",
        entity_type="event",
        entity_id=submission["id"],
        summary=f"Public events list from {name} ({email}): {len(created_ids)} pending event(s)",
        meta={"event_ids": created_ids, "org_name": req.org_name, "skipped": skipped},
        actor="public",
    )

    # Best-effort notifications (never block or fail the submission)
    titles_html = "".join(f"<li>{html.escape(t)}</li>" for t in submitted_titles[:30])
    org_line = f" for <strong>{html.escape(req.org_name.strip())}</strong>" if req.org_name.strip() else ""
    admin_to = os.environ.get("ADMIN_NOTIFY_EMAIL") or SENDER_EMAIL
    admin_html = (
        f"<h2>New events list submitted</h2>"
        f"<p><strong>{html.escape(name)}</strong> ({html.escape(email)}) submitted "
        f"{len(created_ids)} event(s){org_line}.</p>"
        f"<ul>{titles_html}</ul>"
        + (f"<p>Notes: {html.escape(req.notes.strip())}</p>" if req.notes.strip() else "")
        + f"<p>All are <strong>pending</strong> — review and approve them in the admin dashboard: "
        f"<a href=\"{PUBLIC_URL.rstrip('/')}/admin\">{PUBLIC_URL.rstrip('/')}/admin</a></p>"
    )
    submitter_html = (
        f"<h2>We've received your events list</h2>"
        f"<p>Hi {html.escape(name)},</p>"
        f"<p>Thanks — we've received {len(created_ids)} event(s){org_line}. "
        f"The Blackrod Now team will check the details and publish them once approved. "
        f"Nothing appears on the site until it has been reviewed.</p>"
        f"<ul>{titles_html}</ul>"
        f"<p>If anything needs correcting, just reply to this email.</p>"
        f"<p>— Blackrod Now</p>"
    )

    def _notify():
        resend_send(admin_to, f"New events list: {len(created_ids)} event(s) from {name}", admin_html)
        resend_send(email, "Your events list is under review — Blackrod Now", submitter_html)

    asyncio.create_task(asyncio.to_thread(_notify))
    return {"ok": True, "created": len(created_ids), "skipped": skipped, "status": "pending_review"}


class AdminCheckReq(BaseModel):
    kind: Literal["event", "org"]
    id: str


CHECK_VERDICTS = {"looks_accurate", "needs_attention", "likely_outdated", "could_not_verify"}


@api.post("/admin/check")
async def admin_check_entity(req: AdminCheckReq):
    """Fact-check an event or organisation against the live web (Claude web search)."""
    if req.kind == "event":
        entity = await db.events.find_one({"id": req.id}, {"_id": 0})
        if not entity:
            raise HTTPException(404, "Event not found")
        org = await db.orgs.find_one({"slug": entity.get("orgSlug")}, {"_id": 0, "name": 1, "website": 1, "socials": 1}) or {}
        facts = {
            "event_title": entity.get("title"),
            "date_start": entity.get("start"),
            "venue": entity.get("venue"),
            "address": entity.get("address"),
            "cost": entity.get("cost"),
            "booking": entity.get("booking"),
            "recurrence": (entity.get("recurrence") or {}).get("freq"),
            "description": (entity.get("description") or "")[:400],
            "organisation": org.get("name"),
            "org_website": org.get("website"),
            "org_socials": org.get("socials"),
        }
        subject = f'the community event "{entity.get("title")}"'
    else:
        entity = await db.orgs.find_one({"slug": req.id}, {"_id": 0})
        if not entity:
            raise HTTPException(404, "Organisation not found")
        facts = {
            "name": entity.get("name"),
            "category": entity.get("category"),
            "description": (entity.get("description") or "")[:400],
            "website": entity.get("website"),
            "socials": entity.get("socials"),
            "email": entity.get("email"),
            "phone": entity.get("phone"),
            "meeting_info": entity.get("meeting_info") or entity.get("meets"),
        }
        subject = f'the community organisation "{entity.get("name")}"'

    if not await _ensure_llm_loaded():
        raise HTTPException(503, "The AI checker isn't available right now — try again shortly")
    from emergentintegrations.llm.chat import LlmChat, UserMessage

    system = (
        "You are a careful fact-checker for Blackrod Now, a community website for Blackrod, Bolton, UK. "
        "Use web search to verify whether the listing below is accurate and still active/running. "
        "Prioritise the organisation's own website/Facebook page, Bolton Council and local news. "
        "Be conservative: if you cannot find clear evidence either way, say so rather than guessing. "
        "Respond with ONLY a JSON object, no markdown: "
        '{"verdict": one of "looks_accurate"|"needs_attention"|"likely_outdated"|"could_not_verify", '
        '"summary": "2-3 plain-English sentences for the site admin", '
        '"issues": ["specific discrepancy or concern", ...] (empty list if none), '
        '"sources": [{"url": "...", "title": "..."}, ...]}'
    )
    chat = (
        LlmChat(api_key=EMERGENT_LLM_KEY, session_id=f"check-{new_id()}", system_message=system)
        .with_model("anthropic", "claude-sonnet-4-5-20250929")
        .with_tools([{"type": "web_search_20250305", "name": "web_search", "max_uses": 4}])
    )
    try:
        raw = await asyncio.wait_for(
            chat.send_message_with_tools(UserMessage(text=f"Verify {subject}. Current listing data:\n{json.dumps(facts, ensure_ascii=False)}")),
            timeout=150,
        )
        content = getattr(raw, "content", None) or str(raw)
        cleaned = re.sub(r"^```(?:json)?|```$", "", content.strip(), flags=re.M).strip()
        m = re.search(r"\{.*\}", cleaned, re.S)
        data = json.loads(m.group(0) if m else cleaned)
    except asyncio.TimeoutError:
        raise HTTPException(504, "The check took too long — please try again")
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("admin_check failed: %s", exc)
        raise HTTPException(502, "The checker couldn't complete — please try again")

    result = {
        "verdict": data.get("verdict") if data.get("verdict") in CHECK_VERDICTS else "could_not_verify",
        "summary": str(data.get("summary") or "")[:1200],
        "issues": [str(i)[:300] for i in (data.get("issues") or [])][:8],
        "sources": [
            {"url": str(s.get("url") or "")[:400], "title": str(s.get("title") or "")[:160]}
            for s in (data.get("sources") or []) if isinstance(s, dict) and s.get("url")
        ][:6],
        "checked_at": now_iso(),
    }
    if req.kind == "event":
        await db.events.update_one({"id": req.id}, {"$set": {"check_result": result}})
    else:
        await db.orgs.update_one({"slug": req.id}, {"$set": {"check_result": result}})
    await _audit(
        action="entity_checked",
        entity_type=req.kind,
        entity_id=req.id,
        summary=f"Web check ({result['verdict']}): {subject}",
        meta={"verdict": result["verdict"]},
        actor="admin",
    )
    return result


@api.get("/sitemap.xml")
async def sitemap_xml(request: Request):
    base = _abs_base_url(request)
    static_paths = ["", "/events", "/organisations", "/volunteering", "/venues", "/local-feed", "/submit-event", "/submit-events-list", "/contact", "/faq"]
    urls: list[str] = [f"<url><loc>{base}{p}</loc><changefreq>daily</changefreq></url>" for p in static_paths]
    horizon = (datetime.now(timezone.utc) - timedelta(days=7)).isoformat()
    events = await db.events.find(
        {"status": "approved"}, {"_id": 0, "id": 1, "start": 1, "recurrence": 1}
    ).to_list(2000)
    for e in events:
        rec = e.get("recurrence") or {}
        if (e.get("start") or "") < horizon and (rec.get("freq") or "none") == "none" and not rec.get("extra_dates"):
            continue
        urls.append(f"<url><loc>{base}/events/{e['id']}</loc><changefreq>weekly</changefreq></url>")
    orgs = await db.orgs.find({"status": {"$ne": "pending"}}, {"_id": 0, "slug": 1}).to_list(500)
    for o in orgs:
        urls.append(f"<url><loc>{base}/organisations/{o['slug']}</loc><changefreq>weekly</changefreq></url>")
    xml = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">' + "".join(urls) + "</urlset>"
    )
    return Response(content=xml, media_type="application/xml", headers={"Cache-Control": "public, max-age=3600"})


@api.get("/admin/documents/template.docx")
async def bulk_import_word_template():
    """Blank Word template matching the labeled-block parser's format."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("Blackrod Now — Bulk Event Upload", 0)
    intro = doc.add_paragraph(
        "Fill in one numbered block per event using the labels below, then upload this "
        "document via Admin → Bulk document import. Every field is read exactly as written — "
        "no AI guesswork. Leave a field as 'Not published' (or delete the line) if unknown."
    )
    intro.runs[0].font.size = Pt(10)

    def block(n, values):
        doc.add_heading(f"{n}. {values['Event title']}", level=2)
        for i, (label, value) in enumerate(values.items(), start=1):
            doc.add_paragraph(f"{i}. {label}: {value}")

    block(1, {
        "Event title": "Example Coffee Morning",
        "Organisation": "Blackrod Community Group",
        "Category": "community",
        "Date": "Tuesday 6 October 2026",
        "Start": "10:00",
        "End": "11:30",
        "Venue": "Blackrod Community Centre",
        "Address": "Blackrod, Bolton, BL6",
        "Description": "Friendly weekly coffee morning — everyone welcome.",
        "Cost": "Free",
        "Age suitability": "All ages",
        "Accessibility": "Step-free access",
        "Booking link": "Not published",
        "Email": "hello@example.org",
        "Phone": "01204 000000",
        "Image URL": "Not published",
        "Repeats": "Yes — Tuesdays",
    })
    block(2, {
        "Event title": "Example Christmas Fair",
        "Organisation": "St Katharine's Church",
        "Category": "family",
        "Date": "Saturday 5 December 2026",
        "Start": "11:00",
        "End": "15:00",
        "Venue": "Church Hall",
        "Address": "Church Street, Blackrod",
        "Description": "Stalls, mulled wine and Santa's grotto.",
        "Cost": "50p entry",
        "Age suitability": "All ages",
        "Accessibility": "Accessible entrance at the rear",
        "Booking link": "https://example.org/fair",
        "Email": "Not published",
        "Phone": "Not published",
        "Image URL": "Not published",
        "Repeats": "No — one-off",
    })
    buf = BytesIO()
    doc.save(buf)
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="blackrod-now-events-template.docx"'},
    )


@api.get("/admin/documents/template.xlsx")
async def bulk_import_template():
    """Blank spreadsheet template matching the structured parser's columns."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill

    wb = Workbook()
    ws = wb.active
    ws.title = "Events"
    headers = [
        "Organisation", "Title", "Date", "End date", "Start time", "End time", "Venue", "Address",
        "Category", "Fee", "Age", "Accessibility", "Booking info", "URL", "Contact email", "Contact phone",
        "Repeats", "Description",
    ]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="0052FF", end_color="0052FF", fill_type="solid")
    ws.append([
        "Blackrod Library", "Summer Craft Morning", "2026-07-21", "", "10am", "12pm", "Blackrod Library",
        "Church Street, Blackrod BL6 5EQ", "Family", "Free", "All ages", "Step-free access", "Just turn up",
        "https://example.org/craft", "library@example.org", "01204 000000", "", "Crafts, stories and songs for families.",
    ])
    ws.append([
        "St Katharine's Church", "Heritage Talk", "05/08/2026", "", "7:30pm", "9pm", "Church Hall",
        "Blackrod Brow, Blackrod BL6 5NA", "Heritage", "£3", "Adults", "Hearing loop", "Book via 01204 000000",
        "", "", "01204 000000", "First Wednesday of the month", "A talk on the history of Blackrod.",
    ])
    widths = [24, 30, 12, 12, 11, 11, 24, 30, 16, 10, 14, 18, 26, 30, 24, 16, 24, 40]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[ws.cell(row=1, column=i).column_letter].width = w
    buf = BytesIO()
    wb.save(buf)
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": 'attachment; filename="blackrod-now-events-template.xlsx"'},
    )


@api.post("/admin/documents/parse", response_model=BulkParseResponse)
async def admin_parse_documents(
    files: List[UploadFile] = File(default=[]),
    source_org_slug: Optional[str] = Form(None),
    urls_json: Optional[str] = Form(None),
    texts_json: Optional[str] = Form(None),
):
    file_sources = await _read_upload_sources(files)
    return await _admin_parse_documents(file_sources, source_org_slug=source_org_slug, urls_json=urls_json, texts_json=texts_json)


@api.post("/admin/documents/parse-jobs")
async def admin_create_parse_job(
    files: List[UploadFile] = File(default=[]),
    source_org_slug: Optional[str] = Form(None),
    urls_json: Optional[str] = Form(None),
    texts_json: Optional[str] = Form(None),
):
    # IMPORTANT: this endpoint does no parsing. It only validates + persists the
    # sources, then returns immediately. A durable Mongo-backed worker processes
    # the job independently of the browser request.
    file_sources = await _read_upload_sources(files)
    url_list = _decode_bulk_sources(urls_json)
    text_list = _decode_bulk_sources(texts_json)
    total = len(file_sources) + len(url_list) + len(text_list)

    if total == 0:
        raise HTTPException(400, "Add files, links or pasted text first")
    if total > MAX_BULK_PARSE_FILES:
        raise HTTPException(400, f"Upload at most {MAX_BULK_PARSE_FILES} total sources per bulk parse request")
    if source_org_slug:
        await _find_org(source_org_slug)

    job_id = new_id()
    job = {
        "id": job_id,
        "status": "queued",
        "total": total,
        "done": 0,
        "current": "Waiting for parser",
        "error": None,
        "source_org_slug": source_org_slug or "",
        "file_count": len(file_sources),
        "url_count": len(url_list),
        "text_count": len(text_list),
        "attempts": 0,
        "recoveries": 0,
        "created_at": now_iso(),
        "updated_at": now_iso(),
    }

    await db.parse_jobs.insert_one(dict(job))
    try:
        await _persist_parse_job_sources(job_id, file_sources, url_list, text_list)
    except Exception as exc:
        await db.parse_jobs.delete_one({"id": job_id})
        await db.parse_job_sources.delete_many({"job_id": job_id})
        logger.exception("Could not persist parse-job sources for %s: %s", job_id, exc)
        raise HTTPException(500, "Could not safely queue the import. Nothing was parsed; please try again.")

    return {
        "job_id": job_id,
        "total": total,
        "status": "queued",
        "message": "Import safely queued",
    }


@api.get("/admin/documents/parse-jobs/{job_id}")
async def admin_get_parse_job(job_id: str):
    job = await db.parse_jobs.find_one({"id": job_id}, {"_id": 0, "partial_docs": 0})
    if not job:
        raise HTTPException(404, "Parse job not found")

    # A stale processing job is recoverable. Re-queue it rather than telling the
    # administrator to upload the source again.
    if _parse_job_is_stale(job):
        attempts = int(job.get("attempts") or 0)
        target_status = "failed" if attempts >= PARSE_JOB_MAX_ATTEMPTS else "queued"
        target_error = (
            f"Parser stopped responding after {attempts} attempts."
            if target_status == "failed"
            else PARSE_JOB_INTERRUPTED_ERROR
        )
        await db.parse_jobs.update_one(
            {
                "id": job_id,
                "status": "processing",
                "updated_at": job.get("updated_at"),
            },
            {
                "$set": {
                    "status": target_status,
                    "current": "" if target_status == "failed" else "Recovering interrupted import",
                    "error": target_error,
                    "worker_id": None,
                    "updated_at": now_iso(),
                },
                "$inc": {"recoveries": 1},
            },
        )
        job = await db.parse_jobs.find_one({"id": job_id}, {"_id": 0}) or job

    return job


# ─────────── Newsletter & broadcast ───────────
def _render_welcome(email: str, unsub_link: str, pref_link: str) -> str:
    return f"""
<!DOCTYPE html><html><body style="font-family:Helvetica,Arial,sans-serif;background:#F4F5F7;padding:24px;color:#0F172A">
  <table role="presentation" width="100%" cellpadding="0" cellspacing="0" style="max-width:560px;margin:0 auto;background:#fff;border-radius:24px;overflow:hidden">
    <tr><td style="padding:32px">
      <h1 style="margin:0;font-size:28px;color:#0F172A">Welcome to <span style="color:#0052FF">Blackrod Now</span> 👋</h1>
      <p style="margin:16px 0 0;font-size:15px;color:#475569">You're on the list! We'll send you a friendly Friday round-up of the best events, freebies and good causes across Blackrod.</p>
      <p style="margin:16px 0 0;font-size:15px;color:#475569">Want it more personalised? <a href="{pref_link}" style="color:#0052FF">Choose which organisations and categories</a> you care about.</p>
    </td></tr>
    <tr><td style="padding:16px 32px 32px;color:#94A3B8;font-size:12px;text-align:center;border-top:1px solid #E2E8F0">
      Blackrod Now · Blackrod, Bolton · <a href="{unsub_link}" style="color:#94A3B8">Unsubscribe</a>
    </td></tr>
  </table>
</body></html>
"""


def _gcal_url(event: dict) -> str:
    """Google Calendar 'add event' URL — works from any email client."""
    try:
        start = datetime.fromisoformat((event.get("start") or "").replace("Z", "+00:00"))
        end = datetime.fromisoformat((event.get("end") or event.get("start") or "").replace("Z", "+00:00"))
    except Exception:
        return ""
    def _fmt(dt: datetime) -> str:
        return dt.astimezone(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    params = {
        "action": "TEMPLATE",
        "text": event.get("title", "Blackrod Now event"),
        "dates": f"{_fmt(start)}/{_fmt(end)}",
        "details": (event.get("description") or "")[:1500],
        "location": " ".join([s for s in [event.get("venue"), event.get("address")] if s]),
    }
    from urllib.parse import urlencode
    return "https://calendar.google.com/calendar/render?" + urlencode(params, safe=":/")


def _outlook_url(event: dict) -> str:
    try:
        start = datetime.fromisoformat((event.get("start") or "").replace("Z", "+00:00"))
        end = datetime.fromisoformat((event.get("end") or event.get("start") or "").replace("Z", "+00:00"))
    except Exception:
        return ""
    from urllib.parse import urlencode
    params = {
        "path": "/calendar/action/compose",
        "rru": "addevent",
        "subject": event.get("title", "Blackrod Now event"),
        "startdt": start.astimezone(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
        "enddt": end.astimezone(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
        "body": (event.get("description") or "")[:1500],
        "location": " ".join([s for s in [event.get("venue"), event.get("address")] if s]),
    }
    return "https://outlook.live.com/calendar/0/deeplink/compose?" + urlencode(params, safe=":/")


def _render_digest(sub: dict, events: List[dict], updates: List[dict]) -> str:
    unsub = f"{PUBLIC_URL}/unsubscribe/{sub['unsub_token']}"
    pref = f"{PUBLIC_URL}/preferences/{sub['pref_token']}"
    def fmt_time(iso):
        try:
            return datetime.fromisoformat(iso.replace("Z", "+00:00")).strftime("%a %d %b · %H:%M")
        except Exception:
            return ""
    def _event_row(e: dict) -> str:
        gcal = _gcal_url(e)
        outlook = _outlook_url(e)
        ics = f"{PUBLIC_URL}/api/events/{e.get('id')}.ics" if e.get("id") else ""
        cal_row = "".join([
            f'<a href="{gcal}" style="display:inline-block;margin-right:8px;padding:6px 12px;border-radius:999px;background:#0052FF;color:#fff;font-size:12px;text-decoration:none;font-weight:700">Google</a>' if gcal else "",
            f'<a href="{outlook}" style="display:inline-block;margin-right:8px;padding:6px 12px;border-radius:999px;background:#0F172A;color:#fff;font-size:12px;text-decoration:none;font-weight:700">Outlook</a>' if outlook else "",
            f'<a href="{ics}" style="display:inline-block;padding:6px 12px;border-radius:999px;background:#475569;color:#fff;font-size:12px;text-decoration:none;font-weight:700">Apple / iCal</a>' if ics else "",
        ])
        return f"""<tr><td style="padding:14px 0;border-bottom:1px solid #E2E8F0">
            <div style="font-weight:700;color:#0F172A;font-size:15px">{e.get('title','')}</div>
            <div style="color:#0052FF;font-size:13px;margin-top:2px">{fmt_time(e.get('start',''))}</div>
            <div style="color:#475569;font-size:13px">{e.get('venue','')}</div>
            <div style="margin-top:8px">{cal_row}</div>
        </td></tr>"""
    ev_rows = "".join(_event_row(e) for e in events[:8]) or "<tr><td style='color:#94A3B8;padding:12px 0'>No events matching your preferences this week.</td></tr>"
    up_rows = "".join(
        f"<li style='margin:8px 0;color:#475569'><b style='color:#0F172A'>{u.get('title','')}</b> — {u.get('body','')[:120]}</li>"
        for u in updates[:5]
    )
    return f"""
<!DOCTYPE html><html><body style="font-family:Helvetica,Arial,sans-serif;background:#F4F5F7;padding:24px;color:#0F172A">
  <table role="presentation" width="100%" cellpadding="0" cellspacing="0" style="max-width:600px;margin:0 auto;background:#fff;border-radius:24px;overflow:hidden">
    <tr><td style="padding:28px 28px 8px">
      <h1 style="margin:0;font-size:24px">Your Blackrod Now digest 📬</h1>
      <p style="color:#475569;margin:8px 0 0">Curated for you — following {len(sub.get('followed_orgs',[]))} orgs and {len(sub.get('followed_categories',[]))} categories.</p>
    </td></tr>
    <tr><td style="padding:8px 28px">
      <h2 style="margin:16px 0 4px;font-size:16px;color:#0F172A">What's on</h2>
      <p style="margin:0 0 8px;font-size:12px;color:#94A3B8">Tap Google / Outlook / Apple to add straight to your calendar.</p>
      <table role="presentation" width="100%">{ev_rows}</table>
    </td></tr>
    {"<tr><td style='padding:8px 28px'><h2 style='margin:16px 0 4px;font-size:16px'>Local updates</h2><ul style='padding-left:16px'>" + up_rows + "</ul></td></tr>" if up_rows else ""}
    <tr><td style="padding:20px 28px 28px;color:#94A3B8;font-size:12px;text-align:center;border-top:1px solid #E2E8F0">
      <a href="{pref}" style="color:#0052FF">Edit preferences</a> · <a href="{unsub}" style="color:#94A3B8">Unsubscribe</a><br/>
      Blackrod Now · Made in Blackrod, Bolton
    </td></tr>
  </table>
</body></html>
"""


async def _events_for_sub(sub: dict) -> List[dict]:
    """Personalised upcoming events for a subscriber."""
    nowi = now_iso()
    q: Dict[str, Any] = {"status": "approved"}
    ors = []
    if sub.get("followed_orgs"):
        ors.append({"orgSlug": {"$in": sub["followed_orgs"]}})
    if sub.get("followed_categories"):
        ors.append({"category": {"$in": sub["followed_categories"]}})
    if ors:
        q["$or"] = ors
    events = await db.events.find(q, {"_id": 0}).to_list(200)
    events = [e for e in events if (e.get("end") or e.get("start", "")) >= nowi]
    events.sort(key=lambda e: e.get("start", ""))
    if not events:  # fallback: any upcoming
        events = await db.events.find({"status": "approved"}, {"_id": 0}).to_list(200)
        events = [e for e in events if (e.get("end") or e.get("start", "")) >= nowi]
        events.sort(key=lambda e: e.get("start", ""))
    return events


async def _updates_for_sub(sub: dict) -> List[dict]:
    q: Dict[str, Any] = {}
    if sub.get("followed_orgs"):
        q["orgSlug"] = {"$in": sub["followed_orgs"]}
    return await db.feed.find(q, {"_id": 0}).sort("time", -1).to_list(20)


@api.get("/admin/newsletter/preview")
async def newsletter_preview(email: Optional[str] = None):
    """Returns the personalised HTML for a subscriber (or a generic one)."""
    sub = None
    if email:
        sub = await db.subscribers.find_one({"email": email.lower(), "unsubscribed": False}, {"_id": 0})
    if not sub:
        # generic preview: no follows
        sub = {"email": "preview@example.com", "unsub_token": "PREVIEW", "pref_token": "PREVIEW", "followed_orgs": [], "followed_categories": []}
    events = await _events_for_sub(sub)
    updates = await _updates_for_sub(sub)
    return {"subject": "Your Blackrod Now digest 📬", "html": _render_digest(sub, events, updates), "sub": sub, "matched_events": len(events), "matched_updates": len(updates)}


class NewsletterEditReq(BaseModel):
    subject: str
    body_intro: Optional[str] = ""
    scheduled_for: Optional[str] = None
    audience: Literal["subscribers", "orgs_all", "orgs_selected"] = "subscribers"
    org_slugs: List[str] = []


async def _newsletter_recipients(req: NewsletterEditReq) -> List[Dict[str, Any]]:
    recipients: List[Dict[str, Any]] = []
    seen: set[str] = set()

    if req.audience == "subscribers":
        subs = await db.subscribers.find({"unsubscribed": False, "digest": True}, {"_id": 0}).to_list(5000)
        for sub in subs:
            email = (sub.get("email") or "").strip().lower()
            if not email or email in seen:
                continue
            seen.add(email)
            recipients.append({"email": email, "kind": "subscriber", "subscriber": sub})
        return recipients

    org_query: Dict[str, Any] = {"status": {"$ne": "rejected"}}
    if req.audience == "orgs_selected":
        clean_slugs = [slug.strip() for slug in (req.org_slugs or []) if slug.strip()]
        if not clean_slugs:
            raise HTTPException(400, "Select at least one organisation slug")
        org_query["slug"] = {"$in": clean_slugs}

    orgs = await db.orgs.find(org_query, {"_id": 0, "slug": 1, "name": 1, "email": 1}).to_list(5000)
    for org in orgs:
        email = (org.get("email") or "").strip().lower()
        if not email or not EMAIL_RE.match(email) or email in seen:
            continue
        seen.add(email)
        recipients.append({"email": email, "kind": "org", "org": org})
    return recipients


@api.post("/admin/newsletter/edition")
async def upsert_edition(req: NewsletterEditReq):
    ed = NewsletterEdition(subject=req.subject, body_intro=req.body_intro or "", scheduled_for=req.scheduled_for, body_html="")
    await db.newsletter.insert_one(ed.model_dump())
    return ed


@api.post("/admin/newsletter/send")
async def send_newsletter(req: NewsletterEditReq):
    recipients = await _newsletter_recipients(req)
    sent, failed = 0, 0
    for recipient in recipients:
        if recipient["kind"] == "subscriber":
            sub = recipient["subscriber"]
            events = await _events_for_sub(sub)
            updates = await _updates_for_sub(sub)
            html = _render_digest(sub, events, updates)
        else:
            intro = (req.body_intro or "").strip() or "Community update from Blackrod Now."
            html = _render_admin_email_html(req.subject, intro, SENDER_EMAIL)
        result = await asyncio.to_thread(resend_send, recipient["email"], req.subject, html)
        if result.get("ok"):
            sent += 1
        else:
            failed += 1
    await db.newsletter.insert_one(
        NewsletterEdition(subject=req.subject, body_intro=req.body_intro or "", sent_at=now_iso(), body_html="").model_dump()
    )
    await _record_analytics_event("newsletter_send", entity_type="site", count=sent)
    return {
        "ok": True,
        "audience": req.audience,
        "sent": sent,
        "failed": failed,
        "recipient_count": len(recipients),
        "mocked": not RESEND_API_KEY,
    }


class BroadcastReq(BaseModel):
    subject: str
    html: str


@api.post("/admin/broadcast")
async def broadcast(req: BroadcastReq):
    subs = await db.subscribers.find({"unsubscribed": False}, {"_id": 0}).to_list(5000)
    sent, failed = 0, 0
    for sub in subs:
        unsub = f"{PUBLIC_URL}/unsubscribe/{sub['unsub_token']}"
        html = req.html + f"<hr style='margin:20px 0;border:none;border-top:1px solid #E2E8F0'/><p style='color:#94A3B8;font-size:12px;text-align:center'>Blackrod Now · <a href='{unsub}' style='color:#94A3B8'>Unsubscribe</a></p>"
        result = await asyncio.to_thread(resend_send, sub["email"], req.subject, html)
        if result.get("ok"):
            sent += 1
        else:
            failed += 1
    await _record_analytics_event("broadcast_send", entity_type="site", count=sent)
    return {"ok": True, "sent": sent, "failed": failed, "mocked": not RESEND_API_KEY}


# ─────────── Admin free-form email compose ───────────
import re as _re


EMAIL_RE = _re.compile(r"^[^\s@]+@[^\s@]+\.[^\s@]+$")


class AdminEmailReq(BaseModel):
    """Free-form email compose. Recipients accepted as a list or a
    comma/newline-separated string. `from_email` is validated against
    ADMIN_SENDER_EMAILS whitelist. `body` is treated as plain text with
    minimal auto-formatting (newlines → paragraphs, URLs auto-linked)."""
    to: str  # comma/newline-separated
    subject: str
    body: str
    from_email: Optional[str] = None
    from_name: Optional[str] = None
    reply_to: Optional[str] = None


def _coerce_admin_email_req(data: Dict[str, Any]) -> AdminEmailReq:
    return AdminEmailReq(
        to=data.get("to", ""),
        subject=data.get("subject", ""),
        body=data.get("body", ""),
        from_email=data.get("from_email") or None,
        from_name=data.get("from_name") or None,
        reply_to=data.get("reply_to") or None,
    )


async def _read_admin_email_request(request: Request) -> tuple[AdminEmailReq, list[dict[str, Any]]]:
    content_type = request.headers.get("content-type", "")
    if content_type.startswith("multipart/form-data"):
        form = await request.form()
        req = _coerce_admin_email_req({key: form.get(key) for key in ("to", "subject", "body", "from_email", "from_name", "reply_to")})
        attachments: list[dict[str, Any]] = []
        for upload in [item for item in form.getlist("attachments") if getattr(item, "filename", None)]:
            data = await upload.read()
            if len(data) > 10 * 1024 * 1024:
                raise HTTPException(413, f"Attachment too large (max 10 MB): {upload.filename or 'attachment'}")
            attachments.append(
                {
                    "filename": upload.filename or "attachment",
                    "content_type": upload.content_type or "application/octet-stream",
                    "size": len(data),
                    "content": base64.b64encode(data).decode("ascii"),
                }
            )
        return req, attachments

    payload = await request.json()
    req = _coerce_admin_email_req(payload or {})
    attachments: list[dict[str, Any]] = []
    for item in (payload or {}).get("attachments", []) if isinstance(payload, dict) else []:
        if not isinstance(item, dict):
            continue
        content = item.get("content")
        if not content:
            continue
        if isinstance(content, str):
            try:
                raw = base64.b64decode(content)
            except Exception:
                raw = content.encode("utf-8")
        else:
            raw = bytes(content)
        attachments.append(
            {
                "filename": item.get("filename") or item.get("name") or "attachment",
                "content_type": item.get("content_type") or item.get("type") or "application/octet-stream",
                "size": len(raw),
                "content": base64.b64encode(raw).decode("ascii"),
            }
        )
    return req, attachments


def _parse_recipients(raw: str) -> tuple[list[str], list[str]]:
    seen: set[str] = set()
    valid: list[str] = []
    invalid: list[str] = []
    for chunk in _re.split(r"[,;\n]+", raw or ""):
        addr = chunk.strip()
        if not addr:
            continue
        low = addr.lower()
        if low in seen:
            continue
        seen.add(low)
        (valid if EMAIL_RE.match(addr) else invalid).append(addr)
    return valid, invalid


def _auto_link(text: str) -> str:
    """Escape HTML, convert URLs to <a> links, preserve newlines/paragraphs."""
    escaped = _html_lib.escape(text or "")
    # Don't include common sentence-ending punctuation in the linked URL —
    # e.g. `visit https://blackrodnow.com.` should link "https://blackrodnow.com"
    # and leave the trailing period as prose.
    url_re = _re.compile(r"(https?://[^\s<]+?)([.,;:!?)\]]*)(?=\s|$|<)")
    linked = url_re.sub(r'<a href="\1" style="color:#0052FF">\1</a>\2', escaped)
    # split into paragraphs on blank line, single newlines become <br>
    paragraphs = [p.strip() for p in _re.split(r"\n\s*\n", linked) if p.strip()]
    return "\n".join(f"<p>{p.replace(chr(10), '<br />')}</p>" for p in paragraphs) or "<p></p>"


_EMAIL_SIGNATURE = """<table role="presentation" cellpadding="0" cellspacing="0" border="0" style="width:100%;max-width:610px;border-collapse:collapse;font-family:Arial,Helvetica,sans-serif;color:#0F1729;margin-top:32px">
  <tr>
    <td valign="middle" style="width:175px;padding:18px 24px 18px 0">
      <div style="font-size:23px;line-height:24px;font-weight:800;letter-spacing:-0.6px;color:#0F1729;margin:0">BLACKROD</div>
      <div style="font-size:31px;line-height:32px;font-weight:900;letter-spacing:1px;color:#004DFF;margin:0">NOW</div>
      <div style="margin-top:8px;font-size:9px;line-height:13px;font-weight:700;letter-spacing:0.35px;color:#667085;text-transform:uppercase">What&#39;s on. What&#39;s new.<br>What&#39;s next.</div>
    </td>
    <td valign="middle" style="width:5px;padding:0">
      <table role="presentation" cellpadding="0" cellspacing="0" border="0" style="height:96px">
        <tr><td style="width:3px;height:67px;background:#004DFF;font-size:1px;line-height:1px">&nbsp;</td></tr>
        <tr><td style="width:3px;height:29px;background:#D5FF00;font-size:1px;line-height:1px">&nbsp;</td></tr>
      </table>
    </td>
    <td valign="middle" style="padding:17px 0 17px 25px">
      <div style="font-size:18px;line-height:22px;font-weight:800;color:#0F1729">Blackrod Now</div>
      <div style="margin-top:2px;font-size:12px;line-height:17px;font-weight:700;color:#004DFF">Administration</div>
      <div style="width:42px;height:3px;background:#D5FF00;margin-top:10px;margin-bottom:10px;font-size:1px;line-height:1px">&nbsp;</div>
      <table role="presentation" cellpadding="0" cellspacing="0" border="0" style="border-collapse:collapse">
        <tr>
          <td style="padding:0 11px 3px 0;font-size:10px;line-height:17px;font-weight:700;color:#98A2B3;text-transform:uppercase">E</td>
          <td style="padding:0 0 3px 0;font-size:12px;line-height:17px"><a href="mailto:blackrodnow@communityalliances.co.uk" style="color:#344054;text-decoration:none;font-weight:500">blackrodnow@communityalliances.co.uk</a></td>
        </tr>
        <tr>
          <td style="padding:0 11px 0 0;font-size:10px;line-height:17px;font-weight:700;color:#98A2B3;text-transform:uppercase">W</td>
          <td style="padding:0;font-size:12px;line-height:17px"><a href="https://blackrodnow.com/" style="color:#004DFF;text-decoration:none;font-weight:700">blackrodnow.com</a></td>
        </tr>
      </table>
    </td>
  </tr>
  <tr>
    <td colspan="3" style="padding:0;border-top:1px solid #E4E7EC;font-size:1px;line-height:1px">&nbsp;</td>
  </tr>
  <tr>
    <td colspan="3" style="padding-top:9px;font-size:10px;line-height:15px;color:#98A2B3">
      Blackrod&#39;s community platform for <span style="color:#667085">events, organisations, volunteering, venues and local information.</span>
    </td>
  </tr>
</table>"""


def _render_admin_email_html(subject: str, body_text: str, from_email: str) -> str:
    body_html = _auto_link(body_text)
    return f"""<!doctype html>
<html><body style="margin:0;padding:0;background:#F9FAFB;font-family:system-ui,-apple-system,Segoe UI,sans-serif;color:#111827">
<table role="presentation" cellspacing="0" cellpadding="0" width="100%" bgcolor="#F9FAFB">
<tr><td align="center">
<table role="presentation" cellspacing="0" cellpadding="0" width="620" style="max-width:620px;padding:32px 24px">
  <tr><td style="padding-bottom:16px">
    <div style="font-size:11px;letter-spacing:0.2em;text-transform:uppercase;color:#0052FF;font-weight:800">Blackrod Now</div>
    <h1 style="font-size:24px;line-height:1.25;color:#111827;margin:8px 0 0 0">{_html_lib.escape(subject)}</h1>
  </td></tr>
  <tr><td style="font-size:15px;line-height:1.6;color:#111827;padding-bottom:24px">
    {body_html}
  </td></tr>
  <tr><td style="padding-bottom:24px">
    {_EMAIL_SIGNATURE}
  </td></tr>
  <tr><td style="border-top:1px solid #E5E7EB;padding-top:16px;color:#9CA3AF;font-size:12px">
    Sent by Blackrod Now via <a href="mailto:{_html_lib.escape(from_email)}" style="color:#9CA3AF">{_html_lib.escape(from_email)}</a>. This is a one-off message from the Blackrod Now admin team.
  </td></tr>
</table>
</td></tr>
</table>
</body></html>"""


@api.get("/admin/email/senders")
async def admin_email_senders():
    """Return the whitelisted sender addresses for the admin compose UI."""
    return {"senders": ADMIN_SENDER_EMAILS, "default": SENDER_EMAIL}


@api.post("/admin/email/preview")
async def admin_email_preview(request: Request):
    req, attachments = await _read_admin_email_request(request)
    from_addr = (req.from_email or SENDER_EMAIL).strip().lower()
    if from_addr not in [s.lower() for s in ADMIN_SENDER_EMAILS]:
        raise HTTPException(400, f"Sender must be one of: {', '.join(ADMIN_SENDER_EMAILS)}")
    valid, invalid = _parse_recipients(req.to)
    html = _render_admin_email_html(req.subject, req.body, from_addr)
    return {
        "html": html,
        "subject": req.subject,
        "from": f'"{req.from_name or SENDER_NAME}" <{from_addr}>',
        "recipients": valid,
        "invalid_recipients": invalid,
        "count": len(valid),
        "attachments": [{"filename": a["filename"], "content_type": a["content_type"], "size": a["size"]} for a in attachments],
        "attachment_count": len(attachments),
    }


@api.post("/admin/email/send")
async def admin_email_send(request: Request):
    req, attachments = await _read_admin_email_request(request)
    from_addr = (req.from_email or SENDER_EMAIL).strip().lower()
    if from_addr not in [s.lower() for s in ADMIN_SENDER_EMAILS]:
        raise HTTPException(400, f"Sender must be one of: {', '.join(ADMIN_SENDER_EMAILS)}")
    valid, invalid = _parse_recipients(req.to)
    if not valid:
        raise HTTPException(400, "No valid recipients")
    if not req.subject.strip():
        raise HTTPException(400, "Subject is required")
    if not req.body.strip():
        raise HTTPException(400, "Body is required")

    html = _render_admin_email_html(req.subject, req.body, from_addr)
    sent, failed, results = 0, 0, []
    for addr in valid:
        r = await asyncio.to_thread(
            resend_send, addr, req.subject, html, from_addr, req.from_name or SENDER_NAME, None, attachments
        )
        if r.get("ok"):
            sent += 1
        else:
            failed += 1
        results.append({"to": addr, **{k: r.get(k) for k in ("ok", "id", "error", "mocked", "attachments")}})
    await _record_analytics_event("admin_email_send", entity_type="site", count=sent)
    return {
        "ok": failed == 0,
        "sent": sent,
        "failed": failed,
        "invalid_recipients": invalid,
        "results": results,
        "mocked": not RESEND_API_KEY,
        "attachments": [{"filename": a["filename"], "content_type": a["content_type"], "size": a["size"]} for a in attachments],
    }


# ─────────── Startup ───────────

# ─────────── Reminder emails for saved events ───────────
_REMINDER_LOOP_INTERVAL_SECONDS = int(os.environ.get("REMINDER_LOOP_INTERVAL", "900"))  # 15 min
_REMINDER_WINDOW_MINUTES = int(os.environ.get("REMINDER_WINDOW_MINUTES", "20"))


def _render_reminder_email(sub_email: str, event: dict, kind: str) -> tuple[str, str]:
    """Return (subject, html) for a 24h or 2h reminder."""
    when_iso = event.get("start") or ""
    try:
        start = datetime.fromisoformat(when_iso.replace("Z", "+00:00"))
        when_pretty = start.strftime("%A %d %B · %H:%M")
    except Exception:
        when_pretty = when_iso
    label = "starts tomorrow" if kind == "24h" else "starts in 2 hours"
    subject = f"Reminder: {event.get('title','Your saved event')} {label}"
    gcal = _gcal_url(event)
    outlook = _outlook_url(event)
    ics = f"{PUBLIC_URL}/api/events/{event.get('id')}.ics" if event.get("id") else ""
    cal_row = "".join([
        f'<a href="{gcal}" style="display:inline-block;margin-right:8px;padding:8px 14px;border-radius:999px;background:#0052FF;color:#fff;font-size:13px;text-decoration:none;font-weight:700">Google</a>' if gcal else "",
        f'<a href="{outlook}" style="display:inline-block;margin-right:8px;padding:8px 14px;border-radius:999px;background:#0F172A;color:#fff;font-size:13px;text-decoration:none;font-weight:700">Outlook</a>' if outlook else "",
        f'<a href="{ics}" style="display:inline-block;padding:8px 14px;border-radius:999px;background:#475569;color:#fff;font-size:13px;text-decoration:none;font-weight:700">Apple / iCal</a>' if ics else "",
    ])
    event_url = f"{PUBLIC_URL}/events/{event.get('id')}" if event.get("id") else PUBLIC_URL
    html = f"""<!DOCTYPE html><html><body style="font-family:Helvetica,Arial,sans-serif;background:#F4F5F7;padding:24px;color:#0F172A">
<table role="presentation" width="100%" cellpadding="0" cellspacing="0" style="max-width:600px;margin:0 auto;background:#fff;border-radius:24px;overflow:hidden">
  <tr><td style="padding:28px">
    <div style="font-size:12px;color:#0052FF;letter-spacing:2px;text-transform:uppercase;font-weight:800">Reminder</div>
    <h1 style="margin:8px 0 0;font-size:24px;color:#0F172A">{event.get('title','Your saved event')}</h1>
    <p style="margin:16px 0 0;font-size:15px;color:#475569">You saved this event on Blackrod Now — it {label}.</p>
    <div style="margin-top:20px;padding:16px;border-radius:16px;background:#F1F5F9">
      <div style="font-size:14px;color:#0F172A"><b>When:</b> {when_pretty}</div>
      <div style="font-size:14px;color:#0F172A;margin-top:4px"><b>Where:</b> {event.get('venue','')}{(", " + event.get('address','')) if event.get('address') else ""}</div>
      {"<div style='margin-top:8px;font-size:13px;color:#475569'>" + (event.get('description') or "")[:280] + "</div>" if event.get('description') else ""}
    </div>
    <div style="margin-top:16px">{cal_row}</div>
    <p style="margin:20px 0 0;font-size:13px"><a href="{event_url}" style="color:#0052FF;font-weight:600">View event page →</a></p>
    <p style="margin-top:24px;font-size:11px;color:#94A3B8">You're receiving this because you saved this event. You can un-save it any time from the event page or your Blackrod Now saved list.</p>
  </td></tr>
</table></body></html>"""
    return subject, html


async def _find_events_in_window(minutes_from_now: int, window_minutes: int) -> List[dict]:
    now = datetime.now(timezone.utc)
    target = now + timedelta(minutes=minutes_from_now)
    lo = (target - timedelta(minutes=window_minutes / 2)).isoformat()
    hi = (target + timedelta(minutes=window_minutes / 2)).isoformat()
    return await db.events.find(
        {"status": "approved", "start": {"$gte": lo, "$lte": hi}}, {"_id": 0}
    ).to_list(200)


async def _process_reminders(kind: str, minutes_from_now: int) -> int:
    events = await _find_events_in_window(minutes_from_now, _REMINDER_WINDOW_MINUTES)
    if not events:
        return 0
    sent = 0
    for event in events:
        event_id = event.get("id")
        if not event_id:
            continue
        # Every subscriber who saved this event and is still active.
        cursor = db.subscribers.find(
            {"saved_events": event_id, "unsubscribed": {"$ne": True}, "digest": {"$ne": False}},
            {"_id": 0, "email": 1},
        )
        async for sub in cursor:
            email = sub.get("email")
            if not email:
                continue
            key = {"email": email, "event_id": event_id, "kind": kind}
            already = await db.event_reminders_sent.find_one(key)
            if already:
                continue
            subject, html = _render_reminder_email(email, event, kind)
            try:
                await asyncio.to_thread(resend_send, email, subject, html)
                await db.event_reminders_sent.insert_one({**key, "sent_at": now_iso()})
                sent += 1
            except Exception as e:
                logger.warning("reminder send failed for %s / %s: %s", email, event_id, e)
    if sent:
        logger.info("Sent %d %s reminders", sent, kind)
    return sent


async def _event_reminder_loop() -> None:
    """Run forever — fires 24h + 2h reminders for saved events."""
    if not RESEND_API_KEY:
        logger.info("Reminder loop skipped: RESEND_API_KEY not set.")
        return
    logger.info("Event reminder loop started (interval=%ds, window=±%dmin)", _REMINDER_LOOP_INTERVAL_SECONDS, _REMINDER_WINDOW_MINUTES // 2)
    while True:
        try:
            await _process_reminders("24h", 24 * 60)
            await _process_reminders("2h", 120)
        except Exception as e:
            logger.exception("reminder loop iteration failed: %s", e)
        await asyncio.sleep(_REMINDER_LOOP_INTERVAL_SECONDS)


@api.post("/admin/reminders/run-now")
async def admin_reminders_run_now(
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    """Force the reminder loop to run once (for smoke testing)."""
    ok = read_admin_from_request(request, authorization) is not None or (
        admin_code and hmac.compare_digest(admin_code, ADMIN_LAUNCH_CODE)
    )
    if not ok:
        raise HTTPException(403, "Admin authentication required")
    sent_24 = await _process_reminders("24h", 24 * 60)
    sent_2 = await _process_reminders("2h", 120)
    return {"ok": True, "sent_24h": sent_24, "sent_2h": sent_2}




# ─────────── Batch D: Scheduled broadcasts + moderation ───────────
class ScheduledBroadcastCreate(BaseModel):
    to: str
    subject: str
    body: str
    from_email: Optional[str] = None
    from_name: Optional[str] = None
    scheduled_for: str  # ISO datetime, UTC preferred


@api.post("/admin/broadcasts/schedule")
async def schedule_broadcast(
    req: ScheduledBroadcastCreate,
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    ok = read_admin_from_request(request, authorization) is not None or (
        admin_code and hmac.compare_digest(admin_code, ADMIN_LAUNCH_CODE)
    )
    if not ok:
        raise HTTPException(403, "Admin authentication required")
    if not req.subject.strip() or not req.body.strip() or not req.to.strip():
        raise HTTPException(400, "to, subject and body are required")
    try:
        when = datetime.fromisoformat(req.scheduled_for.replace("Z", "+00:00"))
        if when.tzinfo is None:
            when = when.replace(tzinfo=timezone.utc)
    except Exception:
        raise HTTPException(400, "scheduled_for must be an ISO datetime")
    if when < datetime.now(timezone.utc) - timedelta(minutes=1):
        raise HTTPException(400, "scheduled_for must be in the future")
    doc = {
        "id": new_id(),
        "to": req.to,
        "subject": req.subject.strip(),
        "body": req.body,
        "from_email": req.from_email or SENDER_EMAIL,
        "from_name": req.from_name or SENDER_NAME,
        "scheduled_for": when.isoformat(),
        "status": "scheduled",
        "created_at": now_iso(),
        "sent_at": None,
        "result": None,
    }
    await db.scheduled_broadcasts.insert_one(doc)
    return {k: v for k, v in doc.items() if k != "_id"}


class BroadcastPreviewReq(BaseModel):
    """Send a one-off preview of a broadcast to a single address (usually the
    admin's own inbox) before scheduling — the classic "email it to me first"
    safety check."""
    subject: str
    body: str
    from_email: Optional[str] = None
    from_name: Optional[str] = None
    preview_to: str  # single email address


@api.post("/admin/broadcasts/preview")
async def broadcast_preview(
    req: BroadcastPreviewReq,
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    ok = read_admin_from_request(request, authorization) is not None or (
        admin_code and hmac.compare_digest(admin_code, ADMIN_LAUNCH_CODE)
    )
    if not ok:
        raise HTTPException(403, "Admin authentication required")
    if not req.subject.strip() or not req.body.strip() or not req.preview_to.strip():
        raise HTTPException(400, "subject, body and preview_to are required")
    addr = req.preview_to.strip()
    if not EMAIL_RE.match(addr):
        raise HTTPException(400, "preview_to must be a valid email address")
    from_addr = req.from_email or SENDER_EMAIL
    # Prepend a small "PREVIEW" banner so it's obvious in the inbox.
    preview_body = "🔍 PREVIEW — this email has been sent only to you.\n\n" + req.body
    html = _render_admin_email_html(f"[PREVIEW] {req.subject}", preview_body, from_addr)
    result = await asyncio.to_thread(
        resend_send, addr, f"[PREVIEW] {req.subject}", html,
        from_addr, req.from_name or SENDER_NAME, None, None,
    )
    await _record_analytics_event("admin_email_send", entity_type="site", count=1 if result.get("ok") else 0)
    return {
        "ok": bool(result.get("ok")),
        "mocked": bool(result.get("mocked")),
        "preview_to": addr,
        "id": result.get("id"),
    }


@api.get("/admin/broadcasts/scheduled")
async def list_scheduled_broadcasts(
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    ok = read_admin_from_request(request, authorization) is not None or (
        admin_code and hmac.compare_digest(admin_code, ADMIN_LAUNCH_CODE)
    )
    if not ok:
        raise HTTPException(403, "Admin authentication required")
    docs = await db.scheduled_broadcasts.find({}, {"_id": 0}).sort("scheduled_for", 1).to_list(200)
    return docs


@api.delete("/admin/broadcasts/scheduled/{broadcast_id}")
async def cancel_scheduled_broadcast(
    broadcast_id: str,
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    ok = read_admin_from_request(request, authorization) is not None or (
        admin_code and hmac.compare_digest(admin_code, ADMIN_LAUNCH_CODE)
    )
    if not ok:
        raise HTTPException(403, "Admin authentication required")
    res = await db.scheduled_broadcasts.update_one(
        {"id": broadcast_id, "status": "scheduled"},
        {"$set": {"status": "cancelled"}},
    )
    return {"ok": True, "cancelled": res.modified_count}


async def _process_scheduled_broadcasts() -> int:
    now = datetime.now(timezone.utc).isoformat()
    cursor = db.scheduled_broadcasts.find({"status": "scheduled", "scheduled_for": {"$lte": now}}, {"_id": 0})
    sent = 0
    async for doc in cursor:
        try:
            valid, invalid = _parse_recipients(doc.get("to", ""))
            html = _render_admin_email_html(doc.get("subject", ""), doc.get("body", ""), doc.get("from_email"))
            results = []
            for addr in valid:
                r = await asyncio.to_thread(
                    resend_send, addr, doc.get("subject", ""), html,
                    doc.get("from_email"), doc.get("from_name") or SENDER_NAME, None, None,
                )
                results.append({"to": addr, "ok": r.get("ok"), "id": r.get("id")})
            await db.scheduled_broadcasts.update_one(
                {"id": doc["id"]},
                {"$set": {
                    "status": "sent",
                    "sent_at": now_iso(),
                    "result": {"sent": sum(1 for x in results if x["ok"]), "failed": sum(1 for x in results if not x["ok"]), "invalid": invalid},
                }},
            )
            await _record_analytics_event("admin_email_send", entity_type="site", count=sum(1 for x in results if x["ok"]))
            sent += 1
        except Exception as e:
            logger.exception("Scheduled broadcast failed: %s", e)
            await db.scheduled_broadcasts.update_one({"id": doc["id"]}, {"$set": {"status": "failed", "result": {"error": str(e)}}})
    return sent


async def _scheduled_broadcast_loop() -> None:
    if not RESEND_API_KEY:
        logger.info("Scheduled broadcast loop skipped: RESEND_API_KEY not set.")
        return
    logger.info("Scheduled broadcast loop started")
    while True:
        try:
            await _process_scheduled_broadcasts()
        except Exception as e:
            logger.exception("scheduled broadcast loop iteration failed: %s", e)
        await asyncio.sleep(int(os.environ.get("SCHEDULED_BROADCAST_INTERVAL", "300")))


@api.post("/admin/broadcasts/scheduled/run-now")
async def run_scheduled_broadcasts_now(
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    ok = read_admin_from_request(request, authorization) is not None or (
        admin_code and hmac.compare_digest(admin_code, ADMIN_LAUNCH_CODE)
    )
    if not ok:
        raise HTTPException(403, "Admin authentication required")
    sent = await _process_scheduled_broadcasts()
    return {"ok": True, "sent": sent}


# ── Content moderation reports ──
_REPORT_RATE_LIMIT_MINUTES = 10


class ModerationReportReq(BaseModel):
    kind: Literal["event", "org", "feed", "venue", "volunteer"]
    target_id: str
    reason: Literal["spam", "inappropriate", "inaccurate", "outdated", "duplicate", "other"]
    notes: Optional[str] = None
    reporter_email: Optional[EmailStr] = None
    reporter_device: Optional[str] = None


@api.post("/reports")
async def submit_moderation_report(req: ModerationReportReq, request: Request):
    if not req.reporter_device and not req.reporter_email:
        raise HTTPException(400, "reporter_device or reporter_email required")
    # Basic rate limit — same reporter can't file more than 5 reports per 10 min.
    since = (datetime.now(timezone.utc) - timedelta(minutes=_REPORT_RATE_LIMIT_MINUTES)).isoformat()
    key: Dict[str, Any] = {"created_at": {"$gte": since}}
    if req.reporter_device:
        key["reporter_device"] = req.reporter_device
    else:
        key["reporter_email"] = req.reporter_email.lower() if req.reporter_email else None
    recent = await db.moderation_reports.count_documents(key)
    if recent >= 5:
        raise HTTPException(429, "Too many reports — please wait a few minutes")
    doc = {
        "id": new_id(),
        "kind": req.kind,
        "target_id": req.target_id,
        "reason": req.reason,
        "notes": (req.notes or "")[:1000],
        "reporter_email": req.reporter_email.lower() if req.reporter_email else None,
        "reporter_device": req.reporter_device,
        "status": "open",
        "created_at": now_iso(),
        "resolved_at": None,
        "resolution": None,
    }
    await db.moderation_reports.insert_one(doc)
    return {"ok": True, "id": doc["id"]}


@api.get("/admin/reports")
async def list_moderation_reports(
    request: Request,
    status: Optional[str] = "open",
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    ok = read_admin_from_request(request, authorization) is not None or (
        admin_code and hmac.compare_digest(admin_code, ADMIN_LAUNCH_CODE)
    )
    if not ok:
        raise HTTPException(403, "Admin authentication required")
    q: Dict[str, Any] = {}
    if status and status != "all":
        q["status"] = status
    reports = await db.moderation_reports.find(q, {"_id": 0}).sort("created_at", -1).to_list(300)
    return reports


class ModerationResolveReq(BaseModel):
    status: Literal["dismissed", "actioned"]
    resolution: Optional[str] = None


@api.post("/admin/reports/{report_id}/resolve")
async def resolve_moderation_report(
    report_id: str,
    req: ModerationResolveReq,
    request: Request,
    authorization: Optional[str] = Header(None),
    admin_code: Optional[str] = Header(None, alias="X-Admin-Code"),
):
    ok = read_admin_from_request(request, authorization) is not None or (
        admin_code and hmac.compare_digest(admin_code, ADMIN_LAUNCH_CODE)
    )
    if not ok:
        raise HTTPException(403, "Admin authentication required")
    res = await db.moderation_reports.update_one(
        {"id": report_id},
        {"$set": {"status": req.status, "resolution": req.resolution, "resolved_at": now_iso()}},
    )
    if not res.matched_count:
        raise HTTPException(404, "Report not found")
    return {"ok": True}



@app.on_event("startup")
async def startup():
    # Preload the heavy LLM libraries (litellm etc.) in a daemon thread so the
    # first AI call never does a multi-minute synchronous import on the event
    # loop thread (which froze the whole API / failed health probes on
    # CPU-limited production pods).
    threading.Thread(target=_preload_llm_libs, daemon=True).start()
    init_storage()
    # useful indexes (idempotent)
    try:
        await db.orgs.create_index("slug", unique=True)
        await db.events.create_index("id", unique=True)
        await db.subscribers.create_index("email", unique=True)
        await db.subscribers.create_index("unsub_token", unique=True)
        await db.subscribers.create_index("pref_token", unique=True)
        await db.org_passwords.create_index("slug", unique=True)
        await db.notifications.create_index([("org_slug", 1), ("created_at", -1)])
        await db.follows.create_index("device_id", unique=True)
        await db.analytics_events.create_index([("created_at", -1)])
        await db.analytics_events.create_index([("kind", 1), ("created_at", -1)])
        await db.analytics_events.create_index([("org_slug", 1), ("created_at", -1)])
        await db.analytics_events.create_index([("entity_id", 1), ("created_at", -1)])
        await db.messages.create_index([("delivery.provider", 1), ("delivery.event_id", 1)])
        await db.users.create_index("email", unique=True)
        await db.login_attempts.create_index("identifier")
        await db.event_reminders_sent.create_index([("email", 1), ("event_id", 1), ("kind", 1)], unique=True)
    except Exception as e:
        logger.warning("Index setup: %s", e)
    await _seed_admin_user()
    # Durable bulk-import queue indexes + recovery. Sources are stored in
    # separate documents so jobs survive browser refreshes and process restarts.
    try:
        await db.parse_jobs.create_index("id", unique=True)
        await db.parse_jobs.create_index([("status", 1), ("created_at", 1)])
        await db.parse_jobs.create_index([("status", 1), ("updated_at", 1)])
        await db.parse_job_sources.create_index([("job_id", 1), ("order", 1)])
        await db.parse_job_sources.create_index("expires_at", expireAfterSeconds=0)
        recovered = await _recover_stale_parse_jobs()
        if recovered:
            logger.warning("Recovered %s parse job(s) after startup", recovered)
    except Exception as e:
        logger.warning("Parse queue startup setup failed: %s", e)

    _parse_worker_shutdown.clear()
    for worker_number in range(PARSE_JOB_WORKER_CONCURRENCY):
        _parse_worker_tasks.append(asyncio.create_task(_parse_job_worker(worker_number + 1)))

    # Background reminder loop for saved events (24h + 2h heads-up).
    asyncio.create_task(_event_reminder_loop())
    # Background scheduled-broadcast loop.
    asyncio.create_task(_scheduled_broadcast_loop())
    # Warm the OCR engine so the first flyer upload isn't penalised by model load.
    def _warm_ocr():
        try:
            _get_ocr_engine()
        except Exception as e:
            logger.warning("OCR warm-up failed: %s", e)
    asyncio.get_running_loop().run_in_executor(None, _warm_ocr)


@app.on_event("shutdown")
async def shutdown():
    # Stop parser workers cleanly. Any job interrupted mid-flight keeps its
    # persisted sources and will be recovered/re-queued on the next process.
    _parse_worker_shutdown.set()
    for task in list(_parse_worker_tasks):
        task.cancel()
    if _parse_worker_tasks:
        await asyncio.gather(*_parse_worker_tasks, return_exceptions=True)
    _parse_worker_tasks.clear()
    client.close()


app.include_router(api)
app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get("CORS_ORIGINS", "*").split(","),
    allow_methods=["*"],
    allow_headers=["*"],
)
